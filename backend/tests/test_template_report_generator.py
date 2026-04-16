"""Tests for the template report generator."""
import io

import pytest
from docx import Document
from docx.oxml.ns import qn

from core.template_report_generator import generate_from_template


# ── Shared fixtures ───────────────────────────────────────────────────────────

def _make_template(
    sections=None,
    account_grouping=None,
    page_break_after=None,
    negative_format="(X,XXX)",
    currency_format="#,##0",
):
    """Build a realistic template dict."""
    if sections is None:
        sections = [
            {"title": "Independent Auditors' Report", "level": 1,
             "content_type": "narrative", "start_page": 1},
            {"title": "Statement of Financial Position", "level": 1,
             "content_type": "table", "start_page": 4},
            {"title": "Statement of Profit or Loss", "level": 1,
             "content_type": "table", "start_page": 5},
            {"title": "Notes to the Financial Statements", "level": 1,
             "content_type": "narrative", "start_page": 8},
        ]
    return {
        "document_structure": {
            "title": "Financial Statements",
            "date_range": "2024",
            "company_name": "Template Co",
            "auditor_name": "Big4 LLP",
            "pages": 22,
            "sections": sections,
        },
        "account_grouping": account_grouping or {},
        "terminology": {
            "headings_seen": [],
            "common_phrases": [],
            "currency": "AED",
        },
        "formatting_rules": {
            "page_break_after_sections": page_break_after or [],
            "table_formatting": {
                "currency_format": currency_format,
                "negative_number_format": negative_format,
            },
            "font_hierarchy": {
                "heading_1_bold": True,
                "table_header_bold": True,
            },
        },
    }


def _make_current_data(rows=None, draft_content=""):
    return {
        "company_name": "Test Company LLC",
        "location": "Dubai - United Arab Emirates",
        "period_end": "December 31, 2024",
        "opinion_type": "unqualified",
        "draft_content": draft_content or "We have audited the financial statements.",
        "rows": rows or [
            {
                "account_name": "Cash and Bank Balances",
                "section": "Statement of Financial Position",
                "indent_level": 1,
                "net": 500000.0,
                "prior_year": 450000.0,
                "is_total": False,
                "is_subtotal": False,
                "notes_ref": "5",
            },
            {
                "account_name": "Total Assets",
                "section": "Statement of Financial Position",
                "indent_level": 0,
                "net": 1500000.0,
                "prior_year": 1400000.0,
                "is_total": True,
                "is_subtotal": False,
                "notes_ref": None,
            },
        ],
    }


# ═══════════════════════════════════════════════════════════════════════════════
# Tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestGenerateBasicReport:
    """test_generate_basic_report — minimal data + template → valid DOCX."""

    def test_returns_bytes(self):
        result = generate_from_template(_make_current_data(), _make_template())
        assert isinstance(result, bytes)
        assert len(result) > 1000  # valid DOCX is never tiny

    def test_contains_company_name(self):
        result = generate_from_template(_make_current_data(), _make_template())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "TEST COMPANY LLC" in all_text

    def test_contains_period_end(self):
        result = generate_from_template(_make_current_data(), _make_template())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DECEMBER 31, 2024" in all_text

    def test_has_tables(self):
        result = generate_from_template(_make_current_data(), _make_template())
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) > 0


class TestSectionsFollowTemplateOrder:
    """test_sections_follow_template_order — sections appear in template order."""

    def test_section_order_matches_template(self):
        template = _make_template(sections=[
            {"title": "Statement of Financial Position", "level": 1,
             "content_type": "table", "start_page": 1},
            {"title": "Statement of Profit or Loss", "level": 1,
             "content_type": "table", "start_page": 2},
            {"title": "Notes to the Financial Statements", "level": 1,
             "content_type": "narrative", "start_page": 3},
        ])
        result = generate_from_template(_make_current_data(), template)
        doc = Document(io.BytesIO(result))

        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 3

        # Verify ordering
        idx_position = next(i for i, h in enumerate(headings)
                           if "Financial Position" in h)
        idx_profit = next(i for i, h in enumerate(headings)
                         if "Profit or Loss" in h)
        idx_notes = next(i for i, h in enumerate(headings)
                        if "Notes" in h)

        assert idx_position < idx_profit < idx_notes


class TestNegativeNumberFormatting:
    """test_negative_number_formatting — parenthetical format."""

    def test_parenthetical_negatives(self):
        rows = [
            {
                "account_name": "Net Loss",
                "section": "Statement of Profit or Loss",
                "indent_level": 0,
                "net": -250000.0,
                "prior_year": -100000.0,
                "is_total": True,
                "is_subtotal": False,
                "notes_ref": None,
            },
        ]
        template = _make_template(negative_format="(X,XXX)")
        result = generate_from_template(
            _make_current_data(rows=rows), template,
        )
        doc = Document(io.BytesIO(result))

        # Find the table and check for parenthetical negatives
        all_cell_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text)

        assert any("(250,000)" in t for t in all_cell_texts), (
            f"Expected '(250,000)' in table cells, got: {all_cell_texts}"
        )

    def test_dash_negatives(self):
        rows = [
            {
                "account_name": "Net Loss",
                "section": "Statement of Profit or Loss",
                "indent_level": 0,
                "net": -250000.0,
                "prior_year": None,
                "is_total": True,
                "is_subtotal": False,
                "notes_ref": None,
            },
        ]
        template = _make_template(negative_format="-X,XXX")
        result = generate_from_template(
            _make_current_data(rows=rows), template,
        )
        doc = Document(io.BytesIO(result))

        all_cell_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text)

        assert any("-250,000" in t for t in all_cell_texts), (
            f"Expected '-250,000' in table cells, got: {all_cell_texts}"
        )


class TestPriorYearColumn:
    """test_prior_year_column — prior year data in the correct column."""

    def test_prior_year_from_row_data(self):
        result = generate_from_template(_make_current_data(), _make_template())
        doc = Document(io.BytesIO(result))

        # Find a table with "Prior Year" header
        for table in doc.tables:
            header_texts = [cell.text for cell in table.rows[0].cells]
            if any("Prior Year" in h for h in header_texts):
                # Check that prior year column (last) has data
                for row in table.rows[1:]:
                    prior_cell = row.cells[-1].text
                    if prior_cell and prior_cell != "-":
                        assert "450,000" in prior_cell or "1,400,000" in prior_cell
                        return
        pytest.fail("No Prior Year data found in any table")

    def test_prior_year_from_lookup(self):
        rows = [
            {
                "account_name": "Cash and Bank Balances",
                "section": "Statement of Financial Position",
                "indent_level": 1,
                "net": 500000.0,
                "prior_year": None,
                "is_total": False,
                "is_subtotal": False,
                "notes_ref": "5",
            },
        ]
        prior = [{"account_name": "Cash and Bank Balances", "prior_year_value": 320000.0}]
        template = _make_template()
        result = generate_from_template(
            _make_current_data(rows=rows), template, prior_year_rows=prior,
        )
        doc = Document(io.BytesIO(result))

        all_cell_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text)

        assert any("320,000" in t for t in all_cell_texts), (
            f"Expected '320,000' in table cells, got: {all_cell_texts}"
        )


class TestEmptyTemplateFallback:
    """test_empty_template_fallback — empty template → still generates report."""

    def test_none_template(self):
        result = generate_from_template(_make_current_data(), None)
        assert isinstance(result, bytes)
        assert len(result) > 1000

    def test_empty_dict_template(self):
        result = generate_from_template(_make_current_data(), {})
        assert isinstance(result, bytes)
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "TEST COMPANY LLC" in all_text

    def test_fallback_has_default_sections(self):
        result = generate_from_template(_make_current_data(), {})
        doc = Document(io.BytesIO(result))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 2  # at least some headings rendered


class TestPageBreaks:
    """test_page_breaks — sections in page_break_after_sections get page breaks."""

    def test_page_break_after_specified_section(self):
        template = _make_template(
            page_break_after=[
                "Statement of Financial Position",
                "Statement of Profit or Loss",
            ],
        )
        result = generate_from_template(_make_current_data(), template)
        doc = Document(io.BytesIO(result))

        # Page breaks are represented as <w:br w:type="page"/> in runs,
        # or as paragraph-level page-break-before.  python-docx adds them
        # as a paragraph with a page-break run.
        break_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                br_elements = run._element.findall(qn("w:br"))
                for br in br_elements:
                    if br.get(qn("w:type")) == "page":
                        break_count += 1

        # Cover page break + TOC break + at least 2 from our sections
        assert break_count >= 4, (
            f"Expected ≥4 page breaks (cover+TOC+2 sections), found {break_count}"
        )

    def test_no_extra_breaks_without_config(self):
        template = _make_template(page_break_after=[])
        result = generate_from_template(_make_current_data(), template)
        doc = Document(io.BytesIO(result))

        break_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                br_elements = run._element.findall(qn("w:br"))
                for br in br_elements:
                    if br.get(qn("w:type")) == "page":
                        break_count += 1

        # Only cover page + TOC breaks (2 total)
        assert break_count == 2, (
            f"Expected exactly 2 page breaks (cover+TOC), found {break_count}"
        )


def test_formatting_rules_applied_to_docx():
    """When formatting_rules specify #,##0 (no decimals), generated DOCX amounts must have no decimal places."""
    import re
    from core.template_report_generator import generate_from_template

    template = {
        "document_structure": {"title": "Test", "date_range": "2024", "company_name": "Test Co", "auditor_name": "", "pages": 5, "sections": []},
        "account_grouping": {},
        "terminology": {"currency": "AED", "common_phrases": [], "headings_seen": []},
        "formatting_rules": {
            "table_formatting": {"currency_format": "#,##0", "negative_number_format": "(X,XXX)"},
            "font_hierarchy": {"heading_1_bold": True, "table_header_bold": True},
            "page_break_after_sections": [],
        },
    }

    current_data = {
        "company_name": "Test Co LLC",
        "location": "Dubai, UAE",
        "period_end": "31 December 2024",
        "opinion_type": "unqualified",
        "draft_content": "Draft report narrative here.",
        "rows": [
            {"account": "Cash", "mappedTo": "Current Assets", "amount": 50000.75},
        ],
    }

    docx_bytes = generate_from_template(current_data, template)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0

    import io
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"
    assert "50,000.75" not in full_text, "Decimal amount found when currency_format is #,##0"
