"""
Export Converter — converts markdown text to Word, PDF, or Excel.
Used by the chat export endpoint to let users download LLM responses.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


def _strip_inline_md(text: str) -> str:
    """Strip inline markdown markers so plain-text contexts (Excel cells, etc.) are clean."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text


def _add_formatted_runs(paragraph, text: str, default_bold: bool = False) -> None:
    """Parse **bold**, *italic*, and `code` inline markers and add formatted runs to a
    python-docx paragraph. Falls back to plain text for everything else."""
    # Split on the three marker patterns (longest first so ** isn't confused with *)
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            run = paragraph.add_run(part)
        if default_bold:
            run.bold = True


def _parse_markdown_tables(text: str) -> list[list[list[str]]]:
    """Extract all markdown tables from text."""
    tables = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.count("|") >= 2:
            header_cells = [c.strip() for c in line.split("|") if c.strip()]
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    row_line = lines[j].strip()
                    if not row_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in row_line.split("|") if c.strip()]
                    if row_cells:
                        rows.append(row_cells)
                    j += 1
                if len(rows) > 1:
                    tables.append(rows)
                i = j
                continue
        i += 1
    return tables


def to_word(markdown_text: str) -> bytes:
    """Convert markdown text to a Word DOCX file. Returns bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Match heading levels: ^#{1,6}\s
        hdr_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if hdr_match:
            hash_count = len(hdr_match.group(1))
            heading_text = _strip_inline_md(hdr_match.group(2).strip())
            level = min(hash_count, 3)  # python-docx only supports levels 1-3
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        stripped = line.strip()

        if stripped.startswith("|") and stripped.count("|") >= 2:
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                header_cells = [c.strip() for c in stripped.split("|") if c.strip()]
                table_rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    r_line = lines[j].strip()
                    if not r_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in r_line.split("|") if c.strip()]
                    if row_cells:
                        table_rows.append(row_cells)
                    j += 1

                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            # Clear the default empty paragraph and re-add with formatting
                            para = cell.paragraphs[0]
                            para.clear()
                            _add_formatted_runs(para, cell_text, default_bold=(r_idx == 0))

                doc.add_paragraph()
                i = j
                continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(markdown_text: str) -> bytes:
    """Convert markdown text to PDF. Returns bytes."""
    try:
        return _to_pdf_weasyprint(markdown_text)
    except Exception as exc:
        logger.warning(f"weasyprint PDF failed ({exc}), trying reportlab")
        try:
            return _to_pdf_reportlab(markdown_text)
        except Exception as exc2:
            logger.error(f"Both PDF methods failed: {exc2}")
            raise RuntimeError("PDF export unavailable — install weasyprint or reportlab") from exc2


def _to_pdf_weasyprint(markdown_text: str) -> bytes:
    import markdown as md_lib
    from weasyprint import HTML, CSS

    # Use full set of extensions: tables for GFM tables, extra for inline formatting, nl2br for line breaks
    html_body = md_lib.markdown(markdown_text, extensions=["tables", "extra", "nl2br"])
    css = CSS(string="""
        @page { size: A4; margin: 2cm; }
        body { font-family: "Times New Roman", Times, serif; font-size: 11pt; line-height: 1.6; }
        h1 { font-size: 16pt; margin-top: 20pt; }
        h2 { font-size: 13pt; margin-top: 14pt; }
        h3 { font-size: 12pt; margin-top: 10pt; }
        table { border-collapse: collapse; width: 100%; margin: 10pt 0; }
        th, td { border: 1px solid #ccc; padding: 6pt 8pt; text-align: left; }
        th { background: #f0f0f0; font-weight: bold; }
        strong { font-weight: bold; }
        em { font-style: italic; }
    """)
    full_html = f"<html><body>{html_body}</body></html>"
    return HTML(string=full_html).write_pdf(stylesheets=[css])


def _to_pdf_reportlab(markdown_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        
        # Match heading levels: ^#{1,6}\s
        hdr_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if hdr_match:
            hash_count = len(hdr_match.group(1))
            heading_text = _strip_inline_md(hdr_match.group(2))
            if hash_count == 1:
                story.append(Paragraph(heading_text, styles["Heading1"]))
            elif hash_count == 2:
                story.append(Paragraph(heading_text, styles["Heading2"]))
            else:
                story.append(Paragraph(heading_text, styles["Heading3"]))
            continue
        
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = _strip_inline_md(stripped[2:])
            story.append(Paragraph(f"&bull; {bullet_text}", styles["Normal"]))
        elif stripped.startswith("|"):
            # Skip table rows — reportlab can't render markdown tables
            continue
        else:
            clean = _strip_inline_md(stripped)
            story.append(Paragraph(clean, styles["Normal"]))

    doc_rl.build(story)
    return buf.getvalue()


def to_excel(markdown_text: str) -> bytes:
    """Convert markdown tables to Excel. Returns empty bytes if no tables found."""
    tables = _parse_markdown_tables(markdown_text)
    if not tables:
        return b""

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    lines = markdown_text.split("\n")
    table_line_indices = []
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("|") and stripped.count("|") >= 2:
            if idx + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[idx + 1].strip()):
                table_line_indices.append(idx)

    for t_idx, table_rows in enumerate(tables):
        sheet_name = f"Table {t_idx + 1}"
        if t_idx < len(table_line_indices):
            tli = table_line_indices[t_idx]
            for li in range(tli - 1, max(tli - 10, -1), -1):
                if lines[li].strip().startswith("#"):
                    raw = re.sub(r"^#+\s*", "", lines[li].strip())
                    sheet_name = _strip_inline_md(raw)[:31]
                    break

        ws = wb.create_sheet(title=sheet_name)
        header_fill = PatternFill("solid", fgColor="E0E7FF")
        header_font = Font(bold=True)

        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_value in enumerate(row):
                clean_value = _strip_inline_md(cell_value)
                cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=clean_value)
                if r_idx == 0:
                    cell.font = header_font
                    cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True)

                if r_idx > 0:
                    try:
                        numeric = float(clean_value.replace(",", "").replace("%", ""))
                        cell.value = numeric
                        if "%" in clean_value:
                            cell.number_format = "0.00%"
                        else:
                            cell.number_format = "#,##0.00"
                    except (ValueError, AttributeError):
                        pass

        for col in ws.columns:
            max_len = max(
                (len(str(c.value)) for c in col if c.value is not None),
                default=10,
            )
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
