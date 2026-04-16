"""
Audit Report DOCX Formatter.

Takes structured audit data and produces a professional DOCX report with:
  - Cover page (company name, location, year, report title)
  - Table of contents
  - Independent Auditors' Report (narrative from draft)
  - Statement of Financial Position (two-column AED table)
  - Statement of Profit or Loss (two-column AED table)
  - Notes placeholder
"""
import io
import logging
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _add_run(para, text: str, bold: bool = False, size_pt: int = 11,
             color: Optional[tuple] = None, italic: bool = False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _build_cover_page(doc: Document, company_name: str, location: str, period_end: str):
    """Insert cover page with company name, location, report title, and year."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_company, company_name.upper(), bold=True, size_pt=16)

    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_loc, location.upper(), bold=True, size_pt=13)

    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, "FINANCIAL STATEMENTS AND", bold=True, size_pt=13)

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title2, "INDEPENDENT AUDITOR'S REPORT", bold=True, size_pt=13)

    doc.add_paragraph()

    p_period = doc.add_paragraph()
    p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_period, f"FOR THE YEAR ENDED {period_end.upper()}", bold=True, size_pt=13)

    doc.add_page_break()


def _build_toc(doc: Document, sections: list):
    """Insert a simple table of contents."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(heading, "Table of Contents", bold=False, size_pt=12)

    doc.add_paragraph()

    toc_items = [
        ("Independent Auditors' Report", "1 - 3"),
        ("Statement of Financial Position", "4"),
        ("Statement of Profit or Loss and Other Comprehensive Income", "5"),
        ("Statement of Changes in Shareholders' Equity", "6"),
        ("Statement of Cash Flows", "7"),
        ("Notes to the Financial Statements", "8 - 22"),
    ]

    for title, pages in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        run_title = p.add_run(title)
        run_title.font.size = Pt(11)
        run_page = p.add_run(f"\t{pages}")
        run_page.font.size = Pt(11)
        run_page.bold = True

    doc.add_page_break()


def _build_narrative_section(doc: Document, title: str, content: str):
    """Add a narrative section (e.g. Auditors' Report) from Markdown-style text."""
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Strip markdown heading markers and add paragraphs
    content = re.sub(r'^#{1,3}\s+', '', content, flags=re.MULTILINE)
    for para_text in content.split('\n\n'):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        # Bold text wrapped in **
        parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                _add_run(p, part[2:-2], bold=True, size_pt=11)
            else:
                _add_run(p, part, size_pt=11)


def _build_financial_table(doc: Document, title: str, rows: list):
    """
    Build a 4-column financial statement table:
    Account Name | Notes | Current Year (AED) | Prior Year (AED)
    """
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = ''
    hdr[1].text = 'Notes'
    hdr[2].text = 'Current Year\nAED'
    hdr[3].text = 'Prior Year\nAED'

    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    last_category = None
    for row in rows:
        account = row.get('account', '')
        category = row.get('category', '')
        amount = row.get('amount')
        prior = row.get('prior_year')
        is_total = 'total' in account.lower() or row.get('is_total', False)

        # Insert category subheading if category changed
        if category and category != last_category:
            cat_row = table.add_row()
            cat_row.cells[0].text = category
            for p in cat_row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
            last_category = category

        data_row = table.add_row()
        cells = data_row.cells

        cells[0].text = f"  {account}" if not is_total else account
        cells[1].text = str(row.get('notes_ref', ''))
        cells[2].text = f"{amount:,.0f}" if amount is not None else 'Not provided'
        cells[3].text = f"{prior:,.0f}" if prior is not None else 'Not provided'

        for i, cell in enumerate(cells):
            for p in cell.paragraphs:
                for r in p.runs:
                    if is_total:
                        r.bold = True
                    r.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def format_audit_report(report_data: dict) -> bytes:
    """
    Main entry point. Takes report_data dict and returns DOCX as bytes.

    Expected keys:
    - company_name: str
    - location: str (defaults to "Dubai - United Arab Emirates")
    - period_end: str
    - opinion_type: str
    - draft_content: str (Markdown narrative from LLM)
    - rows: list[{account, category, amount, prior_year, notes_ref?, is_total?}]
    """
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    company_name = report_data.get("company_name", "Company")
    location = report_data.get("location", "Dubai - United Arab Emirates")
    period_end = report_data.get("period_end", "")
    draft_content = report_data.get("draft_content", "")
    rows = report_data.get("rows", [])

    _build_cover_page(doc, company_name, location, period_end)
    _build_toc(doc, [])

    # Auditors' Report — extract from draft content
    if draft_content:
        _build_narrative_section(doc, "Independent Auditors' Report", draft_content[:3000])
        doc.add_page_break()

    # Financial Tables — group rows by section
    balance_sheet = [r for r in rows if r.get('category', '').lower() in
                     ('current assets', 'non-current assets', 'current liabilities',
                      'non-current liabilities', 'equity', 'assets', 'liabilities')]
    income_rows = [r for r in rows if r.get('category', '').lower() in
                   ('revenue', 'operating expenses', 'cost of sales', 'other income')]

    if balance_sheet:
        _build_financial_table(doc, "Statement of Financial Position", balance_sheet)
        doc.add_page_break()

    if income_rows:
        _build_financial_table(doc, "Statement of Profit or Loss and Other Comprehensive Income", income_rows)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
