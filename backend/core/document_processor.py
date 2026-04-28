"""
Document Processor – Parses PDF, Word, Excel, and text files into chunks.

Each chunk includes text content and metadata (page number, source file, etc.)
ready for embedding and indexing in the RAG vector store.
"""

import hashlib
import logging
import os
import shutil
from pathlib import Path

from config import settings

logger = logging.getLogger(__name__)


class DocumentChunk:
    """A chunk of text extracted from a document."""

    def __init__(self, text: str, metadata: dict):
        self.text = text
        self.metadata = metadata

    def to_dict(self):
        return {"text": self.text, "metadata": self.metadata}


class DocumentProcessor:
    """Processes documents into text chunks for the RAG pipeline."""

    SUPPORTED_TYPES = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".csv", ".md"}

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        if chunk_overlap >= chunk_size:
            raise ValueError(
                f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size}). "
                "Check CHUNK_OVERLAP and CHUNK_SIZE in your .env."
            )
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._ocr_dependency_warning_emitted = False
        self._ocr_runtime_warning_emitted = False
        self._cv2_dependency_warning_emitted = False
        self._tesseract_cmd_checked = False
        self._tesseract_cmd: str | None = None

    def _resolve_tessdata_dir(self) -> str | None:
        """Resolve a tessdata directory that can include additional OCR languages."""
        backend_dir = Path(__file__).resolve().parent.parent
        configured = settings.pdf_ocr_tessdata_dir.strip()
        candidates = []

        if configured:
            configured_path = Path(configured)
            if not configured_path.is_absolute():
                configured_path = backend_dir / configured_path
            candidates.append(configured_path)

        candidates.append(backend_dir / "tessdata")

        for candidate in candidates:
            if candidate.exists() and candidate.is_dir():
                return str(candidate)

        return None

    def _resolve_tesseract_cmd(self) -> str | None:
        """Resolve tesseract executable path when it's not available on PATH."""
        if self._tesseract_cmd_checked:
            return self._tesseract_cmd

        self._tesseract_cmd_checked = True
        candidates = []

        configured = settings.pdf_ocr_tesseract_cmd.strip()
        if configured:
            configured_path = Path(configured)
            if configured_path.exists():
                candidates.append(configured_path)

        discovered = shutil.which("tesseract")
        if discovered:
            candidates.append(Path(discovered))

        if os.name == "nt":
            candidates.extend([
                Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
                Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
                Path.home() / r"AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
            ])

        for candidate in candidates:
            if candidate.exists() and candidate.is_file():
                self._tesseract_cmd = str(candidate)
                return self._tesseract_cmd

        return None

    def get_file_type(self, filepath: str) -> str:
        """Determine file type from extension."""
        return Path(filepath).suffix.lower()

    def is_supported(self, filepath: str) -> bool:
        """Check if the file type is supported."""
        return self.get_file_type(filepath) in self.SUPPORTED_TYPES

    async def process(self, filepath: str, doc_id: str = "") -> list[DocumentChunk]:
        """
        Process a document file into chunks.

        Args:
            filepath: Path to the document file.
            doc_id: Document ID for metadata tagging.

        Returns:
            List of DocumentChunk objects.
        """
        file_type = self.get_file_type(filepath)
        filename = Path(filepath).name

        logger.info(f"Processing document: {filename} (type: {file_type})")

        # Extract raw text based on file type
        if file_type == ".pdf":
            raw_text = self._extract_pdf(filepath)
        elif file_type == ".docx":
            raw_text = self._extract_docx(filepath)
        elif file_type in {".xlsx", ".xls"}:
            raw_text = self._extract_excel(filepath)
        elif file_type in {".txt", ".md", ".csv"}:
            raw_text = self._extract_text(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Split into chunks
        chunks = self._split_text(raw_text, filename, doc_id)
        logger.info(f"Document '{filename}' split into {len(chunks)} chunks")
        return chunks

    # ── Extractors ────────────────────────────────────────────────

    def _extract_pdf(self, filepath: str) -> list[dict]:
        """Extract text from PDF using PyMuPDF with OCR fallback for scanned pages."""
        import fitz  # PyMuPDF

        filename = Path(filepath).name
        pages = []
        ocr_pages = 0
        doc = fitz.open(filepath)

        try:
            for page_num, page in enumerate(doc, 1):
                extracted_text = page.get_text("text") or ""
                extracted_len = len(extracted_text.strip())
                selected_text = extracted_text

                # For scanned/image-only pages, use OCR and prefer it when it yields richer text.
                if extracted_len < settings.pdf_ocr_min_chars:
                    ocr_text = self._ocr_pdf_page(page, filename, page_num)
                    if len(ocr_text.strip()) > extracted_len:
                        selected_text = ocr_text
                        ocr_pages += 1

                # If neither native text nor OCR yielded enough content, notify the user.
                if len(selected_text.strip()) < 50:
                    selected_text = (
                        f"[Page {page_num}: Could not be read — provide a "
                        "higher-resolution scan or text-based PDF]"
                    )

                pages.append({"text": selected_text, "page": page_num})
        finally:
            doc.close()

        if ocr_pages:
            logger.info(f"OCR fallback used for '{filename}' on {ocr_pages} page(s)")

        return pages

    def _ocr_pdf_page(self, page, filename: str, page_num: int) -> str:
        """Run OCR on a rendered PDF page with enhanced preprocessing pipeline."""
        if not settings.pdf_ocr_enabled:
            return ""

        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image
        except Exception as exc:
            if not self._ocr_dependency_warning_emitted:
                logger.warning(f"PDF OCR fallback unavailable (install pytesseract + Pillow): {exc}")
                self._ocr_dependency_warning_emitted = True
            return ""

        try:
            tesseract_cmd = self._resolve_tesseract_cmd()
            if tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

            tessdata_dir = self._resolve_tessdata_dir()
            if tessdata_dir:
                os.environ["TESSDATA_PREFIX"] = tessdata_dir

            # Render at 300 DPI minimum for reliable OCR on scanned pages.
            dpi = max(settings.pdf_ocr_dpi, 300)
            zoom = dpi / 72.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            mode = "RGB" if pix.n >= 3 else "L"
            image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)

            # ── Advanced preprocessing (cv2 + deskew + skimage) ──────────────
            try:
                import cv2
                import numpy as np
                from deskew import determine_skew
                from skimage.transform import rotate as sk_rotate

                img_array = np.array(image.convert("RGB"))
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

                # Adaptive (Gaussian) thresholding — handles uneven illumination.
                thresh = cv2.adaptiveThreshold(
                    gray, 255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    31, 11,
                )

                # Deskew — correct page rotation.
                angle = determine_skew(thresh)
                if angle is not None and abs(angle) > 0.1:
                    rotated = sk_rotate(thresh, angle, resize=True, preserve_range=True)
                    thresh = (rotated > 127).astype(np.uint8) * 255

                # Noise removal.
                thresh = cv2.medianBlur(thresh, 3)

                # Signature/stamp masking — mask blobs wider or taller than
                # 15 % of the image dimension (unlikely to be regular text).
                h, w = thresh.shape
                num_labels, _, stats, _ = cv2.connectedComponentsWithStats(
                    cv2.bitwise_not(thresh), connectivity=8
                )
                for i in range(1, num_labels):
                    comp_w = stats[i, cv2.CC_STAT_WIDTH]
                    comp_h = stats[i, cv2.CC_STAT_HEIGHT]
                    if comp_w > 0.15 * w or comp_h > 0.15 * h:
                        x = stats[i, cv2.CC_STAT_LEFT]
                        y = stats[i, cv2.CC_STAT_TOP]
                        thresh[y:y + comp_h, x:x + comp_w] = 255  # white out blob

                processed: Image.Image = Image.fromarray(thresh)

            except Exception as cv2_exc:
                if not self._cv2_dependency_warning_emitted:
                    logger.warning(
                        "Advanced OCR preprocessing unavailable "
                        "(install opencv-python-headless + deskew + scikit-image); "
                        f"falling back to basic grayscale pipeline: {cv2_exc}"
                    )
                    self._cv2_dependency_warning_emitted = True
                # Basic fallback: grayscale only.
                processed = image.convert("L")

            return pytesseract.image_to_string(
                processed,
                lang=settings.pdf_ocr_languages,
                config="--oem 1 --psm 6",
            )
        except Exception as exc:
            if not self._ocr_runtime_warning_emitted:
                logger.warning(f"PDF OCR fallback failed: {exc}")
                self._ocr_runtime_warning_emitted = True
            logger.debug(f"OCR failed for {filename} page {page_num}", exc_info=True)
            return ""

    def _extract_docx(self, filepath: str) -> list[dict]:
        """Extract text from Word documents."""
        from docx import Document

        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        return [{"text": "\n".join(full_text), "page": 1}]

    def _extract_excel(self, filepath: str) -> list[dict]:
        """Extract text from Excel spreadsheets."""
        import openpyxl

        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        sheets = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cell_values = [str(v) if v is not None else "" for v in row]
                row_text = " | ".join(cell_values).strip()
                if row_text.replace("|", "").strip():
                    rows.append(row_text)
            if rows:
                sheet_text = f"Sheet: {sheet_name}\n" + "\n".join(rows)
                sheets.append({"text": sheet_text, "page": sheet_name})

        wb.close()
        return sheets

    def _extract_text(self, filepath: str) -> list[dict]:
        """Extract text from plain text / CSV / markdown files."""
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return [{"text": content, "page": 1}]

    # ── Chunking ──────────────────────────────────────────────────

    @staticmethod
    def _extract_heading_from_text(text: str) -> str:
        """Return the most likely section heading from the start of a text block.

        Heuristic: a heading is a line that is:
        - Entirely uppercase, OR
        - Title Case and <= 80 chars and ends without a sentence terminator
        Returns empty string if no heading detected.
        """
        first_line = text.strip().split("\n")[0].strip()
        if not first_line:
            return ""
        # All-caps heading (e.g. "WILLS AND INHERITANCE")
        if first_line.isupper() and 3 <= len(first_line) <= 120:
            return first_line
        # Title Case heading (e.g. "Article 4 – Distribution of Estate")
        words = first_line.split()
        if (
            2 <= len(words) <= 12
            and not first_line[-1] in ".,:;"
            and sum(1 for w in words if w and w[0].isupper()) >= len(words) * 0.6
        ):
            return first_line
        return ""

    def _split_text(
        self,
        pages: list[dict],
        filename: str,
        doc_id: str,
    ) -> list[DocumentChunk]:
        """Split extracted pages/sheets into overlapping chunks."""
        chunks = []
        chunk_index = 0
        current_section = ""

        for page_data in pages:
            text = page_data["text"]
            page_ref = page_data["page"]

            # Update running section from the page's leading text
            page_heading = self._extract_heading_from_text(text)
            if page_heading:
                current_section = page_heading

            # Sliding window chunking
            start = 0
            while start < len(text):
                end = start + self.chunk_size
                chunk_text = text[start:end].strip()

                if chunk_text:
                    chunk_heading = self._extract_heading_from_text(chunk_text)
                    if chunk_heading:
                        current_section = chunk_heading
                    section = chunk_heading if chunk_heading else current_section

                    chunks.append(DocumentChunk(
                        text=chunk_text,
                        metadata={
                            "doc_id": doc_id,
                            "source": filename,
                            "page": str(page_ref),
                            "chunk_index": chunk_index,
                            "section": section,
                            "word_count": len(chunk_text.split()),
                            "total_chunks": 0,  # backfilled below
                        },
                    ))
                    chunk_index += 1

                start += self.chunk_size - self.chunk_overlap
                if start >= len(text):
                    break

        # Second pass: backfill total_chunks with the final count
        total = len(chunks)
        for chunk in chunks:
            chunk.metadata["total_chunks"] = total

        return chunks


# Module-level convenience instance
document_processor = DocumentProcessor()


async def ingest_text(text: str, source: str | None = None, source_type: str = "research", category: str | None = None) -> None:
    """Index a raw text blob into the RAG engine with a source tag."""
    # Local import to avoid a circular dependency (rag_engine imports DocumentChunk from here)
    from core.rag_engine import rag_engine as _rag_engine  # noqa: PLC0415

    doc_id = "research_" + hashlib.md5((source or text[:50]).encode()).hexdigest()[:8]
    meta: dict = {"source": source or "", "source_type": source_type, "page": 1}
    if category:
        meta["category"] = category
    chunk = DocumentChunk(
        text=text[:8000],
        metadata=meta,
    )
    await _rag_engine.ingest_chunks([chunk], doc_id=doc_id)
