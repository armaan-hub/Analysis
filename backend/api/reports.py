"""
Reports API – Triggers financial and tax report generation.
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse, Response
from pydantic import BaseModel
from typing import Any, List, Optional
import io
import json
import os
import re
import uuid
import shutil
import tempfile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc, delete

import logging

logger = logging.getLogger(__name__)

from core.report_generator import report_generator
from core.trial_balance_mapper import map_trial_balance, get_column_suggestions, get_column_suggestions_with_llm
from core.agents.trial_balance_classifier import classify_risks
from core.llm_manager import get_llm_provider
from db.database import get_db
from db.models import GeneratedReport, SavedReport, AuditTemplate
from config import settings

router = APIRouter(prefix="/api/reports", tags=["Reports"])

# ── Report-type → mapper fields (mirrors frontend REPORT_TYPE_CONFIG) ─────────

REPORT_MAPPER_FIELDS: dict[str, list[str]] = {
    "audit": [
        "Current Assets", "Non-Current Assets", "Current Liabilities",
        "Non-Current Liabilities", "Equity", "Revenue", "Cost of Sales",
        "Operating Expenses", "Input VAT", "Output VAT",
        "Cash and Cash Equivalents", "Retained Earnings",
    ],
    "vat": [
        "Standard Rated Sales", "Zero Rated Sales", "Exempt Sales",
        "Output Tax", "Standard Rated Purchases", "Input Tax",
        "Reverse Charge", "Adjustments",
    ],
    "corporate_tax": [
        "Gross Revenue", "Operating Expenses", "Depreciation",
        "Interest Income", "Interest Expense", "Exempt Income",
        "Disallowed Expenses", "Prior Year Losses",
    ],
    "mis": [
        "Revenue", "Cost of Sales", "Gross Profit", "Operating Expenses",
        "EBITDA", "Net Profit", "Cash Position", "Trade Receivables",
    ],
    "ifrs": [
        "Current Assets", "Non-Current Assets", "Current Liabilities",
        "Non-Current Liabilities", "Equity", "Revenue", "Cost of Sales",
        "Operating Expenses", "Finance Costs", "Tax Expense",
        "Cash and Cash Equivalents", "Retained Earnings", "Share Capital",
    ],
    "cash_flow": [
        "Operating Activities", "Investing Activities", "Financing Activities",
        "Opening Balance", "Closing Balance",
    ],
    "budget_vs_actual": [
        "Budgeted Revenue", "Actual Revenue", "Budgeted Expenses",
        "Actual Expenses", "Capital Budget", "Actual Capital",
    ],
    "compliance": [
        "Revenue", "Operating Expenses", "Tax Paid",
        "Regulatory Fees", "Penalties", "Interest on Late Payment",
    ],
    "financial_analysis": [
        "Current Assets", "Current Liabilities", "Total Assets",
        "Total Liabilities", "Revenue", "Net Profit", "Equity", "Operating Cash Flow",
    ],
}

# ── Schemas ───────────────────────────────────────────────────────

class IFRSRequest(BaseModel):
    company_name: str
    revenue: float
    cogs: float
    operating_expenses: float
    assets: float
    liabilities: float

class VatTransaction(BaseModel):
    id: str
    type: str  # "sale" | "purchase"
    amount: float
    tax_rate: float  # e.g., 0.05
    emirate: Optional[str] = None  # e.g. "dubai", "abu_dhabi", "sharjah", etc.

class VATRequest(BaseModel):
    company_name: str
    trn: str
    transactions: List[VatTransaction]

class ReportResponse(BaseModel):
    status: str
    message: str
    download_url: str

# ── Endpoints ─────────────────────────────────────────────────────

@router.get("/list")
async def list_reports(db: AsyncSession = Depends(get_db)):
    """Return all generated reports ordered newest first."""
    result = await db.execute(
        select(GeneratedReport).order_by(desc(GeneratedReport.created_at))
    )
    reports = result.scalars().all()
    return {
        "reports": [
            {
                "id": r.id,
                "report_type": r.report_type,
                "title": r.title,
                "status": r.status,
                "output_path": r.output_path,
                "error_message": r.error_message,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in reports
        ]
    }

@router.get("/download/{filename}")
async def download_report(filename: str):
    """Endpoint to download a generated report."""
    filepath = os.path.join(report_generator.output_dir, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@router.post("/extract-template")
async def extract_template(file: UploadFile = File(...)):
    """Extract plain text from a template file (txt, docx, doc, pdf, xlsx, xls)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    allowed = {"txt", "docx", "doc", "pdf", "xlsx", "xls"}
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        if ext == "txt":
            with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        elif ext in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    line = "\t".join(str(c) for c in row if c is not None)
                    if line.strip():
                        lines.append(line)
            text = "\n".join(lines)
        elif ext in ("docx", "doc"):
            import docx as _docx
            doc = _docx.Document(tmp_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext == "pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(tmp_path)
                text = "\n".join(page.get_text() for page in doc)
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF extraction requires PyMuPDF (pip install pymupdf).")
        else:
            text = ""
        return {"text": text}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {e}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ── Audit Company Docs ────────────────────────────────────────────────────────

@router.post("/extract-company-docs")
async def extract_company_docs(files: list[UploadFile] = File(...)):
    """Extract company name, address, shareholders, activity from uploaded docs."""
    all_text: list[str] = []
    tmp_paths: list[str] = []

    for f in files:
        if not f.filename:
            continue
        ext = f.filename.rsplit(".", 1)[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            shutil.copyfileobj(f.file, tmp)
            tmp_paths.append(tmp.name)

    try:
        for path in tmp_paths:
            ext = path.rsplit(".", 1)[-1].lower()
            try:
                if ext == "txt":
                    with open(path, encoding="utf-8", errors="replace") as fh:
                        all_text.append(fh.read()[:5000])
                elif ext in ("xlsx", "xls"):
                    import openpyxl as _xl
                    wb = _xl.load_workbook(path, read_only=True, data_only=True)
                    lines = []
                    for ws in wb.worksheets:
                        for row in ws.iter_rows(values_only=True):
                            line = " ".join(str(c) for c in row if c is not None)
                            if line.strip():
                                lines.append(line)
                    all_text.append("\n".join(lines)[:5000])
                elif ext in ("docx", "doc"):
                    import docx as _docx
                    doc = _docx.Document(path)
                    all_text.append("\n".join(p.text for p in doc.paragraphs if p.text.strip())[:5000])
                elif ext == "pdf":
                    import fitz
                    doc = fitz.open(path)
                    all_text.append("\n".join(page.get_text() for page in doc)[:5000])
            except Exception:
                pass

        combined = "\n\n---\n\n".join(all_text)
        if not combined.strip():
            return {"company_name": "", "address": "", "shareholders": "", "activity": "",
                    "trade_license_number": "", "registration_number": "", "incorporation_date": ""}

        system_prompt = (
            "You are a UAE corporate document analyst. Extract key information from these business documents. "
            "Return ONLY valid JSON with these exact keys: "
            '{"company_name": "...", "address": "...", "shareholders": "...", "activity": "...", '
            '"trade_license_number": "...", "registration_number": "...", "incorporation_date": "..."}'
            " — trade_license_number is the UAE DED or free-zone trade/commercial license number; "
            "registration_number is the MOE/MOEC company registration number (use trade_license_number value if not separately stated); "
            "incorporation_date is the company incorporation/establishment date in DD/MM/YYYY or YYYY-MM-DD format."
        )
        llm = get_llm_provider(None)
        response = await llm.chat(
            [{"role": "system", "content": system_prompt},
             {"role": "user", "content": f"Documents:\n{combined}\n\nExtract the information."}],
            temperature=0.1,
            max_tokens=500,
        )
        raw = response.content.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        return json.loads(raw.strip())
    except Exception as e:
        return {"company_name": "", "address": "", "shareholders": "", "activity": "",
                "trade_license_number": "", "registration_number": "", "incorporation_date": "", "error": str(e)}
    finally:
        for path in tmp_paths:
            try:
                os.unlink(path)
            except OSError:
                pass


def _has_placeholders(text: str) -> bool:
    """Return True if the text contains any placeholder patterns."""
    if re.search(r'\[.*?\]', text):
        return True
    if re.search(r'INSERT', text, re.IGNORECASE):
        return True
    if re.search(r'\bTBD\b', text, re.IGNORECASE):
        return True
    if '____' in text:
        return True
    if re.search(r'placeholder', text, re.IGNORECASE):
        return True
    return False


# ── Audit Pass 1: Draft Report ────────────────────────────────────────────────

class DraftReportRequest(BaseModel):
    grouped_rows: List[dict]
    company_info: Optional[dict] = None
    company_name: Optional[str] = None
    auditor_name: Optional[str] = None
    period_end: Optional[str] = None
    report_mode: Optional[str] = "standalone"  # standalone | comparative
    prior_year_content: Optional[str] = None   # extracted text from prior year PDF/TB
    ca_questions: Optional[List[dict]] = None  # list of {"id", "question", "risk", "answered"}
    risk_flags: Optional[List[dict]] = None    # list of {"flag", "triggered", "detail"}


@router.post("/draft")
async def generate_draft_report(req: DraftReportRequest):
    """
    Pass 1: Generate a complete, format-agnostic audit draft in Markdown.
    No token limit. No timeout.
    """
    is_comparative = (req.report_mode or "standalone") == "comparative"

    company_block = ""
    if req.company_info and any(req.company_info.values()):
        ci = req.company_info
        company_block = (
            "\n\n**Company Information (from registered documents):**\n"
            f"- Company Name: {ci.get('company_name') or req.company_name or 'N/A'}\n"
            f"- Registered Address: {ci.get('address', 'N/A')}\n"
            f"- Business Activity: {ci.get('activity', 'N/A')}\n"
            f"- Shareholders: {ci.get('shareholders', 'N/A')}\n"
        )

    tb_block = ""
    if req.grouped_rows:
        tb_block = "\n\n**Grouped Trial Balance (Current Year):**\n| Account | Category | Amount (AED) |\n|---------|----------|-------------|\n"
        for r in req.grouped_rows:
            amt = r.get("amount", 0)
            fmt_amt = f"{amt:,.2f}" if isinstance(amt, (int, float)) else str(amt)
            tb_block += f"| {r.get('account', '')} | {r.get('mappedTo', '')} | {fmt_amt} |\n"

    prior_year_block = ""
    if is_comparative and req.prior_year_content:
        prior_year_block = (
            "\n\n**Prior Year Reference (Signed Audit Report / Trial Balance):**\n"
            + req.prior_year_content[:8000]  # cap to avoid token overflow
        )

    comparative_instruction = ""
    if is_comparative:
        comparative_instruction = (
            "IMPORTANT — COMPARATIVE AUDIT: This is a two-period comparative audit per IAS 1.38. "
            "All financial statements MUST include a comparative column for the prior year. "
            "Extract prior year figures from the prior year reference provided below. "
            "Label columns clearly as 'Current Year (AED)' and 'Prior Year (AED)'. "
            "Highlight material variances in the notes and findings sections.\n\n"
        )

    # ── Opinion gate ─────────────────────────────────────────────────────────
    opinion_type = "unqualified"

    # 1. Evidence check: signed/prior-year document required
    has_prior_doc = bool(req.prior_year_content and req.prior_year_content.strip())

    # 2. ISA-critical CA questions must be answered (going_concern, fraud, related_party risks)
    isa_critical_answered = True
    if req.ca_questions:
        isa_critical_answered = all(
            q.get("answered") for q in req.ca_questions
            if q.get("risk") in ("going_concern", "fraud", "related_party")
        )

    # 3. Determine opinion based on risk flags
    if req.risk_flags:
        going_concern_triggered = any(
            f.get("triggered") for f in req.risk_flags
            if f.get("flag") == "going_concern"
        )
        negative_equity_triggered = any(
            f.get("triggered") for f in req.risk_flags
            if f.get("flag") == "negative_equity"
        )
        if going_concern_triggered or negative_equity_triggered:
            opinion_type = "qualified"

    # Do NOT downgrade opinion just because prior year data is missing.
    # Prior year absence = first year audit or user didn't upload — not a scope limitation.
    # Keep whatever opinion was selected (default: unqualified).
    if not has_prior_doc:
        prior_year_note = (
            "Comparative figures not available — "
            "first year of audit or prior year report not uploaded."
        )
    else:
        prior_year_note = ""

    if not isa_critical_answered:
        opinion_type = "disclaimer"
    # ── End opinion gate ──────────────────────────────────────────────────────

    opinion_instructions = {
        "unqualified": (
            "Issue an UNQUALIFIED (clean) audit opinion. State that financial statements present "
            "a true and fair view in all material respects."
        ),
        "qualified": (
            "Issue a QUALIFIED opinion. State that EXCEPT FOR the matters identified in the going "
            "concern or equity section, the financial statements present a true and fair view."
        ),
        "disclaimer": (
            "Issue a DISCLAIMER OF OPINION. State that due to insufficient audit evidence, you are "
            "unable to express an opinion on the financial statements."
        ),
    }

    system_prompt = (
        "You are a senior UAE-qualified auditor preparing a complete audit working paper. "
        + comparative_instruction
        + "Generate a COMPREHENSIVE audit report in Markdown. Include ALL of these sections — "
        "do not omit or summarise any:\n"
        "1. Statement of Financial Position (Balance Sheet)\n"
        "2. Statement of Profit or Loss and Other Comprehensive Income\n"
        "3. Statement of Changes in Equity\n"
        "4. Statement of Cash Flows (IAS 7 indirect method)\n"
        "5. Notes to the Financial Statements (accounting policies, significant estimates, "
        "related party transactions, contingencies, going concern assessment)\n"
        "6. Supporting Schedules (fixed assets, receivables aging, payables aging)\n"
        "7. Audit Findings and Observations\n"
        "8. Management Recommendations\n\n"
        "Use actual figures from the trial balance. Reference IFRS and UAE Federal Law where applicable. "
        "Produce a complete working document — no truncation.\n\n"
        "CRITICAL — COMPLETENESS MANDATE: Every section must be fully written with specific figures, "
        "dates, and names from the trial balance data. Never write placeholder text such as "
        "[INSERT ...], [COMPANY NAME], [DATE], TBD, or ____. If data is genuinely unavailable, "
        "state \"N/A\" explicitly. A report with any placeholder will be rejected."
        + f"\n\nOPINION TYPE: {opinion_type.upper()}. "
        + opinion_instructions[opinion_type]
    )
    user_prompt = (
        f"Company: {req.company_name or 'Unknown'}\n"
        f"Auditor: {req.auditor_name or 'Unknown'}\n"
        f"Period End: {req.period_end or 'Unknown'}\n"
        f"Audit Mode: {'Comparative (IAS 1.38)' if is_comparative else 'Standalone'}"
        + company_block
        + tb_block
        + prior_year_block
        + "\n\nGenerate the complete audit report."
    )

    llm = get_llm_provider(None)
    messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]
    response = await llm.chat(messages, temperature=0.2, max_tokens=None)
    draft_text = response.content

    max_retries = 2
    attempt = 0
    while _has_placeholders(draft_text) and attempt < max_retries:
        attempt += 1
        retry_messages = messages + [
            {"role": "assistant", "content": draft_text},
            {
                "role": "user",
                "content": (
                    "The previous draft contained placeholder text which is not acceptable. "
                    "Rewrite the COMPLETE report with ALL placeholders replaced by actual content "
                    "from the trial balance data. Do not leave any section incomplete."
                ),
            },
        ]
        response = await llm.chat(retry_messages, temperature=0.2, max_tokens=None)
        draft_text = response.content

    if _has_placeholders(draft_text):
        draft_text = (
            "⚠ AUDITOR REVIEW REQUIRED: The following sections may contain incomplete placeholder "
            "text that could not be automatically resolved. Please review and complete before "
            "finalizing.\n\n"
        ) + draft_text

    return {"report_text": draft_text}


# ── Evidence Checklist ────────────────────────────────────────────────────────

class EvidenceChecklistRequest(BaseModel):
    grouped_rows: List[dict]

@router.post("/evidence-checklist")
async def get_evidence_checklist(req: EvidenceChecklistRequest):
    """
    Returns a structured evidence checklist derived from the TB grouped rows.
    Each item maps to an ISA 500-series procedure category.
    """
    ISA_PROCEDURE_MAP = {
        "Current Assets": "ISA 501 — Existence & completeness of assets",
        "Non-Current Assets": "ISA 501 — Physical inspection / third-party confirmation",
        "Current Liabilities": "ISA 505 — External confirmation of balances",
        "Non-Current Liabilities": "ISA 505 — Loan agreement / bank confirmation",
        "Equity": "ISA 550 — Related party review / shareholder register",
        "Revenue": "ISA 240 — Revenue recognition cut-off testing",
        "Cost of Sales": "ISA 500 — Purchase invoice / goods receipt verification",
        "Operating Expenses": "ISA 500 — Invoice & payment voucher inspection",
        "Input VAT": "ISA 500 — FTA VAT return reconciliation",
        "Output VAT": "ISA 500 — FTA VAT return reconciliation",
        "Cash and Cash Equivalents": "ISA 505 — Bank confirmation letter",
        "Retained Earnings": "ISA 500 — Prior year audit report agreement",
    }
    items = []
    for row in req.grouped_rows:
        category = row.get("mappedTo", "")
        procedure = ISA_PROCEDURE_MAP.get(category, "ISA 500 — Substantive testing")
        items.append({
            "id": str(row.get("id", len(items))),
            "account": row.get("account", ""),
            "category": category,
            "amount": row.get("amount", 0),
            "procedure": procedure,
        })
    return {"items": items}


# ── Audit Pass 2: Format Report ───────────────────────────────────────────────

_FORMAT_PROMPTS: dict[str, str] = {
    "big4": (
        "You are a Big 4 audit partner preparing a publication-ready audit report. "
        "Do not deviate from this structure. Do not reorder sections. Do not paraphrase prescribed paragraphs. "
        "Insert the audited entity's data into the correct positions.\n\n"
        "REQUIRED SECTIONS IN ORDER:\n"
        "1. Independent Auditor's Report — use exact ISA 700 prescribed wording for the opinion paragraph\n"
        "2. Basis for Opinion — per ISA 700.28, state compliance with ISAs and ethical requirements\n"
        "3. Key Audit Matters — per ISA 701: for each matter state (a) why it was a KAM and (b) how it was addressed. Identify minimum 3 KAMs from the financial data.\n"
        "4. Management Responsibility Statement — per ISA 700.33\n"
        "5. Auditor's Responsibilities — per ISA 700.38–40\n"
        "6. Going Concern — per ISA 570, include if material uncertainty exists or confirm none identified\n"
        "7. Other Information — per ISA 720\n"
        "8. Financial Statements — all four primary statements per IAS 1 with comparative columns\n"
        "9. Notes to Financial Statements — accounting policies, estimates, related parties (IAS 24), contingencies\n"
        "10. Schedules — Fixed Asset Register, Receivables Aging, Payables Aging\n"
        "11. Management Letter — specific observations, risk ratings, recommendations\n"
        "12. Big 4 Firm Disclosures — independence statement, engagement partner details, PCAOB/IAASB registration\n\n"
        "Reference: ISA 700, ISA 701, ISA 705, ISA 706, ISA 570, IFRS, UAE Federal Decree-Law No. 2 of 2015. "
        "Preserve all financial figures exactly. Do not use placeholders. Do not write [INSERT] or [TBD]."
    ),
    "isa": (
        "You are an ISA-qualified auditor. Produce a complete audit report strictly following ISA 700. "
        "Do not deviate from this structure. Do not reorder sections. Do not paraphrase prescribed paragraphs.\n\n"
        "REQUIRED SECTIONS IN ORDER (all mandatory per ISA 700):\n"
        "1. Title: 'Independent Auditor's Report'\n"
        "2. Addressee\n"
        "3. Auditor's Opinion — use verbatim prescribed language from ISA 700.35: 'In our opinion, the accompanying financial statements present fairly, in all material respects...'\n"
        "4. Basis for Opinion — ISA 700.28: reference ISAs, ethical requirements, and audit evidence obtained\n"
        "5. Key Audit Matters — ISA 701: description and how addressed for each matter\n"
        "6. Going Concern — ISA 570: either material uncertainty paragraph or confirmation none identified\n"
        "7. Other Information — ISA 720\n"
        "8. Responsibilities of Management and Those Charged with Governance — ISA 700.33\n"
        "9. Auditor's Responsibilities for the Audit of Financial Statements — ISA 700.38–40 verbatim\n"
        "10. Report on Other Legal and Regulatory Requirements (UAE Federal Law)\n"
        "11. Engagement partner name, firm name, address, date\n\n"
        "Modified opinions: if evidence was limited or contradictory, apply ISA 705 — use 'Except for' (qualified), "
        "adverse, or disclaimer language as appropriate.\n"
        "Emphasis of matter: apply ISA 706 if a matter is fundamental to user understanding.\n"
        "Reference: ISA 700, 701, 705, 706, 570, 720. Do not use placeholders. Preserve all figures."
    ),
    "fta": (
        "You are a UAE FTA compliance auditor producing a formal FTA audit report. "
        "Do not deviate from this structure. Do not reorder sections. Do not paraphrase regulatory text.\n\n"
        "REQUIRED SECTIONS IN ORDER:\n"
        "1. Cover Page — Engagement title, company name, TRN, period under review, report date\n"
        "2. VAT Compliance Opinion — opinion on compliance with Federal Decree-Law No. 8 of 2017\n"
        "3. VAT 201 Return Analysis — output tax, input tax, net payable/recoverable per Emirate\n"
        "4. Corporate Tax Compliance Opinion — compliance with Federal Decree-Law No. 47 of 2022\n"
        "5. Corporate Tax Computation — taxable income, add-backs, exemptions, 9% calculation\n"
        "6. AML/CFT Findings — compliance with Federal Decree-Law No. 20 of 2018\n"
        "7. FTA Filing Status — all returns filed, due dates, any late submissions\n"
        "8. Penalties and Assessments — any FTA penalties assessed, grounds, amounts\n"
        "9. Recommendations — prioritized remediation actions with timeline\n"
        "10. Auditor's Declaration — signed statement of independence\n\n"
        "Reference specific UAE Federal Decree-Laws and Cabinet Decisions. "
        "Do not use placeholders. Preserve all figures."
    ),
    "internal": (
        "You are an Internal Audit Manager following IIA International Standards. "
        "Produce a formal Internal Audit Report. Do not deviate from this structure.\n\n"
        "REQUIRED SECTIONS IN ORDER:\n"
        "1. Executive Summary — overall audit opinion (Satisfactory/Partially Satisfactory/Unsatisfactory), scope, period\n"
        "2. Audit Objectives and Scope\n"
        "3. Methodology — audit procedures performed\n"
        "4. Detailed Findings — for EACH finding:\n"
        "   a. Finding title\n"
        "   b. Risk Rating: Critical / High / Medium / Low (with justification)\n"
        "   c. Criteria — what should be happening\n"
        "   d. Condition — what was found\n"
        "   e. Cause — root cause analysis\n"
        "   f. Effect — business impact\n"
        "   g. Recommendation\n"
        "5. Management Responses — management's response to each finding\n"
        "6. Management Action Plan — owner, action, target date for each finding\n"
        "7. Summary of Findings by Risk Rating — table\n"
        "8. Appendix — supporting schedules\n\n"
        "Reference IIA International Standards for the Professional Practice of Internal Auditing. "
        "Do not use placeholders. Preserve all figures."
    ),
}


class FormatReportRequest(BaseModel):
    draft: str
    format: str  # big4 | isa | fta | internal | custom
    custom_instructions: Optional[str] = None
    prior_year_template: Optional[str] = None  # extracted text from signed prior year audit report


@router.post("/format")
async def format_report(req: FormatReportRequest):
    """
    Pass 2: Reformat draft to a selected audit standard.
    No token limit. No timeout.
    """
    prior_template_block = ""
    if req.prior_year_template and req.prior_year_template.strip():
        prior_template_block = (
            "\n\n---\n**SIGNED PRIOR YEAR AUDIT REPORT (use as structural/formatting template):**\n"
            + req.prior_year_template[:6000]
            + "\n---\n"
            + "IMPORTANT: Match the section structure, heading style, and layout of the signed prior year report above. "
            "Preserve all current year financial figures from the draft exactly — do not alter numbers.\n\n"
        )

    if req.format == "custom" and req.custom_instructions:
        system_prompt = (
            "You are a senior auditor reformatting a financial report to EXACTLY match a provided template.\n\n"
            "STRICT RULES — follow every one without exception:\n"
            "1. REPLICATE the template's structure exactly: use the same section headings, sub-headings, and order.\n"
            "2. REPLICATE the template's table layout exactly: same column names, same row labels, same indentation.\n"
            "3. REPLICATE the template's formatting: if the template uses bold for totals, bold them; if it uses AED, use AED; if it shows dates in DD/MM/YYYY, use that format.\n"
            "4. DO NOT add any sections, headings, or paragraphs that do not appear in the template.\n"
            "5. DO NOT change any financial figures from the audit draft — only reformat, never recalculate.\n"
            "6. If the template shows a two-column comparative layout (Current Year / Prior Year), reproduce that exact layout.\n"
            "7. Output in Markdown. Use Markdown tables to represent any tabular data in the template.\n\n"
            "HERE IS THE TEMPLATE TO MATCH:\n\n"
            + req.custom_instructions
        )
    else:
        system_prompt = _FORMAT_PROMPTS.get(req.format, _FORMAT_PROMPTS["big4"])

    user_prompt = (
        "Here is the complete audit working paper draft. Reformat and enhance it according to your instructions. "
        "Preserve all financial figures exactly. Add all sections required by your standard.\n\n"
        + req.draft
        + prior_template_block
    )

    llm = get_llm_provider(None)
    response = await llm.chat(
        [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=None,  # No limit
    )
    return {"report_text": response.content}


# ── Generic AI Report Generation (handles all report types) ──────────────────

class GenericReportRequest(BaseModel):
    mapped_data: Optional[List[dict]] = None
    requirements: Optional[dict] = None  # Answers from the Requirements Q&A step
    model_config = {"extra": "allow"}  # Accept any extra fields (company_name, period_end, etc.)


AUDIT_FORMAT_PROMPTS: dict[str, str] = {
    "big4": (
        "Generate a Big 4 firm-style audit report following ISA standards in Markdown. "
        "Include an Independent Auditor's Report, Key Audit Matters section, and Management Letter. "
        "Reference IFRS, ISA 700, and UAE Federal Law No. 2 of 2015."
    ),
    "isa": (
        "Generate an audit report strictly following ISA 700 (Forming an Opinion and Reporting on Financial Statements) "
        "in Markdown. Sections: Auditor's Opinion, Basis for Opinion, Key Audit Matters, "
        "Other Information, Responsibilities of Management, Auditor's Responsibilities."
    ),
    "internal": (
        "Generate an Internal Audit report following IIA standards in Markdown. "
        "Include: Executive Summary, Findings with risk ratings (High/Medium/Low), "
        "Root Cause Analysis, and Management Action Plan."
    ),
    "fta": (
        "Generate a UAE FTA-compliant audit report in Markdown. "
        "Reference Federal Law No. 2 of 2015 and UAE GAAP. "
        "Include: Auditor's Opinion on FTA compliance, VAT/CT findings, and recommendations."
    ),
}

REPORT_TYPE_PROMPTS: dict[str, str] = {
    "ifrs": (
        "You are an IFRS-qualified accountant. Generate IFRS Financial Statements in Markdown. "
        "Include: Statement of Financial Position, Statement of Profit or Loss and Other Comprehensive Income, "
        "Statement of Changes in Equity, Statement of Cash Flows (IAS 7 indirect method), and key accounting policy notes. "
        "Reference applicable IFRS standards (IFRS 15, IAS 16, IAS 36, IAS 38, etc.) where relevant."
    ),
    "vat": (
        "You are a UAE VAT specialist. Generate a UAE VAT 201 Return analysis in Markdown. "
        "Apply Federal Decree-Law No. 8 of 2017. Include: Standard Rated Supplies, Zero Rated Supplies, "
        "Exempt Supplies, Output Tax, Standard Rated Purchases, Input Tax, per-Emirate breakdown (Box 1a–1g), "
        "Net VAT payable or recoverable, and any partial exemption or reverse charge notes."
    ),
    "audit": (
        "You are a senior UAE-qualified auditor. Generate a professional audit report in Markdown. "
        "Include: Audit Opinion, Basis of Opinion, Key Audit Matters, Responsibilities of Management, "
        "Auditor's Responsibilities, and a Management Letter summary. "
        "Reference IFRS and UAE Federal Law No. 2 of 2015 where appropriate."
    ),
    "corporate_tax": (
        "You are a UAE Corporate Tax specialist. Generate a detailed UAE Corporate Tax computation in Markdown. "
        "Apply Federal Decree-Law No. 47 of 2022. Include: Accounting Profit, Add-backs, "
        "Exempt Income, Taxable Income, Small Business Relief eligibility check (AED 3M threshold), "
        "9% tax rate calculation, and key notes on applicable exemptions."
    ),
    "compliance": (
        "You are a UAE regulatory compliance expert. Generate a comprehensive compliance report in Markdown. "
        "Cover: VAT compliance, Corporate Tax registration, AML/CFT obligations, "
        "regulatory findings (if any), risk ratings (High/Medium/Low), and recommended remediation actions. "
        "Reference FTA regulations and UAE Federal laws."
    ),
    "mis": (
        "You are a CFO-level management accountant. Generate a Management Information System (MIS) report in Markdown. "
        "Include: Executive Summary, Revenue vs Budget, Gross Margin Analysis, EBITDA, "
        "Key Performance Indicators (KPIs), Trend Analysis, and Management Commentary. "
        "Use tables where appropriate."
    ),
    "financial_analysis": (
        "You are a financial analyst. Generate a detailed financial analysis report in Markdown. "
        "Include: Liquidity Ratios (Current Ratio, Quick Ratio), Profitability Ratios (Gross Margin, Net Margin, ROE, ROA), "
        "Leverage Ratios (Debt-to-Equity, Interest Coverage), Efficiency Ratios, "
        "Trend Analysis, Benchmarking commentary, and an overall assessment."
    ),
    "budget_vs_actual": (
        "You are a management accountant. Generate a Budget vs Actual Variance Analysis report in Markdown. "
        "Include: Revenue Variance (Favorable/Unfavorable), Expense Variance analysis, "
        "Overall Budget Performance, Variance explanations, and Recommendations for corrective action. "
        "Use tables to show Budget, Actual, Variance (AED and %) columns."
    ),
    "cash_flow": (
        "You are an IFRS-qualified accountant. Generate a Cash Flow Statement in Markdown following IFRS IAS 7. "
        "Include: Operating Activities (Indirect Method), Investing Activities, Financing Activities, "
        "Net Change in Cash, Opening and Closing Cash Balances, and Notes on significant non-cash transactions."
    ),
    "custom": (
        "You are a professional financial analyst and accountant. Generate a comprehensive financial report in Markdown "
        "based on the provided data and any specific instructions given. Apply relevant UAE accounting standards, "
        "VAT laws, and IFRS where applicable. Structure the report clearly with sections and tables."
    ),
}


# ── LLM Column Mapping Suggestions ──────────────────────────────────────────

class SuggestMappingsRequest(BaseModel):
    columns: list[str]
    report_type: str
    system_fields: list[str]


@router.post("/suggest-mappings")
async def suggest_column_mappings(req: SuggestMappingsRequest):
    """
    Use LLM to map raw column names to report system fields.
    Returns suggested mappings and plain-English questions for ambiguous columns.
    """
    system_prompt = (
        "You are a financial data mapping assistant. Given a list of raw spreadsheet column names "
        "and a list of target system fields, map each raw column to the most appropriate system field. "
        "For columns you cannot confidently map, provide a simple plain-English question (no jargon) "
        "with 3-4 short answer options so a non-finance user can clarify.\n\n"
        "Respond ONLY with valid JSON in this exact format:\n"
        "{\n"
        '  "mappings": [\n'
        '    {"column": "raw col name", "field": "System Field Name", "confidence": 95},\n'
        '    {"column": "ambiguous col", "field": null, "confidence": 0, '
        '"question": "What does X represent?", "options": ["Option A", "Option B", "Option C"]}\n'
        "  ]\n"
        "}"
    )

    user_prompt = (
        f"Report type: {req.report_type}\n\n"
        f"Raw columns to map:\n{chr(10).join(f'- {c}' for c in req.columns)}\n\n"
        f"Available system fields:\n{chr(10).join(f'- {f}' for f in req.system_fields)}\n\n"
        "Map each raw column to the best matching system field, or ask a clarifying question if unsure."
    )

    try:
        llm = get_llm_provider(None)
        response = await llm.chat(
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            max_tokens=1500,
        )
        raw = response.content.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw.strip())
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM mapping failed: {e}")


@router.post("/generate/{report_type}")
async def generate_ai_report(report_type: str, req: GenericReportRequest):
    """
    AI-powered report generation for all report types.
    Routes through AGENT_REGISTRY when available; falls back to inline prompt selection.
    """
    # Normalize aliases
    if report_type in ("corptax",):
        report_type = "corporate_tax"

    extra_fields = {k: v for k, v in (req.model_extra or {}).items() if k not in ("mapped_data", "requirements")}

    # Try the agent registry first
    from core.agents.registry import get_agent
    agent = get_agent(report_type)

    if agent is not None and report_type != "audit":
        # Non-audit agents: use registry generate()
        try:
            report_text = await agent.generate(
                tb_data=req.mapped_data or [],
                answers={},
                report_type=report_type,
                extra_fields=extra_fields,
                requirements=req.requirements or {},
            )
            return {"report_text": report_text, "report_type": report_type}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI report generation failed: {e}")

    # Audit: handle opinion gate + format selection via AuditAgent
    if report_type == "audit":
        auditor_format = extra_fields.pop("auditor_format", "big4")
        custom_instructions = extra_fields.pop("custom_format_instructions", None)
        opinion = extra_fields.pop("opinion", "unqualified")
        disclaimer_text = extra_fields.pop("disclaimer_text", "")
        company_info = extra_fields.pop("company_info", None)
        prior_year_content = extra_fields.pop("prior_year_content", "")

        if auditor_format == "custom" and custom_instructions:
            # Custom format: bypass agent, use inline custom prompt
            system_prompt = (
                "You are a senior auditor. Generate an audit report in Markdown. "
                f"Follow this custom format template as a guide:\n\n{custom_instructions}"
            )
        else:
            try:
                audit_agent = get_agent("audit")
                report_text = await audit_agent.generate(
                    tb_data=req.mapped_data or [],
                    answers={},
                    opinion=opinion,
                    disclaimer_text=disclaimer_text,
                    report_fields=extra_fields,
                    company_info=company_info,
                    audit_format=auditor_format,
                    prior_year_content=prior_year_content,
                )
                return {"report_text": report_text, "report_type": report_type}
            except ValueError as e:
                raise HTTPException(status_code=422, detail=str(e))
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"AI report generation failed: {e}")

        # Custom format fallback path (inline prompt)
        context_lines = []
        for k, v in extra_fields.items():
            label = k.replace("_", " ").title()
            context_lines.append(f"**{label}**: {v}")
        if req.mapped_data:
            context_lines.append("\n**Mapped Financial Data:**")
            context_lines.append("| Account | Mapped To | Amount (AED) |")
            context_lines.append("|---------|-----------|-------------|")
            for row in req.mapped_data:
                account = row.get("account", row.get("account_name", ""))
                mapped_to = row.get("mapped_to", row.get("mappedTo", row.get("category", "")))
                amount = row.get("amount", row.get("net", row.get("debit", 0)))
                context_lines.append(f"| {account} | {mapped_to} | {amount:,.2f} |" if isinstance(amount, (int, float)) else f"| {account} | {mapped_to} | {amount} |")
        if req.requirements:
            context_lines.append("\n**Report Requirements (user-specified):**")
            for q, a in req.requirements.items():
                context_lines.append(f"- {q}: {a}")
        user_prompt = "Generate the Audit report for the following data:\n\n" + "\n".join(context_lines)
        try:
            llm = get_llm_provider(None)
            response = await llm.chat(
                [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                temperature=0.3, max_tokens=settings.max_tokens,
            )
            return {"report_text": response.content, "report_type": report_type}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI report generation failed: {e}")

    # Unknown report type: generic fallback
    system_prompt = REPORT_TYPE_PROMPTS.get(
        report_type,
        "You are a professional financial and accounting expert. Generate a detailed, "
        "well-structured report in Markdown based on the provided financial data."
    )
    context_lines = []
    for k, v in extra_fields.items():
        label = k.replace("_", " ").title()
        context_lines.append(f"**{label}**: {v}")
    if req.mapped_data:
        context_lines.append("\n**Mapped Financial Data:**")
        context_lines.append("| Account | Mapped To | Amount (AED) |")
        context_lines.append("|---------|-----------|-------------|")
        for row in req.mapped_data:
            account = row.get("account", row.get("account_name", ""))
            mapped_to = row.get("mapped_to", row.get("mappedTo", row.get("category", "")))
            amount = row.get("amount", row.get("net", row.get("debit", 0)))
            context_lines.append(f"| {account} | {mapped_to} | {amount:,.2f} |" if isinstance(amount, (int, float)) else f"| {account} | {mapped_to} | {amount} |")
    if req.requirements:
        context_lines.append("\n**Report Requirements (user-specified):**")
        for q, a in req.requirements.items():
            context_lines.append(f"- {q}: {a}")
    user_prompt = (
        f"Generate the {report_type.replace('_', ' ').title()} report for the following data:\n\n"
        + "\n".join(context_lines)
    )
    try:
        llm = get_llm_provider(None)
        response = await llm.chat(
            [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=0.3, max_tokens=settings.max_tokens,
        )
        return {"report_text": response.content, "report_type": report_type}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI report generation failed: {e}")


AUDIT_GROUPING_CATEGORIES = [
    "Current Assets", "Non-Current Assets", "Current Liabilities",
    "Non-Current Liabilities", "Equity", "Revenue", "Cost of Sales",
    "Operating Expenses", "Input VAT", "Output VAT",
    "Cash and Cash Equivalents", "Retained Earnings",
]


async def _extract_and_group_pdf_tb(pdf_text: str) -> list[dict]:
    """
    Use LLM to extract trial balance / financial statement data from raw PDF text
    and classify into audit categories in one pass.
    """
    AUDIT_CATEGORIES = ", ".join(AUDIT_GROUPING_CATEGORIES)

    system_prompt = (
        "You are a senior UAE-qualified auditor analyzing a financial document. "
        "Extract ALL financial line items (account names and their amounts) from the provided text, "
        "then classify each into one of these audit categories:\n"
        f"{AUDIT_CATEGORIES}\n\n"
        "Rules:\n"
        "- Extract every account/line item that has a monetary amount\n"
        "- Use the net/balance amount (positive for assets/expenses, can be negative for liabilities/income)\n"
        "- If amounts appear in thousands, use as-is (don't multiply)\n"
        "- For P&L items: Revenue is positive, Expenses are positive costs\n"
        "- For Balance Sheet: Assets positive, Liabilities/Equity positive\n"
        "- Ignore totals/subtotals rows, only extract individual account lines\n\n"
        "Return ONLY valid JSON (no markdown fences, no explanation):\n"
        '{"rows": [{"account": "Account Name", "mappedTo": "Category", "amount": 0.00}]}\n'
        "Include ALL line items found. If no financial data found, return {\"rows\": []}."
    )

    # Cap text to avoid token overflow — take first 12000 chars
    text_sample = pdf_text[:12000].strip()

    user_prompt = (
        f"Financial document text:\n\n{text_sample}\n\n"
        "Extract all financial line items and classify them."
    )

    llm = get_llm_provider(None)
    response = await llm.chat(
        [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
        max_tokens=None,
    )
    raw = response.content.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    try:
        result = json.loads(raw.strip())
        rows = result.get("rows", [])
        # Validate and clean rows
        clean_rows = []
        for r in rows:
            account = str(r.get("account", "")).strip()
            mapped_to = str(r.get("mappedTo", "")).strip()
            amount = r.get("amount", 0)
            if not account or not mapped_to:
                continue
            if mapped_to not in AUDIT_GROUPING_CATEGORIES:
                mapped_to = "Operating Expenses"  # safe fallback
            try:
                amount = float(amount)
            except (TypeError, ValueError):
                amount = 0.0
            clean_rows.append({"account": account, "mappedTo": mapped_to, "amount": amount})
        return clean_rows
    except json.JSONDecodeError:
        # If JSON parsing fails, return empty (will be handled gracefully in frontend)
        return []


async def _group_tb_with_llm(rows: list[dict]) -> list[dict]:
    """
    Classify every trial balance row into an audit category.

    Hits the account-mapping cache first so known accounts never reach
    the LLM.  Unknown accounts are classified by the LLM using a stable
    ID-based protocol (the LLM returns IDs, not echoed account strings)
    and then saved back to the cache for future uploads.
    """
    from core.agents.account_cache import lookup_many, save_many

    valid_rows = [
        r for r in rows
        if (r.get('account_name') or r.get('account_code', '')).strip()
        and r.get('account_name', '') not in ('', 'nan', 'None', 'Total', 'TOTAL', 'Grand Total')
    ]
    if not valid_rows:
        valid_rows = rows

    # Attach stable indices + extract display name / amount
    indexed: list[tuple[int, str, float]] = [
        (
            i,
            (r.get('account_name') or r.get('account_code', 'Unknown')).strip(),
            float(r.get('net', 0)),
        )
        for i, r in enumerate(valid_rows)
    ]

    # ── Cache lookup ──────────────────────────────────────────────
    cached = await lookup_many([name for _, name, _ in indexed])

    result_map: dict[int, dict] = {}
    unknowns: list[tuple[int, str, float]] = []  # (original_idx, name, amount)

    for idx, name, amount in indexed:
        mapped_to = cached.get(name.strip().lower())
        if mapped_to:
            result_map[idx] = {"account": name, "mappedTo": mapped_to, "amount": amount}
        else:
            unknowns.append((idx, name, amount))

    # ── LLM for unknowns only ─────────────────────────────────────
    if unknowns:
        llm_results = await _classify_unknowns_with_llm(unknowns)
        await save_many(
            [(name, llm_results[idx]["mappedTo"]) for idx, name, _ in unknowns if idx in llm_results],
            source="llm",
        )
        result_map.update(llm_results)

    return [result_map[i] for i in range(len(valid_rows)) if i in result_map]


async def _classify_unknowns_with_llm(
    unknowns: list[tuple[int, str, float]],
) -> dict[int, dict]:
    """
    Classify a list of (original_idx, account_name, amount) triples via LLM.

    Uses an ID-based request/response protocol so the LLM never needs to
    echo account text back — the original names are always taken from the
    input, eliminating cache-key mismatches from paraphrasing.

    Chunks automatically at 300 rows to stay within token limits.
    Returns ``{original_idx: {account, mappedTo, amount}}``.
    """
    CHUNK_SIZE = 300
    if len(unknowns) > CHUNK_SIZE:
        result: dict[int, dict] = {}
        for start in range(0, len(unknowns), CHUNK_SIZE):
            result.update(await _classify_unknowns_with_llm(unknowns[start:start + CHUNK_SIZE]))
        return result

    system_prompt = (
        "You are a senior UAE-qualified auditor classifying a trial balance. "
        "Classify EVERY account into exactly one of these categories:\n"
        + ", ".join(AUDIT_GROUPING_CATEGORIES)
        + "\n\nEach input row has a numeric ID. "
        "Return ONLY valid JSON (no markdown fences, no explanation):\n"
        '{"rows": [{"id": <number>, "mappedTo": "<Category>"}]}\n'
        "Include a result for EVERY ID. Do not skip any row."
    )
    lines = "\n".join(
        f"{idx} | {name} | {amount:.2f}"
        for idx, name, amount in unknowns
    )
    user_prompt = (
        f"Classify these {len(unknowns)} trial balance accounts:\n"
        "ID | Account Name | Amount\n"
        "---|--------------|-------\n"
        + lines
    )

    llm = get_llm_provider(None)
    response = await llm.chat(
        [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
        max_tokens=None,
    )
    raw = response.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    try:
        id_to_mapped: dict[int, str] = {}
        for r in json.loads(raw.strip()).get("rows", []):
            try:
                row_id = int(r["id"])
                mapped_to = str(r.get("mappedTo", "Operating Expenses")).strip()
                if mapped_to not in AUDIT_GROUPING_CATEGORIES:
                    mapped_to = "Operating Expenses"
                id_to_mapped[row_id] = mapped_to
            except (KeyError, TypeError, ValueError):
                continue
    except json.JSONDecodeError:
        id_to_mapped = {}

    # Reconstruct from original row data — never trust LLM-echoed account text
    return {
        idx: {
            "account": name,
            "mappedTo": id_to_mapped.get(idx, "Operating Expenses"),
            "amount": amount,
        }
        for idx, name, amount in unknowns
    }


# ── Trial Balance Upload ──────────────────────────────────────────────────────

@router.post("/upload-trial-balance")
async def upload_trial_balance(
    file: UploadFile = File(...),
    report_type: Optional[str] = Form(None),
    report_mode: Optional[str] = Form(None),
    prior_year_file: Optional[UploadFile] = File(None),
):
    """
    Upload an Excel/CSV Trial Balance.
    Returns detected column headers + auto-suggested field mappings for the UI.
    Optionally runs LLM fallback mapping when report_type is supplied.
    For audit comparative mode, also accepts and extracts the prior year file.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("xlsx", "xls", "csv", "pdf"):
        raise HTTPException(status_code=400, detail="Only .xlsx, .xls, .csv, or .pdf files are accepted.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    # Extract prior year content if provided (PDF, Excel, or CSV)
    prior_year_content = ""
    prior_tmp_path = None
    if prior_year_file and prior_year_file.filename:
        prior_ext = prior_year_file.filename.rsplit(".", 1)[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{prior_ext}") as ptmp:
            shutil.copyfileobj(prior_year_file.file, ptmp)
            prior_tmp_path = ptmp.name
        try:
            if prior_ext == "pdf":
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(prior_tmp_path)
                    prior_year_content = "\n".join(page.get_text() for page in doc)
                    doc.close()
                except ImportError:
                    prior_year_content = ""
            elif prior_ext in ("xlsx", "xls", "csv"):
                from core.document_processor import map_trial_balance as _map_tb
                try:
                    py_rows = _map_tb(prior_tmp_path)
                    lines = ["| Account | Debit | Credit | Net |", "|---------|-------|--------|-----|"]
                    for r in py_rows[:200]:
                        lines.append(f"| {r.get('account_name','')} | {r.get('debit',0):,.2f} | {r.get('credit',0):,.2f} | {r.get('net',0):,.2f} |")
                    prior_year_content = "\n".join(lines)
                except Exception:
                    prior_year_content = ""
        except Exception:
            prior_year_content = ""
        finally:
            if prior_tmp_path:
                try:
                    os.unlink(prior_tmp_path)
                except OSError:
                    pass

    try:
        mapper_fields = REPORT_MAPPER_FIELDS.get(report_type or "", [])

        # PDF path: extract text and use LLM to parse into trial balance rows
        if ext == "pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(tmp_path)
                pdf_text = "\n".join(page.get_text() for page in doc)
                doc.close()
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF extraction requires PyMuPDF. Install with: pip install pymupdf")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to read PDF: {e}")

            # Use LLM to extract trial balance data from PDF text
            grouped = await _extract_and_group_pdf_tb(pdf_text)
            grouped = classify_risks(grouped)
            return {
                "filename": file.filename,
                "columns": [],
                "suggestions": {},
                "llm_suggestions": {},
                "llm_questions": [],
                "grouped_rows": grouped,
                "prior_year_content": prior_year_content,
            }

        # Audit type: parse all rows + LLM group them (no token limit, no timeout)
        if report_type == "audit":
            rows = map_trial_balance(tmp_path)
            grouped = await _group_tb_with_llm(rows)
            grouped = classify_risks(grouped)
            return {
                "filename": file.filename,
                "columns": [],
                "suggestions": {},
                "llm_suggestions": {},
                "llm_questions": [],
                "grouped_rows": grouped,
                "prior_year_content": prior_year_content,
            }

        # All other types: existing column suggestion flow
        if report_type and mapper_fields:
            result = await get_column_suggestions_with_llm(tmp_path, report_type, mapper_fields)
        else:
            basic = get_column_suggestions(tmp_path)
            result = {**basic, "llm_suggestions": {}, "llm_questions": []}

        return {
            "filename": file.filename,
            "columns": result["columns"],
            "suggestions": result["suggestions"],
            "llm_suggestions": result.get("llm_suggestions", {}),
            "llm_questions": result.get("llm_questions", []),
        }
    except ValueError as e:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Server error processing file: {str(e)}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@router.post("/map-trial-balance")
async def map_trial_balance_endpoint(file: UploadFile = File(...)):
    """
    Upload an Excel/CSV Trial Balance and return fully normalised rows.
    Output: [{ account_code, account_name, category, debit, credit, net }]
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("xlsx", "xls", "csv"):
        raise HTTPException(status_code=400, detail="Only .xlsx, .xls, or .csv files are accepted.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        rows = map_trial_balance(tmp_path)
        totals = {
            "total_debit": round(sum(r["debit"] for r in rows), 2),
            "total_credit": round(sum(r["credit"] for r in rows), 2),
            "net": round(sum(r["net"] for r in rows), 2),
            "is_balanced": abs(sum(r["net"] for r in rows)) < 0.01,
        }
        return {"rows": rows, "totals": totals, "row_count": len(rows)}
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Mapping failed: {str(e)}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass  # Windows file-lock; GC will release it


# ── Report Requirements Q&A ──────────────────────────────────────────────────

class ReportQuestionsRequest(BaseModel):
    mapped_data: Optional[List[dict]] = None
    extra_fields: Optional[dict] = None


# Hardcoded law-driven questions for VAT and Corporate Tax
_VAT_QUESTIONS = [
    {"id": "format", "question": "Which report format do you need?", "type": "choice",
     "options": ["Standard FTA VAT201", "Custom Format"]},
    {"id": "reverse_charge", "question": "Are there any reverse charge transactions?", "type": "boolean",
     "options": ["Yes", "No"]},
    {"id": "partial_exemption", "question": "Is partial exemption applicable?", "type": "boolean",
     "options": ["Yes", "No"]},
    {"id": "tax_period", "question": "Which tax period does this return cover?", "type": "text"},
]

_CT_QUESTIONS = [
    {"id": "sbr", "question": "Should Small Business Relief be applied? (threshold: AED 3M revenue)", "type": "boolean",
     "options": ["Yes", "No"]},
    {"id": "exempt_income", "question": "Is there any exempt income to declare?", "type": "boolean",
     "options": ["Yes", "No"]},
    {"id": "free_zone", "question": "Is this a Free Zone entity?", "type": "boolean",
     "options": ["Yes", "No"]},
    {"id": "transfer_pricing", "question": "Is transfer pricing applicable?", "type": "boolean",
     "options": ["Yes", "No"]},
]


@router.post("/questions/{report_type}")
async def get_report_questions(report_type: str, req: ReportQuestionsRequest):
    """
    Return dynamic questions for the Requirements step.
    VAT and Corporate Tax: hardcoded law-driven questions.
    All others: LLM-generated data-driven questions.
    """
    if report_type in ("corptax",):
        report_type = "corporate_tax"

    if report_type == "vat":
        return {"questions": _VAT_QUESTIONS}

    if report_type == "corporate_tax":
        return {"questions": _CT_QUESTIONS}

    # For all other report types: LLM generates 3-7 data-driven questions
    data_summary = ""
    if req.mapped_data:
        field_names = list({r.get("mapped_to", r.get("mappedTo", "")) for r in req.mapped_data if r.get("mapped_to") or r.get("mappedTo")})
        data_summary = f"Mapped fields present: {', '.join(field_names[:15])}"

    prompt = (
        f"You are generating a requirements Q&A for a {report_type.replace('_', ' ')} report.\n"
        f"Data context: {data_summary or 'No data yet'}\n\n"
        "Generate 3 to 7 relevant, practical questions to understand what the user wants in this report. "
        "Use plain English. For each question choose the best type: 'choice' (multiple select), 'boolean' (yes/no), or 'text' (free).\n"
        "Respond ONLY with valid JSON:\n"
        '{"questions": [{"id": "q1", "question": "...", "type": "choice", "options": ["opt1", "opt2"]}, ...]}'
    )

    try:
        llm = get_llm_provider(None)
        response = await llm.chat(
            [
                {"role": "system", "content": "You are a financial report requirements analyst. Respond with valid JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=800,
        )
        raw = response.content.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw.strip())
        return {"questions": result.get("questions", [])}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate questions: {e}")


# ── VAT Precision Calculations (Phase 2.3) ───────────────────────────────────

class PartialExemptionRequest(BaseModel):
    taxable_supplies: float
    exempt_supplies: float
    total_input_tax: float


class DeregistrationCheckRequest(BaseModel):
    annual_taxable_supplies: float


def calculate_partial_exemption(
    taxable_supplies: float,
    exempt_supplies: float,
    total_input_tax: float,
) -> dict:
    """
    Partial Exemption Ratio = Taxable / (Taxable + Exempt).
    Per FTA Federal Decree-Law No. 8 of 2017, Article 54.
    """
    total = taxable_supplies + exempt_supplies
    if total == 0:
        return {
            "partial_exemption_ratio": 0.0,
            "recoverable_input_tax": 0.0,
            "irrecoverable_input_tax": 0.0,
            "note": "No supplies — ratio cannot be calculated.",
        }
    per = taxable_supplies / total
    return {
        "partial_exemption_ratio": round(per, 4),
        "partial_exemption_ratio_pct": round(per * 100, 2),
        "recoverable_input_tax": round(total_input_tax * per, 2),
        "irrecoverable_input_tax": round(total_input_tax * (1 - per), 2),
        "reference": "Federal Decree-Law No. 8 of 2017, Article 54",
    }


def check_deregistration_threshold(annual_taxable_supplies: float) -> dict:
    """
    Returns True if the business may apply for VAT de-registration.
    Threshold: AED 187,500 (FTA Article 21).
    """
    threshold = 187_500.0
    eligible = annual_taxable_supplies < threshold
    return {
        "annual_taxable_supplies_aed": annual_taxable_supplies,
        "deregistration_threshold_aed": threshold,
        "eligible_for_deregistration": eligible,
        "message": (
            "Annual taxable supplies are below the AED 187,500 de-registration threshold. "
            "The business may apply to de-register from VAT with the FTA."
            if eligible
            else
            f"Annual taxable supplies (AED {annual_taxable_supplies:,.2f}) exceed the "
            "AED 187,500 threshold. The business must remain VAT-registered."
        ),
        "reference": "Federal Decree-Law No. 8 of 2017, Article 21",
    }


@router.post("/vat/partial-exemption")
async def vat_partial_exemption(req: PartialExemptionRequest):
    """Calculate VAT partial exemption ratio per FTA Article 54."""
    return calculate_partial_exemption(
        req.taxable_supplies, req.exempt_supplies, req.total_input_tax
    )


@router.post("/vat/deregistration-check")
async def vat_deregistration_check(req: DeregistrationCheckRequest):
    """Check whether a business qualifies for VAT de-registration per FTA Article 21."""
    return check_deregistration_threshold(req.annual_taxable_supplies)


# ── Saved Reports CRUD ────────────────────────────────────────────────────────

class SaveReportRequest(BaseModel):
    company_name: str = ""
    report_type: str = "audit"
    format: str = "big4"
    period_end_date: Optional[str] = None
    status: str = "draft"
    wizard_state_json: Optional[dict] = None
    draft_content: Optional[str] = None
    final_content: Optional[str] = None


@router.get("/saved")
async def list_saved_reports(db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(SavedReport).order_by(SavedReport.updated_at.desc()))
    reports = result.scalars().all()
    return [
        {
            "id": r.id,
            "company_name": r.company_name,
            "report_type": r.report_type,
            "format": r.format,
            "period_end_date": r.period_end_date,
            "status": r.status,
            "created_at": r.created_at.isoformat() if r.created_at else None,
            "updated_at": r.updated_at.isoformat() if r.updated_at else None,
        }
        for r in reports
    ]


@router.get("/saved/{report_id}/wizard-state")
async def get_report_wizard_state(report_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(SavedReport).where(SavedReport.id == report_id))
    report = result.scalar_one_or_none()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return {
        "wizard_state": report.wizard_state_json,
        "has_wizard_state": report.wizard_state_json is not None,
        "draft_content": report.draft_content,
        "company_name": report.company_name,
        "report_type": report.report_type,
    }


@router.get("/saved/{report_id}")
async def get_saved_report(report_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(SavedReport).where(SavedReport.id == report_id))
    report = result.scalar_one_or_none()
    if not report:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Report not found")
    return {
        "id": report.id,
        "company_name": report.company_name,
        "report_type": report.report_type,
        "format": report.format,
        "period_end_date": report.period_end_date,
        "status": report.status,
        "draft_content": report.draft_content,
        "final_content": report.final_content,
        "wizard_state_json": report.wizard_state_json,
        "created_at": report.created_at.isoformat() if report.created_at else None,
        "updated_at": report.updated_at.isoformat() if report.updated_at else None,
    }


@router.post("/saved")
async def save_report(req: SaveReportRequest, db: AsyncSession = Depends(get_db)):
    report = SavedReport(
        company_name=req.company_name,
        report_type=req.report_type,
        format=req.format,
        period_end_date=req.period_end_date,
        status=req.status,
        wizard_state_json=req.wizard_state_json,
        draft_content=req.draft_content,
        final_content=req.final_content,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)
    return {"id": report.id}


@router.post("/audit/ca-questions")
async def get_audit_ca_questions(payload: dict):
    """
    Get CA-style clarification questions for high/medium risk TB rows, and run
    the trial balance classifier to surface risk flags.

    Payload: {
        "tb_data": [{"account": str, "mappedTo": str, "amount": float}],
        "trial_balance_rows": [...],  # alias for tb_data, used by the classifier
        "company_info": {...}         # optional, used for going-concern share-capital check
    }
    Returns: {
        "questions": [{"id": str, "question": str, "account": str, "risk": str}],
        "risk_flags": [{"flag": str, "triggered": bool, "detail": str}]
    }
    """
    from core.agents.audit_agent import AuditAgent
    from core.agents.trial_balance_classifier import analyze_trial_balance

    tb_data = payload.get("tb_data", [])
    trial_balance_rows = payload.get("trial_balance_rows") or tb_data
    company_info = payload.get("company_info") or {}

    agent = AuditAgent()
    questions = await agent.ask_questions(tb_data)

    risk_flags: list = []
    if trial_balance_rows:
        analysis = analyze_trial_balance(trial_balance_rows, company_info)
        risk_flags = analysis.get("risk_flags", [])

    return {"questions": questions, "risk_flags": risk_flags}


@router.put("/saved/{report_id}")
async def update_saved_report(report_id: str, req: SaveReportRequest, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(SavedReport).where(SavedReport.id == report_id))
    report = result.scalar_one_or_none()
    if not report:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Report not found")
    if req.company_name:
        report.company_name = req.company_name
    if req.status:
        report.status = req.status
    if req.draft_content is not None:
        report.draft_content = req.draft_content
    if req.final_content is not None:
        report.final_content = req.final_content
    await db.commit()
    await db.refresh(report)
    return {"id": report.id}


@router.delete("/saved/{report_id}")
async def delete_saved_report(report_id: str, db: AsyncSession = Depends(get_db)):
    await db.execute(delete(SavedReport).where(SavedReport.id == report_id))
    await db.commit()
    return {"deleted": report_id}


# ── Export Endpoints ──────────────────────────────────────────────────────────

class ExportReportRequest(BaseModel):
    content: str
    filename: str = "audit_report"
    report_type: str = "general"          # "audit" triggers formatter
    company_name: str = ""
    period_end: str = ""
    location: str = "Dubai - United Arab Emirates"
    opinion_type: str = "unqualified"
    rows: list[dict] = []
    template_id: str = ""


class TemplateReviewRequest(BaseModel):
    action: str  # "approve" or "reject"


class TemplateResponse(BaseModel):
    template_id: str
    template: dict
    rows: list[dict]
    confidence: float
    extraction_method: str


class AnalysisChatRequest(BaseModel):
    message: str
    session_history: list[dict] = []
    trial_balance_summary: str = ""
    prior_year_context: str = ""
    draft_content: str = ""
    company_name: str = ""
    period_end: str = ""

class GenerateFromSessionRequest(BaseModel):
    trial_balance_summary: str = ""
    prior_year_context: str = ""
    analysis_history: list[dict] = []
    company_name: str = ""
    period_end: str = ""
    opinion_type: str = "unqualified"

class AgingScheduleRequest(BaseModel):
    rows: list[dict]
    as_of_date: str = ""


@router.post("/export-docx")
async def export_report_docx(req: ExportReportRequest, db: AsyncSession = Depends(get_db)):
    from core.audit_formatter import format_audit_report

    if req.report_type == "audit" and req.template_id:
        result = await db.execute(select(AuditTemplate).where(AuditTemplate.id == req.template_id))
        tpl = result.scalar_one_or_none()
        if tpl:
            template = {
                "document_structure": tpl.document_structure or {},
                "account_grouping": tpl.account_grouping or {},
                "terminology": tpl.terminology or {},
                "formatting_rules": tpl.formatting_rules or {},
            }
            from core.template_report_generator import generate_from_template
            current_data = {
                "company_name": req.company_name,
                "location": req.location,
                "period_end": req.period_end,
                "opinion_type": req.opinion_type,
                "draft_content": req.content,
                "rows": req.rows,
            }
            docx_bytes = generate_from_template(current_data, template)
        else:
            report_data = {
                "company_name": req.company_name,
                "location": req.location,
                "period_end": req.period_end,
                "opinion_type": req.opinion_type,
                "draft_content": req.content,
                "rows": req.rows,
            }
            docx_bytes = format_audit_report(report_data)
    elif req.report_type == "audit":
        report_data = {
            "company_name": req.company_name,
            "location": req.location,
            "period_end": req.period_end,
            "opinion_type": req.opinion_type,
            "draft_content": req.content,
            "rows": req.rows,
        }
        docx_bytes = format_audit_report(report_data)
    else:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for line in req.content.splitlines():
            if line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

    filename = req.filename if req.filename.endswith('.docx') else req.filename + '.docx'
    safe = filename.replace(" ", "_")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}"'},
    )


@router.post("/export-xlsx")
async def export_report_xlsx(req: ExportReportRequest):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Audit Report"
    for i, line in enumerate(req.content.splitlines(), start=1):
        ws.cell(row=i, column=1, value=line)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe = req.filename.replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{safe}.xlsx"'},
    )


@router.post("/export-pdf")
async def export_report_pdf(req: ExportReportRequest):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    style_h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=14, spaceAfter=6, spaceBefore=12)
    style_h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12, spaceAfter=4, spaceBefore=10)
    style_h3 = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=11, spaceAfter=4, spaceBefore=8)
    style_body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=4)
    style_bold = ParagraphStyle('Bold', parent=style_body, fontName='Helvetica-Bold')

    story = []
    table_rows: list[list[str]] = []
    in_table = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        data = [row for row in table_rows if not all(cell.strip().replace('-', '').replace(':', '') == '' for cell in row)]
        if len(data) >= 1:
            col_count = max(len(r) for r in data)
            padded = [r + [''] * (col_count - len(r)) for r in data]
            t = Table(padded, hAlign='LEFT', repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 6))
        table_rows = []
        in_table = False

    for raw_line in req.content.splitlines():
        line = raw_line.strip()

        if line.startswith('|'):
            in_table = True
            cells = [c.strip() for c in line.strip('|').split('|')]
            table_rows.append(cells)
            continue
        elif in_table:
            flush_table()

        if not line:
            story.append(Spacer(1, 6))
            continue

        if line.startswith('#### '):
            story.append(Paragraph(line[5:], style_h3))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], style_h2))
        elif line.startswith('## '):
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#e2e8f0'), spaceAfter=4))
            story.append(Paragraph(line[3:], style_h1))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], style_h1))
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            story.append(Paragraph(line[2:-2], style_bold))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(f'• {line[2:]}', style_body))
        elif line.startswith('> '):
            callout_style = ParagraphStyle('Callout', parent=style_body, leftIndent=12, textColor=colors.HexColor('#64748b'), borderPadding=(4, 4, 4, 4))
            story.append(Paragraph(line[2:], callout_style))
        else:
            import re as _re
            formatted = _re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(formatted, style_body))

    if in_table:
        flush_table()

    doc.build(story)
    buf.seek(0)
    safe = req.filename.replace(' ', '_')
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe}.pdf"'},
    )


@router.post("/extract-format")
async def extract_format(file: UploadFile = File(...)):
    """
    Extract structured section/heading information from a template file.
    Returns sections list + raw preview for LLM-guided formatting.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    allowed = {"txt", "docx", "doc", "pdf", "xlsx", "xls"}
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        from core.format_extractor import extract_format_structure
        result = extract_format_structure(tmp_path, ext)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract format: {e}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@router.post("/extract-prior-year")
async def extract_prior_year(
    file: UploadFile = File(...),
    session_id: str = Form(default=""),
):
    """Extract prior year financial data from an uploaded PDF (digital or scanned)."""
    import os
    from pathlib import Path as FilePath
    from core.prior_year_extractor import extract_prior_year_from_pdf

    suffix = FilePath(file.filename or "upload.pdf").suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = await extract_prior_year_from_pdf(tmp_path)
    finally:
        os.unlink(tmp_path)

    return result


# ── Analysis Chat (SSE streaming) ──────────────────────────────────────────

@router.post("/analysis-chat")
async def analysis_chat(req: AnalysisChatRequest):
    """Stream an analysis-discussion response from the LLM given audit context."""
    system_prompt = (
        "You are a senior chartered accountant reviewing a client's financial statements. "
        "You have access to the trial balance, prior year comparatives, and draft audit report. "
        "Answer the auditor's questions concisely and professionally."
    )

    # Trim context to avoid exceeding model token limit (prevents mid-stream TCP reset)
    tb_summary = (req.trial_balance_summary or "")[:2500]
    py_context = (req.prior_year_context or "")[:1000]
    draft_excerpt = (req.draft_content or "")[:2000]

    context_block = ""
    if tb_summary:
        context_block += f"\n\n### Trial Balance Summary\n{tb_summary}"
    if py_context:
        context_block += f"\n\n### Prior Year Context\n{py_context}"
    if draft_excerpt:
        context_block += f"\n\n### Draft Report (excerpt)\n{draft_excerpt}"

    messages = [{"role": "system", "content": system_prompt + context_block}]
    # Strip leading assistant messages — NVIDIA NIM requires first non-system message to be 'user'
    history = list(req.session_history[-8:])
    while history and history[0].get("role") == "assistant":
        history.pop(0)
    for h in history:
        messages.append(h)
    messages.append({"role": "user", "content": req.message})

    async def generate():
        try:
            llm = get_llm_provider()
            async for chunk in llm.chat_stream(messages, temperature=0.3, max_tokens=2000):
                yield f"data: {json.dumps({'type': 'chunk', 'content': chunk})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            logger.error(f"Analysis chat error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── Generate Report from Analysis Session ──────────────────────────────────

@router.post("/generate-from-session")
async def generate_from_session(req: GenerateFromSessionRequest, db: AsyncSession = Depends(get_db)):
    """Generate a full audit draft from a completed analysis chat session."""
    analysis_summary = "\n".join(
        f"{m['role'].upper()}: {m['content']}" for m in req.analysis_history[-10:]
    )

    prompt = (
        f"Company: {req.company_name}\n"
        f"Period End: {req.period_end}\n"
        f"Opinion Type: {req.opinion_type}\n\n"
        f"Trial Balance Summary:\n{req.trial_balance_summary}\n\n"
        f"Prior Year Context:\n{req.prior_year_context}\n\n"
        f"Analysis Discussion:\n{analysis_summary}\n\n"
        "Based on the above, produce a complete Independent Auditor's Report."
    )

    llm = get_llm_provider()
    response = await llm.chat([{"role": "user", "content": prompt}])
    draft = response.content
    return {"draft_content": draft, "company_name": req.company_name, "period_end": req.period_end}


# ── Aging Schedule ──────────────────────────────────────────────────────────

@router.post("/aging-schedule")
async def compute_aging_schedule(req: AgingScheduleRequest):
    """
    Compute a simple aging schedule from a list of receivable rows.
    Each row: {invoice_date, amount, customer}
    Returns buckets: current, 30, 60, 90, 90+
    """
    from datetime import datetime

    as_of = datetime.today()
    if req.as_of_date:
        try:
            as_of = datetime.strptime(req.as_of_date, "%Y-%m-%d")
        except ValueError:
            pass

    buckets = {"current": 0.0, "1_30": 0.0, "31_60": 0.0, "61_90": 0.0, "over_90": 0.0}

    for row in req.rows:
        try:
            inv_date = datetime.strptime(row["invoice_date"], "%Y-%m-%d")
            amount = float(row.get("amount", 0))
            age = (as_of - inv_date).days
            if age <= 0:
                buckets["current"] += amount
            elif age <= 30:
                buckets["1_30"] += amount
            elif age <= 60:
                buckets["31_60"] += amount
            elif age <= 90:
                buckets["61_90"] += amount
            else:
                buckets["over_90"] += amount
        except (KeyError, ValueError):
            continue

    total = sum(buckets.values())
    return {"buckets": buckets, "total": total, "as_of_date": as_of.strftime("%Y-%m-%d")}


# ── Template Management ────────────────────────────────────────────────────

@router.post("/extract-audit-template", response_model=TemplateResponse)
async def extract_audit_template(file: UploadFile = File(...), db: AsyncSession = Depends(get_db)):
    """Extract audit report template from an uploaded prior year PDF."""
    from pathlib import Path as FilePath
    from core.prior_year_extractor import extract_prior_year_from_pdf

    suffix = FilePath(file.filename or "upload.pdf").suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = await extract_prior_year_from_pdf(tmp_path)
    finally:
        os.unlink(tmp_path)

    template_data = result.get("template", {})

    tpl = AuditTemplate(
        company_name=template_data.get("document_structure", {}).get("company_name", "Unknown"),
        document_structure=template_data.get("document_structure"),
        account_grouping=template_data.get("account_grouping"),
        terminology=template_data.get("terminology"),
        formatting_rules=template_data.get("formatting_rules"),
        extraction_method=result.get("extraction_method", "unknown"),
        confidence=result.get("confidence", 0.0),
        status="draft",
        source_filename=file.filename or "unknown.pdf",
    )
    db.add(tpl)
    await db.flush()

    return {
        "template_id": tpl.id,
        "template": template_data,
        "rows": result.get("rows", []),
        "confidence": result.get("confidence", 0.0),
        "extraction_method": result.get("extraction_method", "unknown"),
    }


@router.post("/review-template/{template_id}")
async def review_template(template_id: str, req: TemplateReviewRequest, db: AsyncSession = Depends(get_db)):
    """Approve or reject an extracted audit template."""
    result = await db.execute(select(AuditTemplate).where(AuditTemplate.id == template_id))
    tpl = result.scalar_one_or_none()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")

    if req.action not in ("approve", "reject"):
        raise HTTPException(status_code=400, detail="Action must be 'approve' or 'reject'")

    tpl.status = "approved" if req.action == "approve" else "rejected"
    return {"status": tpl.status, "template_id": template_id}


@router.get("/template/{template_id}")
async def get_template(template_id: str, db: AsyncSession = Depends(get_db)):
    """Retrieve a stored audit template by ID."""
    result = await db.execute(select(AuditTemplate).where(AuditTemplate.id == template_id))
    tpl = result.scalar_one_or_none()
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")

    return {
        "template_id": tpl.id,
        "company_name": tpl.company_name,
        "status": tpl.status,
        "confidence": tpl.confidence,
        "extraction_method": tpl.extraction_method,
        "source_filename": tpl.source_filename,
        "document_structure": tpl.document_structure,
        "account_grouping": tpl.account_grouping,
        "terminology": tpl.terminology,
        "formatting_rules": tpl.formatting_rules,
        "created_at": tpl.created_at.isoformat() if tpl.created_at else None,
    }
