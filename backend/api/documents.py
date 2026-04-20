"""
Documents API – Upload, manage, and query indexed documents for RAG.
"""

import logging
import io
import shutil
import uuid
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select, desc
from sqlalchemy.ext.asyncio import AsyncSession

from db.database import get_db
from db.models import Document
from core.document_processor import document_processor
from core.rag_engine import rag_engine
from core.documents.summarizer import summarize_document_text
from core.documents.xlsx_injector import should_inject, parse_to_structured_text
from config import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/documents", tags=["Documents"])


# ── Response Schemas ──────────────────────────────────────────────

class DocumentResponse(BaseModel):
    id: str
    filename: str
    original_name: str
    file_type: str
    file_size: int
    chunk_count: int
    status: str
    error_message: Optional[str] = None
    created_at: str
    summary: Optional[str] = None
    key_terms: Optional[list[str]] = None
    source: Optional[str] = None

class DocumentSearchResult(BaseModel):
    text: str
    source: str
    page: str
    score: float

class UploadResponse(BaseModel):
    document: DocumentResponse
    message: str

class ExportSourceRequest(BaseModel):
    text: str
    filename: str = "source"


# ── Endpoints ─────────────────────────────────────────────────────

@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Upload a document, parse it, and index it for RAG search."""

    # Validate file type
    original_name = file.filename or "unknown"
    suffix = Path(original_name).suffix.lower()

    if not document_processor.is_supported(original_name):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {suffix}. "
                   f"Supported: {', '.join(document_processor.SUPPORTED_TYPES)}",
        )

    # Generate unique filename to avoid collisions
    doc_id = str(uuid.uuid4())
    stored_filename = f"{doc_id}{suffix}"
    upload_path = Path(settings.upload_dir) / stored_filename

    # Save file to disk
    try:
        with open(upload_path, "wb") as f:
            content = await file.read()
            f.write(content)
        file_size = len(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

    # Create DB record
    doc = Document(
        id=doc_id,
        filename=stored_filename,
        original_name=original_name,
        file_type=suffix.lstrip("."),
        file_size=file_size,
        status="processing",
    )
    db.add(doc)
    await db.flush()

    # Process and index the document
    try:
        chunks = await document_processor.process(str(upload_path), doc_id=doc_id)
        if not chunks:
            doc.status = "error"
            doc.error_message = "No text could be extracted from this file"
            logger.warning(f"Document '{original_name}' produced zero chunks — marked as error")
        else:
            chunk_count = await rag_engine.ingest_chunks(chunks, doc_id=doc_id)
            doc.chunk_count = chunk_count
            if chunk_count == 0:
                doc.status = "error"
                doc.error_message = "Text extracted but embedding/indexing produced zero chunks"
                logger.warning(f"Document '{original_name}' zero chunks after indexing — marked as error")
            else:
                doc.status = "indexed"
                logger.info(f"Document '{original_name}' indexed: {chunk_count} chunks")

                # Auto-summarize in background (non-blocking)
                try:
                    full_text = " ".join(
                        c.get("text", "") if isinstance(c, dict) else getattr(c, "text", "")
                        for c in chunks
                    )
                    if full_text.strip():
                        result = await summarize_document_text(full_text)
                        doc.summary = result.summary
                        doc.key_terms = result.key_terms
                except Exception as e:
                    logger.warning(f"Auto-summary failed for {doc_id}: {e}")

                # Small xlsx/csv injection
                try:
                    if should_inject(str(upload_path), file_size):
                        structured = parse_to_structured_text(str(upload_path))
                        if structured:
                            existing_meta = doc.metadata_json if isinstance(doc.metadata_json, dict) else {}
                            existing_meta["structured_text"] = structured
                            doc.metadata_json = existing_meta
                except Exception as e:
                    logger.warning(f"Structured text injection failed for {doc_id}: {e}")

    except Exception as e:
        doc.status = "error"
        err_msg = str(e).strip()
        doc.error_message = err_msg if err_msg else f"{type(e).__name__} (empty error message)"
        logger.error(f"Document processing error for '{original_name}': {e}")

    return UploadResponse(
        document=DocumentResponse(
            id=doc.id,
            filename=doc.filename,
            original_name=doc.original_name,
            file_type=doc.file_type,
            file_size=doc.file_size,
            chunk_count=doc.chunk_count,
            status=doc.status,
            error_message=doc.error_message,
            created_at=str(doc.created_at),
            summary=doc.summary,
            key_terms=doc.key_terms,
            source=doc.source,
        ),
        message=f"Document '{original_name}' uploaded and {'indexed' if doc.status == 'indexed' else 'failed'}.",
    )


@router.get("/", response_model=list[DocumentResponse])
async def list_documents(db: AsyncSession = Depends(get_db)):
    """List all uploaded documents."""
    result = await db.execute(
        select(Document).order_by(desc(Document.created_at))
    )
    docs = result.scalars().all()

    return [
        DocumentResponse(
            id=d.id,
            filename=d.filename,
            original_name=d.original_name,
            file_type=d.file_type,
            file_size=d.file_size,
            chunk_count=d.chunk_count,
            status=d.status,
            error_message=d.error_message,
            created_at=str(d.created_at),
            summary=d.summary,
            key_terms=d.key_terms,
            source=d.source,
        )
        for d in docs
    ]


@router.get("/search")
async def search_documents(
    query: str,
    top_k: int = 5,
    doc_id: Optional[str] = None,
):
    """Search indexed documents using semantic similarity."""
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    try:
        results = await rag_engine.search(query, top_k=top_k, doc_id=doc_id)
    except Exception as e:
        raise HTTPException(
            status_code=502,
            detail=f"RAG search failed: {e}. Check NVIDIA_BASE_URL and NVIDIA_API_KEY in .env",
        )

    return {
        "query": query,
        "results": [
            DocumentSearchResult(
                text=r["text"],
                source=r["metadata"].get("source", "Unknown"),
                page=r["metadata"].get("page", "?"),
                score=round(r.get("score", 0), 3),
            )
            for r in results
        ],
        "total_results": len(results),
    }


@router.get("/source-content")
async def get_source_content(source: str, page: Optional[str] = None, db: AsyncSession = Depends(get_db)):
    """
    Return the full stored text for a given source filename (and optional page).
    Looks up all ChromaDB chunks whose metadata.source matches the given filename,
    then concatenates them in order. Falls back to all chunks if page not found.
    If no results found by the given source, falls back to DB lookup by ID,
    original_name, or stored filename to resolve the correct filename for ChromaDB.
    """
    def _query_chroma(src: str):
        return rag_engine.collection.get(
            where={"source": src},
            include=["documents", "metadatas"],
        )

    # 1. Initial direct attempt with provided 'source' string
    try:
        results = _query_chroma(source)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Vector store query failed: {e}")

    docs = results.get("documents") or []
    metas = results.get("metadatas") or []

    # 2. Fallback: try to resolve document record from SQL DB if direct Chroma query failed
    if not docs:
        # Check if 'source' is an ID
        db_result = await db.execute(select(Document).where(Document.id == source))
        doc_record = db_result.scalar_one_or_none()
        
        # Check if 'source' is original_name
        if not doc_record:
            db_result = await db.execute(select(Document).where(Document.original_name == source))
            doc_record = db_result.scalar_one_or_none()
            
        # Check if 'source' is filename (already prefixed)
        if not doc_record:
            db_result = await db.execute(select(Document).where(Document.filename == source))
            doc_record = db_result.scalar_one_or_none()

        if doc_record:
            try:
                results = _query_chroma(doc_record.filename)
                docs = results.get("documents") or []
                metas = results.get("metadatas") or []
            except Exception:
                pass

    if not docs:
        raise HTTPException(
            status_code=404,
            detail=f"No content found for source '{source}'. "
                   "The document may not be indexed yet.",
        )

    if page and page != "?":
        page_str = str(page)
        filtered = [
            d for d, m in zip(docs, metas)
            if str(m.get("page", "")) == page_str
            or str(m.get("page_number", "")) == page_str
        ]
        text = "\n\n".join(filtered) if filtered else "\n\n".join(docs)
    else:
        text = "\n\n".join(docs)

    return {"source": source, "page": page, "text": text}


@router.post("/export-source-docx")
async def export_source_docx(req: ExportSourceRequest):
    """Export source text content as a .docx file."""
    try:
        import docx as _docx
        doc = _docx.Document()
        doc.add_heading(req.filename, level=1)
        for para in req.text.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{req.filename}_source.docx"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export-source-xlsx")
async def export_source_xlsx(req: ExportSourceRequest):
    """Export source text (markdown table) as .xlsx."""
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Source"
        lines = req.text.strip().split('\n')
        for line in lines:
            if line.startswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                ws.append(cells)
            elif line.strip() and not line.startswith('|-') and not line.startswith('|---'):
                ws.append([line])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{req.filename}_source.xlsx"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{doc_id}")
async def delete_document(doc_id: str, db: AsyncSession = Depends(get_db)):
    """Delete a document and remove it from the vector store."""
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove from DB first — commit before touching filesystem so a partial failure
    # doesn't leave a ghost record pointing to a deleted file.
    upload_path = Path(settings.upload_dir) / doc.filename
    await db.delete(doc)
    await db.commit()

    # Remove from vector store and disk (best-effort; orphaned files can be cleaned up separately)
    removed_chunks = await rag_engine.delete_document(doc_id)
    if upload_path.exists():
        upload_path.unlink()

    return {
        "status": "deleted",
        "document_id": doc_id,
        "chunks_removed": removed_chunks,
    }


@router.get("/stats")
async def document_stats(db: AsyncSession = Depends(get_db)):
    """Get document and vector store statistics."""
    result = await db.execute(select(Document))
    docs = result.scalars().all()

    rag_stats = rag_engine.get_stats()

    return {
        "total_documents": len(docs),
        "indexed_documents": sum(1 for d in docs if d.status == "indexed"),
        "processing_documents": sum(1 for d in docs if d.status == "processing"),
        "error_documents": sum(1 for d in docs if d.status == "error"),
        "zero_chunk_documents": sum(1 for d in docs if d.status == "indexed" and (d.chunk_count or 0) == 0),
        "total_chunks": rag_stats["total_chunks"],
        "total_size_bytes": sum(d.file_size for d in docs),
    }


@router.post("/repair-zero-chunks")
async def repair_zero_chunk_documents(db: AsyncSession = Depends(get_db)):
    """
    Mark all 'indexed' documents with zero chunks as 'error' so they can be re-uploaded.
    Safe to call multiple times (idempotent).
    """
    result = await db.execute(select(Document))
    docs = result.scalars().all()
    fixed = 0
    for d in docs:
        if d.status == "indexed" and (d.chunk_count or 0) == 0:
            d.status = "error"
            d.error_message = "No chunks indexed — re-upload to reprocess"
            fixed += 1
    await db.commit()
    return {"status": "ok", "documents_marked_for_reprocess": fixed}
