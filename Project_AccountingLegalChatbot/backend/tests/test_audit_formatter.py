"""Tests for the audit report DOCX formatter."""
import pytest
from io import BytesIO
from core.audit_formatter import (
    format_audit_report,
    _build_cover_page,
    _build_toc,
    _build_financial_table,
)


def test_format_audit_report_returns_bytes():
    report_data = {
        "company_name": "Test Company LLC",
        "location": "Dubai - United Arab Emirates",
        "period_end": "December 31, 2024",
        "opinion_type": "qualified",
        "draft_content": "We have audited the financial statements.",
        "rows": [
            {"account": "Total Assets", "category": "Assets", "amount": 5929549.0, "prior_year": 9489570.0},
            {"account": "Total Liabilities", "category": "Liabilities", "amount": 6100721.0, "prior_year": 6323991.0},
        ],
    }
    result = format_audit_report(report_data)
    assert isinstance(result, bytes)
    assert len(result) > 1000  # Valid DOCX is never tiny


def test_cover_page_contains_company_name():
    from docx import Document
    doc = Document()
    _build_cover_page(doc, "ACME Corp", "Abu Dhabi - UAE", "December 31, 2024")
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ACME CORP" in full_text
    assert "DECEMBER 31, 2024" in full_text


def test_financial_table_has_correct_columns():
    from docx import Document
    doc = Document()
    rows = [
        {"account": "Trade receivables", "category": "Current Assets", "amount": 720277.0, "prior_year": 424857.0},
    ]
    _build_financial_table(doc, "Statement of Financial Position", rows)
    # Table should exist with 4 columns
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 4
