"""
Document Analyzer — unified extraction for PDF, Excel, and DOCX files.

Provides structured extraction results for the NotebookLM-style audit
profile system. Each extractor returns a consistent dict shape.
"""
import logging
import os
from typing import Any

logger = logging.getLogger(__name__)


def analyze_document(file_path: str, doc_type: str = "auto") -> dict:
    """
    Main dispatcher. Extracts structured data from a document.

    Args:
        file_path: Path to the document file.
        doc_type: One of 'pdf', 'excel', 'docx', or 'auto' (detect from extension).

    Returns:
        {
            "doc_type": str,
            "file_name": str,
            "pages": int | None,
            "tables": list[list[list[str]]],  # list of tables, each table is rows of cells
            "text": str,                       # full text content
            "structure": dict,                 # headings, sections, layout info
            "metadata": dict,                  # file size, author, etc.
        }
    """
    # Map file extension to extraction type
    _EXT_MAP = {".pdf": "pdf", ".xlsx": "excel", ".xls": "excel",
                ".docx": "docx", ".doc": "docx"}
    _EXTRACTION_TYPES = {"pdf", "excel", "docx"}

    if doc_type == "auto" or doc_type not in _EXTRACTION_TYPES:
        # Business types (prior_audit, trial_balance, template, etc.) resolve via extension
        ext = os.path.splitext(file_path)[1].lower()
        resolved = _EXT_MAP.get(ext, "unknown")
        if doc_type != "auto":
            logger.info(f"Resolved business doc_type '{doc_type}' → '{resolved}' from extension")
        doc_type = resolved

    try:
        if doc_type == "pdf":
            return _extract_from_pdf(file_path)
        elif doc_type == "excel":
            return _extract_from_excel(file_path)
        elif doc_type == "docx":
            return _extract_from_docx(file_path)
        else:
            logger.warning(f"Unsupported document type: {doc_type} for {file_path}")
            return _empty_result(os.path.basename(file_path), doc_type)
    except Exception as e:
        logger.error(f"Document analysis failed for {file_path}: {e}")
        result = _empty_result(os.path.basename(file_path), doc_type)
        result["error"] = str(e)
        return result


def _empty_result(file_name: str = "", doc_type: str = "unknown") -> dict:
    return {
        "doc_type": doc_type,
        "file_name": file_name,
        "pages": None,
        "tables": [],
        "text": "",
        "structure": {},
        "metadata": {},
    }


def _extract_from_pdf(file_path: str) -> dict:
    """Extract tables, text, and structure from a PDF using PyMuPDF.
    Falls back to OCR (tesseract) for scanned/image-based PDFs."""
    import fitz

    doc = fitz.open(file_path)
    file_name = os.path.basename(file_path)
    page_count = len(doc)

    all_text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []
    headings: list[dict] = []

    for page_num in range(page_count):
        page = doc[page_num]

        # Extract text
        text = page.get_text("text")
        if text.strip():
            all_text_parts.append(text)

        # Extract tables
        try:
            tables = page.find_tables()
            for table in (tables or []):
                extracted = table.extract()
                if extracted and len(extracted) > 1:
                    cleaned = [[cell or "" for cell in row] for row in extracted]
                    all_tables.append(cleaned)
        except Exception:
            pass  # Some pages may not have tables

        # Detect headings by font size
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:  # text block
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 11)
                    text_content = span.get("text", "").strip()
                    if size >= 13 and text_content and len(text_content) > 2:
                        headings.append({
                            "text": text_content,
                            "page": page_num + 1,
                            "font_size": round(size, 1),
                            "bold": "bold" in span.get("font", "").lower(),
                        })

    doc.close()

    full_text = "\n".join(all_text_parts)

    # OCR fallback: if no text extracted, the PDF is likely scanned/image-based
    is_ocr = False
    if not full_text.strip() and page_count > 0:
        logger.info(f"No text from native extraction — attempting OCR on {file_name}")
        ocr_text, ocr_tables = _ocr_extract_pdf(file_path, page_count)
        if ocr_text.strip():
            full_text = ocr_text
            all_tables.extend(ocr_tables)
            is_ocr = True
            logger.info(f"OCR extracted {len(full_text)} chars, {len(ocr_tables)} tables from {file_name}")

    file_size = os.path.getsize(file_path)

    return {
        "doc_type": "pdf",
        "file_name": file_name,
        "pages": page_count,
        "tables": all_tables,
        "text": full_text[:50000],
        "structure": {
            "headings": headings,
            "table_count": len(all_tables),
            "has_tables": len(all_tables) > 0,
        },
        "metadata": {
            "file_size": file_size,
            "page_count": page_count,
            "ocr_used": is_ocr,
        },
    }


def _ocr_extract_pdf(file_path: str, page_count: int) -> tuple:
    """OCR fallback for scanned/image-based PDFs using tesseract."""
    import re as _re

    try:
        import fitz
        import pytesseract  # type: ignore[import]
        from PIL import Image
        import io as _io

        # Set tesseract path on Windows if not in PATH
        tesseract_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for tp in tesseract_paths:
            if os.path.exists(tp):
                pytesseract.pytesseract.tesseract_cmd = tp
                break

        doc = fitz.open(file_path)
        all_text_parts: list[str] = []
        all_tables: list[list[list[str]]] = []

        for page_num in range(min(page_count, 30)):  # Cap at 30 pages for performance
            page = doc[page_num]
            mat = fitz.Matrix(250 / 72, 250 / 72)  # 250 DPI
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img = Image.open(_io.BytesIO(img_bytes))

            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                all_text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            # Try to extract tabular data from OCR text using regex patterns
            table_rows = _parse_financial_rows_from_text(text)
            if table_rows:
                all_tables.append(table_rows)

        doc.close()
        return "\n".join(all_text_parts), all_tables

    except ImportError as e:
        logger.warning(f"OCR dependencies not available: {e}")
        return "", []
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return "", []


def _parse_financial_rows_from_text(text: str) -> list[list[str]]:
    """Parse table-like rows from OCR text (account name + numbers)."""
    import re as _re
    rows: list[list[str]] = []
    # Pattern: text followed by one or more number columns separated by spaces
    pattern = _re.compile(
        r'^(.{3,50}?)\s{2,}([\-\(]?\d[\d,]*(?:\.\d+)?[\)]?)'
        r'(?:\s{2,}([\-\(]?\d[\d,]*(?:\.\d+)?[\)]?))?'
        r'(?:\s{2,}([\-\(]?\d[\d,]*(?:\.\d+)?[\)]?))?\s*$'
    )
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = pattern.match(line)
        if m:
            row = [m.group(1).strip()]
            for g in [m.group(2), m.group(3), m.group(4)]:
                if g:
                    row.append(g.strip())
            rows.append(row)
    return rows


def _extract_from_excel(file_path: str) -> dict:
    """Extract sheet data, tables, and formatting from an Excel file."""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True)
    file_name = os.path.basename(file_path)
    sheet_names = list(wb.sheetnames)

    all_tables: list[list[list[str]]] = []
    sheets_info: dict[str, Any] = {}
    all_text_parts: list[str] = []

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        rows_data: list[list[str]] = []

        for row in ws.iter_rows(values_only=True):
            str_row = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in str_row):  # Skip fully empty rows
                rows_data.append(str_row)

        if rows_data:
            all_tables.append(rows_data)
            for row in rows_data:
                line = " | ".join(c for c in row if c.strip())
                if line.strip():
                    all_text_parts.append(line)

        sheets_info[sheet_name] = {
            "row_count": len(rows_data),
            "col_count": ws.max_column or 0,
        }

    wb.close()
    full_text = "\n".join(all_text_parts)

    return {
        "doc_type": "excel",
        "file_name": file_name,
        "pages": len(sheet_names),
        "tables": all_tables,
        "text": full_text[:50000],
        "structure": {
            "sheets": sheets_info,
            "sheet_names": sheet_names,
            "table_count": len(all_tables),
            "has_tables": len(all_tables) > 0,
        },
        "metadata": {
            "file_size": os.path.getsize(file_path),
            "sheet_count": len(sheets_info),
        },
    }


def _extract_from_docx(file_path: str) -> dict:
    """Extract headings, paragraphs, and tables from a Word document."""
    from docx import Document

    doc = Document(file_path)
    file_name = os.path.basename(file_path)

    all_tables: list[list[list[str]]] = []
    headings: list[dict] = []
    all_text_parts: list[str] = []

    # Extract paragraphs and headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        all_text_parts.append(text)

        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            headings.append({
                "text": text,
                "level": level,
                "style": para.style.name,
            })

    # Extract tables
    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)
        if table_data:
            all_tables.append(table_data)

    full_text = "\n".join(all_text_parts)

    return {
        "doc_type": "docx",
        "file_name": file_name,
        "pages": None,  # DOCX doesn't have page count without rendering
        "tables": all_tables,
        "text": full_text[:50000],
        "structure": {
            "headings": headings,
            "table_count": len(all_tables),
            "has_tables": len(all_tables) > 0,
            "paragraph_count": len(doc.paragraphs),
        },
        "metadata": {
            "file_size": os.path.getsize(file_path),
        },
    }
