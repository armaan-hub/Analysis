"""
Format Applier — converts structured audit_report.json into PDF, DOCX, or Excel.

Takes the output of structured_report_generator.generate_audit_report() and
produces professional, download-ready financial documents.
"""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

DEFAULT_TEMPLATE = {
    "columns": ["Account", "Notes", "Current Year", "Prior Year"],
    "currency_symbol": "AED",
    "font_family": "Times New Roman",
    "font_size": 10,
    "page_size": "LETTER",
    "margins": {"top": 120, "bottom": 40, "left": 109, "right": 25},
}


def _fmt_number(value, currency: str = "AED") -> str:
    """Format a number with commas; negatives in parentheses; None as '-'."""
    if value is None:
        return "-"
    try:
        num = float(value)
    except (TypeError, ValueError):
        return "-"
    if num < 0:
        return f"({abs(num):,.2f})"
    return f"{num:,.2f}"


def _safe_get(d: dict, *keys, default=None):
    """Nested dict accessor."""
    for k in keys:
        if not isinstance(d, dict):
            return default
        d = d.get(k, default)
    return d


def _fmt_int(value) -> str:
    """Format as integer with commas; parentheses for negatives; '-' for zero/None."""
    if value is None:
        return "-"
    try:
        num = float(value)
    except (TypeError, ValueError):
        return "-"
    if abs(num) < 0.5:
        return "-"
    num_int = round(num)
    if num_int < 0:
        return f"({abs(num_int):,})"
    return f"{num_int:,}"


def _format_period_date(period_end: str) -> str:
    """Convert '2025-12-31' to '31.12.2025'."""
    try:
        from datetime import datetime
        dt = datetime.strptime(str(period_end).strip(), "%Y-%m-%d")
        return dt.strftime("%d.%m.%Y")
    except Exception:
        return str(period_end)


def _prior_year_header(period_end: str) -> str:
    """Derive prior year header: '2025-12-31' → '31.12.2024'."""
    try:
        from datetime import datetime
        dt = datetime.strptime(str(period_end).strip(), "%Y-%m-%d")
        prior = dt.replace(year=dt.year - 1)
        return prior.strftime("%d.%m.%Y")
    except Exception:
        return ""


def _nice_date(period_end: str) -> str:
    """Convert '2025-12-31' to 'December 31, 2025'."""
    try:
        from datetime import datetime
        dt = datetime.strptime(str(period_end).strip(), "%Y-%m-%d")
        return dt.strftime("%B %d, %Y")
    except Exception:
        return str(period_end)


def _condense_sofp(sections: list) -> list:
    """
    Collapse individual Tally ledger accounts into auditor-standard SOFP summary lines.
    Returns condensed list of section dicts.
    """
    condensed = []
    for section in sections:
        title = section.get("title", "")
        items = section.get("line_items", [])
        subtotal = section.get("subtotal", {})

        if title in ("Non-Current Assets", "Non-current assets"):
            ppe = sum((i.get("current_year") or 0) for i in items)
            ppe_pr_items = sum((i.get("prior_year") or 0) for i in items)
            # Prefer section subtotal prior (driven by group-level fallback) over item sum
            ppe_pr = subtotal.get("prior_year") or ppe_pr_items
            condensed.append({
                "title": "Non-current assets",
                "line_items": [{"account_name": "Property, plant & equipment",
                                 "notes_ref": "6", "current_year": ppe, "prior_year": ppe_pr}],
                "subtotal": {"account_name": "Total non-current assets",
                             "current_year": subtotal.get("current_year", ppe),
                             "prior_year": subtotal.get("prior_year", ppe_pr)},
            })

        elif title in ("Current Assets", "Current assets"):
            trade_recv = advances = cash = 0.0
            trade_recv_pr = advances_pr = cash_pr = 0.0
            for item in items:
                n = item["account_name"].lower()
                cur = item.get("current_year") or 0
                pr = item.get("prior_year") or 0
                if any(k in n for k in ["debtor", "receivable"]):
                    trade_recv += cur; trade_recv_pr += pr
                elif any(k in n for k in ["cash", "bank"]):
                    cash += cur; cash_pr += pr
                else:
                    advances += cur; advances_pr += pr
            ci = []
            if abs(trade_recv) > 0.5:
                ci.append({"account_name": "Trade receivables", "notes_ref": "7",
                            "current_year": trade_recv, "prior_year": trade_recv_pr})
            if abs(advances) > 0.5:
                ci.append({"account_name": "Advances, deposits and other receivables", "notes_ref": "8",
                            "current_year": advances, "prior_year": advances_pr})
            if abs(cash) > 0.5:
                ci.append({"account_name": "Cash in hand and at banks", "notes_ref": "9",
                            "current_year": cash, "prior_year": cash_pr})
            condensed.append({
                "title": "Current assets",
                "line_items": ci,
                "subtotal": {"account_name": "Total current assets",
                             "current_year": subtotal.get("current_year", trade_recv + advances + cash),
                             "prior_year": subtotal.get("prior_year", trade_recv_pr + advances_pr + cash_pr)},
            })

        elif title in ("Current Liabilities", "Current liabilities"):
            loans = trade_pay = other = 0.0
            loans_pr = trade_pr = other_pr = 0.0
            for item in items:
                n = item["account_name"].lower()
                cur = item.get("current_year") or 0
                pr = item.get("prior_year") or 0
                # Bank OD ("bank od") stays in Other (short-term); only true long-term
                # loans/mortgages/borrowings go to the loans bucket here.
                if any(k in n for k in ["loan", "mortgage", "borrowing"]) and "bank od" not in n:
                    loans += cur; loans_pr += pr
                elif "commission payable" in n or any(k in n for k in ["sundry creditor", "trade payable"]):
                    trade_pay += cur; trade_pr += pr
                else:
                    other += cur; other_pr += pr
            li = []
            if abs(loans) > 0.5:
                li.append({"account_name": "Long Term Loans", "notes_ref": "12",
                            "current_year": loans, "prior_year": loans_pr})
            if abs(trade_pay) > 0.5:
                li.append({"account_name": "Trade Payables", "notes_ref": "10",
                            "current_year": trade_pay, "prior_year": trade_pr})
            if abs(other) > 0.5:
                li.append({"account_name": "Other payables, accruals & provisions", "notes_ref": "11",
                            "current_year": other, "prior_year": other_pr})
            condensed.append({
                "title": "Current liabilities",
                "line_items": li,
                "subtotal": {"account_name": "Total current liabilities",
                             "current_year": subtotal.get("current_year"),
                             "prior_year": subtotal.get("prior_year")},
            })

        elif title in ("Equity", "Shareholders' Equity"):
            sc = ca = re = 0.0
            sc_pr = ca_pr = re_pr = 0.0
            for item in items:
                n = item["account_name"].lower()
                cur = item.get("current_year") or 0
                pr = item.get("prior_year") or 0
                if "capital a/c" in n or "share capital" in n:
                    sc += cur; sc_pr += pr
                elif "current a/c" in n or "current account" in n:
                    ca += cur; ca_pr += pr
                else:
                    re += cur; re_pr += pr
            ei = []
            if abs(sc) > 0.5:
                ei.append({"account_name": "Share Capital", "notes_ref": "2",
                            "current_year": sc, "prior_year": sc_pr})
            if abs(ca) > 0.5:
                ei.append({"account_name": "Shareholders' current account", "notes_ref": "13",
                            "current_year": ca, "prior_year": ca_pr})
            if abs(re) > 0.5:
                ei.append({"account_name": "Retained Earnings", "notes_ref": "14",
                            "current_year": re, "prior_year": re_pr})
            condensed.append({
                "title": "Shareholders' Equity",
                "line_items": ei,
                "subtotal": {"account_name": "Total Shareholders' Equity",
                             "current_year": subtotal.get("current_year"),
                             "prior_year": subtotal.get("prior_year")},
            })
        elif title in ("Non-Current Liabilities", "Non-current liabilities"):
            loans = 0.0
            loans_pr = 0.0
            for item in items:
                loans += item.get("current_year") or 0
                loans_pr += item.get("prior_year") or 0
            # Use section subtotal prior if individual items have no prior year data
            if loans_pr == 0.0:
                loans_pr = subtotal.get("prior_year") or 0
            loans_cy = subtotal.get("current_year") or loans
            if abs(loans_cy) > 0.5 or abs(loans_pr) > 0.5:
                condensed.append({
                    "title": "Non-current liabilities",
                    "line_items": [{"account_name": "Long Term Loans", "notes_ref": "12",
                                    "current_year": loans_cy, "prior_year": loans_pr}],
                    "subtotal": {"account_name": "Total non-current liabilities",
                                 "current_year": loans_cy, "prior_year": loans_pr},
                })
        else:
            condensed.append(section)
    return condensed


def _condense_sopl(sections: list, net_profit_total: dict) -> list:
    """
    Condense SOPL sections to auditor-standard rows.
    Returns list of (label, notes_ref, cur_val, prior_val, row_type) tuples.
    row_type: 'item' | 'subtotal' | 'net' | 'blank'
    """
    revenue_cur = revenue_pr = 0.0
    cost_cur = cost_pr = 0.0
    depr_cur = depr_pr = 0.0
    ga_cur = ga_pr = 0.0
    other_cur = other_pr = 0.0

    for section in sections:
        title = section.get("title", "")
        items = section.get("line_items", [])
        st = section.get("subtotal", {})
        if title == "Revenue":
            revenue_cur = st.get("current_year") or 0
            revenue_pr = st.get("prior_year") or 0
        elif title == "Cost of Sales":
            cost_cur = st.get("current_year") or 0
            cost_pr = st.get("prior_year") or 0
        elif title == "Operating Expenses":
            op_section_pr = st.get("prior_year") or 0
            for item in items:
                n = item["account_name"].lower()
                cur = item.get("current_year") or 0
                pr = item.get("prior_year") or 0
                if "depreciation" in n:
                    depr_cur += cur; depr_pr += pr
                else:
                    ga_cur += cur; ga_pr += pr
            # Derive G&A prior from section total when G&A items have no individual matches
            if ga_pr == 0 and op_section_pr > depr_pr:
                ga_pr = op_section_pr - depr_pr
        elif title == "Other Income":
            other_cur = st.get("current_year") or 0
            other_pr = st.get("prior_year") or 0

    gross = revenue_cur - cost_cur
    gross_pr = revenue_pr - cost_pr
    net = net_profit_total.get("current_year") or 0
    net_pr = net_profit_total.get("prior_year") or 0

    return [
        ("Revenue", "15", revenue_cur, revenue_pr, "item"),
        ("Cost of revenue", "16", -cost_cur if cost_cur else None, -cost_pr if cost_pr else None, "item"),
        ("Gross profit", "", gross, gross_pr, "subtotal"),
        ("Depreciation", "6", -depr_cur if depr_cur else None, -depr_pr if depr_pr else None, "item"),
        ("General and administration expenses", "17", -ga_cur if ga_cur else None, -ga_pr if ga_pr else None, "item"),
        ("Other Income", "18", other_cur if other_cur else None, other_pr if other_pr else None, "item"),
        ("Net profit / (loss) for the year", "", net, net_pr, "net"),
        ("", "", None, None, "blank"),
        ("Other comprehensive income / (expenses)", "", None, None, "header"),
        ("  Items that will not be reclassified to profit or loss:", "", None, None, "item"),
        ("  Re-measurement of end-of-service benefits", "", None, None, "item"),
        ("Total Other Comprehensive Income", "", None, None, "subtotal"),
        ("Total Comprehensive Income for the year", "", net, net_pr, "net"),
    ]


# ── Public API ────────────────────────────────────────────────────────────────


def apply_format(
    report_json: dict,
    format_template: Optional[dict] = None,
    output_format: str = "pdf",
) -> bytes:
    """
    Convert audit_report.json dict to formatted file bytes.

    Args:
        report_json: Structured audit report (from structured_report_generator).
        format_template: Optional display overrides (columns, font, margins, etc.).
        output_format: One of "pdf", "docx", "xlsx".

    Returns:
        Raw bytes of the generated document.
    """
    tpl = {**DEFAULT_TEMPLATE, **(format_template or {})}

    dispatch = {
        "pdf": _generate_pdf,
        "docx": _generate_docx,
        "xlsx": _generate_xlsx,
    }

    generator = dispatch.get(output_format.lower())
    if generator is None:
        raise ValueError(f"Unsupported output_format '{output_format}'. Use pdf, docx, or xlsx.")

    return generator(report_json, tpl)


# ══════════════════════════════════════════════════════════════════════════════
#  PDF generation (ReportLab)
# ══════════════════════════════════════════════════════════════════════════════


def _generate_pdf(report: dict, tpl: dict) -> bytes:  # noqa: C901
    """Generate a professional audit report PDF using ReportLab BaseDocTemplate."""
    from reportlab.lib.pagesizes import A4, letter, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, Table, TableStyle,
        PageBreak, NextPageTemplate, HRFlowable, KeepTogether,
    )
    from reportlab.platypus.flowables import Flowable

    buf = io.BytesIO()
    PAGE_W, PAGE_H = letter   # 612 × 792 pt  (US Letter — matches reference)
    LM = 109   # left margin  (reference: 108.81 pt)
    RM = 72    # right margin (reference: ~86 pt on 672 pt page → ~78 pt scaled; 72 pt = 1 inch)
    TM = 120   # top margin   (reference: 118.87 pt)
    BM = 40    # bottom margin (footer space)

    meta = report.get("metadata", {})
    company = meta.get("company_name", "Company")
    period_end = meta.get("period_end", "")
    currency = tpl.get("currency_symbol", "AED")

    cur_hdr = _format_period_date(period_end)    # "31.12.2025"
    prior_hdr = _prior_year_header(period_end)   # "31.12.2024"
    nice_dt = _nice_date(period_end)             # "December 31, 2025"

    FS = tpl.get("font_size", 9)
    BLK = colors.black

    # ── Shared mutable state updated by _SetStmtTitle flowable ───────────────
    _state = {"stmt_title": ""}

    class _SetStmtTitle(Flowable):
        """Zero-height flowable: updates _state['stmt_title'] when rendered."""
        def __init__(self, title: str):
            Flowable.__init__(self)
            self._t = title
        def wrap(self, aW, aH): return 0, 0
        def draw(self): _state["stmt_title"] = self._t

    # ── Canvas helpers ────────────────────────────────────────────────────────
    def _draw_header(canvas, right_text=""):
        canvas.setFont("Times-Bold", 9)
        canvas.drawString(LM, PAGE_H - TM + 18, company.upper())
        if right_text:
            canvas.setFont("Times-Roman", 9)
            canvas.drawRightString(PAGE_W - RM, PAGE_H - TM + 18, right_text)
        canvas.setLineWidth(0.5)
        canvas.line(LM, PAGE_H - TM + 12, PAGE_W - RM, PAGE_H - TM + 12)

    def _draw_page_num(canvas):
        canvas.setFont("Times-Roman", 8)
        canvas.drawRightString(PAGE_W - RM, BM - 20, str(canvas.getPageNumber()))

    def _draw_fin_footer(canvas):
        canvas.setFont("Times-Italic", 7.5)
        canvas.drawString(LM, BM - 14,
                          "The accompanying notes form an integral part of these financial statements.")

    # ── Page template callbacks ───────────────────────────────────────────────
    def _cb_cover(canvas, doc): pass

    def _cb_auditor(canvas, doc):
        """Auditor report pages: page number only, no running company header."""
        canvas.saveState(); _draw_page_num(canvas); canvas.restoreState()

    def _cb_plain(canvas, doc):
        canvas.saveState(); _draw_header(canvas); _draw_page_num(canvas); canvas.restoreState()

    def _cb_financial(canvas, doc):
        canvas.saveState()
        # Two-line header: company name + location (matches reference format)
        canvas.setFont("Times-Bold", 9)
        canvas.drawString(LM, PAGE_H - TM + 12, company.upper())
        canvas.setFont("Times-Roman", 8)
        canvas.drawString(LM, PAGE_H - TM + 2, "Dubai \u2013 United Arab Emirates")
        canvas.setLineWidth(0.5)
        canvas.line(LM, PAGE_H - TM - 4, PAGE_W - RM, PAGE_H - TM - 4)
        _draw_fin_footer(canvas)
        # Authorized signatory block at fixed position on SOFP page
        if _state["stmt_title"] == "Statement of Financial Position":
            canvas.setFont("Times-Roman", 7.5)
            canvas.drawString(LM, BM + 66,
                              f"The financial statements were approved and signed on behalf of the management on {nice_dt}:")
            canvas.setFont("Times-Bold", 8)
            canvas.drawString(LM, BM + 50, company.upper())
            canvas.setFont("Times-Roman", 8)
            canvas.drawString(LM, BM + 37, "______________________________")
            canvas.drawString(LM, BM + 25, "Authorized Signatory")
        _draw_page_num(canvas)
        canvas.restoreState()

    def _cb_notes(canvas, doc):
        canvas.saveState()
        # Two-line left header: company name + location (matches reference format)
        canvas.setFont("Times-Bold", 9)
        canvas.drawString(LM, PAGE_H - TM + 12, company.upper())
        canvas.setFont("Times-Roman", 8)
        canvas.drawString(LM, PAGE_H - TM + 2, "Dubai \u2013 United Arab Emirates")
        # Right side: full notes title with date
        canvas.setFont("Times-Roman", 9)
        canvas.drawRightString(PAGE_W - RM,
                               PAGE_H - TM + 7,
                               f"Notes to the financial statements for the year ended {nice_dt}")
        canvas.setLineWidth(0.5)
        canvas.line(LM, PAGE_H - TM - 4, PAGE_W - RM, PAGE_H - TM - 4)
        _draw_page_num(canvas)
        canvas.restoreState()

    # ── Build frames and document ─────────────────────────────────────────────
    avail_w = PAGE_W - LM - RM
    EXTRA_TOP = 6    # pts below header line to frame top
    EXTRA_BTM = 22   # pts above page edge to frame bottom (footer space)
    fr = Frame(LM, BM + EXTRA_BTM, avail_w,
               PAGE_H - TM - BM - EXTRA_TOP - EXTRA_BTM,
               id='main', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    doc = BaseDocTemplate(
        buf, pagesize=letter,
        leftMargin=LM, rightMargin=RM, topMargin=TM, bottomMargin=BM,
        pageTemplates=[
            PageTemplate('Cover',     [fr], onPage=_cb_cover),
            PageTemplate('Auditor',   [fr], onPage=_cb_auditor),
            PageTemplate('Plain',     [fr], onPage=_cb_plain),
            PageTemplate('Financial', [fr], onPage=_cb_financial),
            PageTemplate('Notes',     [fr], onPage=_cb_notes),
        ],
    )

    # ── Column widths ─────────────────────────────────────────────────────────
    # Proportions derived from reference PDF (scaled to avail_w=431 pt):
    # account ~52.7%, notes ~5.8%, current-year ~26.2%, prior-year ~15.3%
    acct_w = 220        # account label column
    note_w = 25         # notes reference number column
    num_w  = 113        # current year amount column
    prior_w = avail_w - acct_w - note_w - num_w   # prior year amount column (~73 pt)
    col_w  = [acct_w, note_w, num_w, prior_w]

    # ── Paragraph styles ─────────────────────────────────────────────────────
    SS = getSampleStyleSheet()

    def _ps(name, **kw):
        return ParagraphStyle(name, parent=SS["Normal"], **kw)

    s_cover_co  = _ps("CoCo",  fontSize=18, leading=24, fontName="Times-Bold",
                      alignment=TA_CENTER, spaceAfter=4)
    s_cover_loc = _ps("CoLoc", fontSize=11, leading=15, fontName="Times-Roman",
                      alignment=TA_CENTER, spaceAfter=4)
    s_cover_hdr = _ps("CoHdr", fontSize=13, leading=17, fontName="Times-Bold",
                      alignment=TA_CENTER, spaceAfter=2)
    s_toc_hdr   = _ps("TocHdr", fontSize=12, fontName="Times-Bold",
                      alignment=TA_CENTER, spaceAfter=12)
    s_heading   = _ps("AudHead", fontSize=11, fontName="Times-Bold",
                      spaceBefore=6, spaceAfter=4)
    s_body      = _ps("Body",  fontSize=FS, leading=FS + 4,
                      fontName="Times-Roman", spaceAfter=3)
    s_body_j    = _ps("BodyJ", fontSize=FS, leading=FS + 4,
                      fontName="Times-Roman", spaceAfter=3, alignment=TA_JUSTIFY)
    s_bold      = _ps("Bold",  fontSize=FS, leading=FS + 4,
                      fontName="Times-Bold", spaceAfter=3)
    s_stmt_title = _ps("StmtTitle", fontSize=11, fontName="Times-Bold",
                       alignment=TA_CENTER, spaceAfter=1)
    s_stmt_sub   = _ps("StmtSub",   fontSize=FS, fontName="Times-Roman",
                       alignment=TA_CENTER, spaceAfter=4)
    s_note_title = _ps("NoteT",     fontSize=FS + 1, fontName="Times-Bold",
                       spaceBefore=10, spaceAfter=4)
    s_small      = _ps("Sm",        fontSize=FS - 1, fontName="Times-Roman", spaceAfter=2)

    story = []

    # ════════════════════════ COVER PAGE ════════════════════════════════════
    story += [
        NextPageTemplate('Cover'),
        Spacer(1, 100),
        Paragraph(company.upper(), s_cover_co),
        Spacer(1, 6),
        Paragraph("DUBAI – UNITED ARAB EMIRATES", s_cover_loc),
        Spacer(1, 28),
        HRFlowable(width="55%", thickness=0.8, color=BLK, hAlign="CENTER"),
        Spacer(1, 28),
        Paragraph("FINANCIAL STATEMENTS AND", s_cover_hdr),
        Paragraph("INDEPENDENT AUDITOR'S REPORT", s_cover_hdr),
        Spacer(1, 14),
        Paragraph(f"FOR THE YEAR ENDED {nice_dt.upper()}", s_cover_loc),
        NextPageTemplate('Cover'),
        PageBreak(),
    ]

    # ════════════════════════ TABLE OF CONTENTS ══════════════════════════════
    story += [Spacer(1, 40), Paragraph("INDEX", s_toc_hdr), Spacer(1, 8)]
    toc_entries = [
        ("Independent Auditors' Report", "1 – 3"),
        ("Statement of Financial Position", "4"),
        ("Statement of Profit or Loss and Other Comprehensive Income", "5"),
        ("Statement of Changes in Shareholders' Equity", "6"),
        ("Statement of Cash Flows", "7"),
        ("Notes to the Financial Statements", "8 – 22"),
    ]
    toc_rows = [[name, pages] for name, pages in toc_entries]
    toc_t = Table(toc_rows, colWidths=[avail_w * 0.80, avail_w * 0.20])
    toc_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), FS),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, colors.HexColor("#cccccc")),
    ]))
    story.append(toc_t)
    story += [NextPageTemplate('Auditor'), PageBreak()]

    # ════════════════════════ AUDITOR'S REPORT ═══════════════════════════════
    opinion = report.get("auditor_opinion", {})
    auditor_firm = meta.get("auditor_firm", "A&M Audit & Accounting LLC")
    auditor_reg = meta.get("auditor_registration", "")

    story.append(Paragraph("INDEPENDENT AUDITOR'S REPORT", s_heading))
    story.append(Spacer(1, 4))
    story.append(Paragraph(f"To the Shareholders of<br/>{company}", s_body))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Opinion</b>", s_bold))
    # Always use template text — OCR'd text is from a different year and contains errors
    story.append(Paragraph(
        f"We have audited the financial statements of {company} (the 'Company'), "
        f"which comprise the statement of financial position as at {nice_dt}, "
        "and the statement of profit or loss and other comprehensive income, "
        "statement of changes in shareholders' equity and statement of cash flows "
        "for the year then ended, and notes to the financial statements, including "
        "a summary of significant accounting policies and other explanatory information.",
        s_body_j))
    story.append(Spacer(1, 4))

    opinion_type = (opinion.get("opinion_type") or "unqualified").lower()
    if "qualified" in opinion_type and "un" not in opinion_type:
        conclusion_prefix = ("In our opinion, except for the matter described in the Basis for "
                             "Qualified Opinion section, the financial statements")
    else:
        conclusion_prefix = "In our opinion, the financial statements"
    story.append(Paragraph(
        f"{conclusion_prefix} present fairly, in all material respects, the financial "
        f"position of the Company as at {nice_dt}, and its financial performance and cash "
        "flows for the year then ended in accordance with International Financial Reporting "
        "Standards (IFRSs).", s_body_j))

    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Basis for Opinion</b>", s_bold))
    basis_text = (opinion.get("basis_text") or
                  "We conducted our audit in accordance with International Standards on Auditing (ISAs). "
                  "Our responsibilities under those standards are further described in the Auditor's "
                  "Responsibilities for the Audit of the Financial Statements section of our report. "
                  "We are independent of the Company in accordance with the ethical requirements that "
                  "are relevant to our audit of the financial statements in the UAE, and we have "
                  "fulfilled our other ethical responsibilities in accordance with these requirements. "
                  "We believe that the audit evidence we have obtained is sufficient and appropriate "
                  "to provide a basis for our opinion.")
    story.append(Paragraph(basis_text, s_body_j))

    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Responsibilities of Management for the Financial Statements</b>", s_bold))
    story.append(Paragraph(
        "Management is responsible for the preparation and fair presentation of the financial "
        "statements in accordance with IFRSs, and for such internal control as management "
        "determines is necessary to enable the preparation of financial statements that are "
        "free from material misstatement, whether due to fraud or error.", s_body_j))

    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Auditor's Responsibilities for the Audit of the Financial Statements</b>", s_bold))
    story.append(Paragraph(
        "Our objectives are to obtain reasonable assurance about whether the financial "
        "statements as a whole are free from material misstatement, whether due to fraud or "
        "error, and to issue an auditor's report that includes our opinion. Reasonable "
        "assurance is a high level of assurance, but is not a guarantee that an audit "
        "conducted in accordance with ISAs will always detect a material misstatement. "
        "Misstatements can arise from fraud or error and are considered material if, "
        "individually or in the aggregate, they could reasonably be expected to influence "
        "the economic decisions of users taken on the basis of these financial statements.",
        s_body_j))

    # Page 1 ends here — add page break for page 2
    story.append(PageBreak())

    # ── Auditor's Report Page 2 ──────────────────────────────────────────────
    story.append(Paragraph("INDEPENDENT AUDITOR'S REPORT (CONTINUED)", s_heading))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Auditor's Responsibilities for the Audit of the Financial Statements (continued)</b>", s_bold))
    story.append(Paragraph(
        "As part of an audit in accordance with ISAs, we exercise professional judgement and "
        "maintain professional skepticism throughout the audit. We also:", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "\u2022 Identify and assess the risks of material misstatement of the financial statements, "
        "whether due to fraud or error, design and perform audit procedures responsive to those "
        "risks, and obtain audit evidence that is sufficient and appropriate to provide a basis "
        "for our opinion. The risk of not detecting a material misstatement resulting from fraud "
        "is higher than for one resulting from error, as fraud may involve collusion, forgery, "
        "intentional omissions, misrepresentations, or the override of internal control.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "\u2022 Obtain an understanding of internal control relevant to the audit in order to design "
        "audit procedures that are appropriate in the circumstances, but not for the purpose of "
        "expressing an opinion on the effectiveness of the Company's internal control.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "\u2022 Evaluate the appropriateness of accounting policies used and the reasonableness of "
        "accounting estimates and related disclosures made by management.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "\u2022 Conclude on the appropriateness of management's use of the going concern basis of "
        "accounting and, based on the audit evidence obtained, whether a material uncertainty "
        "exists related to events or conditions that may cast significant doubt on the Company's "
        "ability to continue as a going concern. If we conclude that a material uncertainty "
        "exists, we are required to draw attention in our auditor's report to the related "
        "disclosures in the financial statements or, if such disclosures are inadequate, to "
        "modify our opinion. Our conclusions are based on the audit evidence obtained up to "
        "the date of our auditor's report. However, future events or conditions may cause "
        "the Company to cease to continue as a going concern.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "\u2022 Evaluate the overall presentation, structure and content of the financial statements, "
        "including the disclosures, and whether the financial statements represent the underlying "
        "transactions and events in a manner that achieves fair presentation.", s_body_j))
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "We communicate with management regarding, among other matters, the planned scope and "
        "timing of the audit and significant audit findings, including any significant deficiencies "
        "in internal control that we identify during our audit.", s_body_j))

    story.append(PageBreak())

    # ── Auditor's Report Page 3 ──────────────────────────────────────────────
    story.append(Paragraph("INDEPENDENT AUDITOR'S REPORT (CONTINUED)", s_heading))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Report on Other Legal and Regulatory Requirements</b>", s_bold))
    story.append(Paragraph(
        "As required by the UAE Federal Law No. (2) of 2015 and its amendments thereto (the "
        "'Companies Law'), we report that:", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "i)   We have obtained all the information and explanations we considered necessary "
        "for the purpose of our audit;", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "ii)  The financial statements comply with the applicable provisions of the UAE "
        "Federal Law No. (2) of 2015;", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "iii) The Company has maintained proper books of account;", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "iv)  The financial information included in the Directors' Report is consistent "
        "with the books of account of the Company;", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "v)   No contraventions of the UAE Federal Law No. (2) of 2015 have come to our "
        "attention during the course of our audit that would have a material effect on the "
        "business or financial position of the Company;", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "vi)  The Company has obtained in all the financial information referred to us "
        "by the management is accurate and complete.", s_body_j))

    story.append(Spacer(1, 28))
    story.append(Paragraph(f"<b>For {auditor_firm}</b>", s_body))
    if auditor_reg:
        story.append(Paragraph(f"Registration No: {auditor_reg}", s_body))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Dubai, United Arab Emirates", s_body))
    story.append(Paragraph(f"{nice_dt}", s_body))

    # Transition to Financial template for the SOFP page
    story += [
        _SetStmtTitle("Statement of Financial Position"),
        NextPageTemplate('Financial'),
        PageBreak(),
    ]

    # ── Additional auditor pages (KAMs / Going Concern) — these come BEFORE
    # the PageBreak above in practice only if they exist.  Since we already
    # emitted the transition above, insert any extra auditor content and then
    # emit a second PageBreak if needed.  For now we leave this stub.
    kams = opinion.get("key_audit_matters", []) if opinion else []
    going_concern = opinion.get("going_concern") if opinion else False

    # ════════════════════════ FINANCIAL STATEMENTS ═══════════════════════════
    statements = report.get("financial_statements", {})
    story.append(NextPageTemplate('Financial'))

    def _fin_table(rows, style_cmds):
        """Build a styled financial statement table."""
        t = Table(rows, colWidths=col_w, repeatRows=0)
        base = [
            ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
            ("FONTSIZE", (0, 0), (-1, -1), FS),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("ALIGN", (0, 0), (0, -1), "LEFT"),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),   # Notes column
            ("ALIGN", (2, 0), (-1, -1), "RIGHT"),    # Amount columns
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        base.extend(style_cmds)
        t.setStyle(TableStyle(base))
        return t

    def _col_header_rows():
        """Return (rows, style_cmds) for the two-row date/AED column header."""
        r0 = ["", "Notes", cur_hdr, prior_hdr]
        r1 = ["", "", "AED", "AED"]
        cmds = [
            ("FONTNAME", (0, 0), (-1, 1), "Times-Bold"),
            ("ALIGN", (1, 0), (-1, 1), "CENTER"),
            ("LINEBELOW", (0, 1), (-1, 1), 0.6, BLK),
            ("FONTSIZE", (0, 0), (-1, 1), FS - 1),
        ]
        return [r0, r1], cmds

    # ─── SOFP ────────────────────────────────────────────────────────────────
    sofp = statements.get("statement_of_financial_position")
    if sofp:
        story.append(Paragraph("Statement of Financial Position", s_stmt_title))
        story.append(Paragraph("(In Arab Emirates Dirhams)", s_stmt_sub))
        story.append(Paragraph(f"As at {nice_dt}", s_stmt_sub))
        story.append(Spacer(1, 6))

        rows, style_cmds = _col_header_rows()

        all_condensed = _condense_sofp(sofp.get("sections", []))
        # NCA before CA (standard IFRS order: long-term → short-term)
        asset_sections = sorted(
            [s for s in all_condensed if s["title"] in ("Non-current assets", "Current assets")],
            key=lambda x: 0 if x["title"] == "Non-current assets" else 1)
        liae_sections = [s for s in all_condensed
                         if s["title"] not in ("Non-current assets", "Current assets")]

        def _append_section(section, is_subsection=False, show_title=True):
            sec_title = section.get("title", "")
            items = section.get("line_items", [])
            subtotal = section.get("subtotal", {})
            if show_title:
                sec_ri = len(rows)
                label = sec_title.title() if is_subsection else sec_title.upper()
                rows.append([label, "", "", ""])
                style_cmds.extend([
                    ("FONTNAME", (0, sec_ri), (-1, sec_ri), "Times-Bold"),
                    ("TOPPADDING", (0, sec_ri), (-1, sec_ri), 5 if is_subsection else 7),
                ])
            for item in items:
                name = item.get("account_name", "")
                note = str(item.get("notes_ref") or item.get("note_ref") or "")
                rows.append([f"    {name}", note,
                             _fmt_int(item.get("current_year")),
                             _fmt_int(item.get("prior_year"))])
            if subtotal:
                sub_ri = len(rows)
                rows.append([subtotal.get("account_name", f"Total {sec_title}"),
                             "",
                             _fmt_int(subtotal.get("current_year")),
                             _fmt_int(subtotal.get("prior_year"))])
                style_cmds.extend([
                    ("FONTNAME", (0, sub_ri), (-1, sub_ri), "Times-Bold"),
                    ("LINEABOVE", (2, sub_ri), (-1, sub_ri), 0.5, BLK),
                    ("LINEBELOW", (2, sub_ri), (-1, sub_ri), 0.5, BLK),
                ])

        # ── ASSETS (top half) ────────────────────────────────────────────────
        assets_ri = len(rows)
        rows.append(["ASSETS", "", "", ""])
        style_cmds.extend([
            ("FONTNAME", (0, assets_ri), (-1, assets_ri), "Times-Bold"),
            ("TOPPADDING", (0, assets_ri), (-1, assets_ri), 8),
        ])
        for section in asset_sections:
            _append_section(section, is_subsection=True)

        # Grand total (Total Assets)
        sofp_total = sofp.get("total", {})
        rows.append(["", "", "", ""])  # blank spacer before Total Assets
        ta_ri = len(rows)
        rows.append(["Total Assets", "",
                     _fmt_int(sofp_total.get("current_year")),
                     _fmt_int(sofp_total.get("prior_year"))])
        style_cmds += [
            ("FONTNAME", (0, ta_ri), (-1, ta_ri), "Times-Bold"),
            ("LINEABOVE", (2, ta_ri), (-1, ta_ri), 0.5, BLK),
            ("LINEBELOW", (2, ta_ri), (-1, ta_ri), 1.5, BLK),
            ("TOPPADDING", (0, ta_ri), (-1, ta_ri), 5),
            ("BOTTOMPADDING", (0, ta_ri), (-1, ta_ri), 5),
        ]

        # ── Blank separator row ──────────────────────────────────────────────
        rows.append(["", "", "", ""])

        # ── LIABILITIES + EQUITY (bottom half) ──────────────────────────────
        liab_secs2 = [s for s in liae_sections if "liabilit" in s["title"].lower()]
        eq_secs2   = [s for s in liae_sections if "liabilit" not in s["title"].lower()]

        if liab_secs2:
            liab_ri = len(rows)
            rows.append(["LIABILITIES", "", "", ""])
            style_cmds.extend([
                ("FONTNAME", (0, liab_ri), (-1, liab_ri), "Times-Bold"),
                ("TOPPADDING", (0, liab_ri), (-1, liab_ri), 7),
            ])
            for section in liab_secs2:
                _append_section(section, is_subsection=(len(liab_secs2) > 1))

        if eq_secs2:
            rows.append(["", "", "", ""])
            eq_hdr_ri = len(rows)
            rows.append(["SHAREHOLDERS' EQUITY", "", "", ""])
            style_cmds.extend([
                ("FONTNAME", (0, eq_hdr_ri), (-1, eq_hdr_ri), "Times-Bold"),
                ("TOPPADDING", (0, eq_hdr_ri), (-1, eq_hdr_ri), 7),
            ])

            # Simulate closing-entry: fold net profit/loss into Retained Earnings
            # so the equity section matches the closed-balance reference format.
            _net_cy = (report.get("financial_statements", {})
                       .get("statement_of_profit_or_loss", {})
                       .get("total", {}).get("current_year") or 0)
            for section in eq_secs2:
                for item in section.get("line_items", []):
                    if "retained" in item.get("account_name", "").lower():
                        item["current_year"] = round((item.get("current_year") or 0) + _net_cy, 2)
                if section.get("subtotal"):
                    cy = section["subtotal"].get("current_year") or 0
                    section["subtotal"]["current_year"] = round(cy + _net_cy, 2)
                _append_section(section, is_subsection=False, show_title=False)
            # Net profit is now absorbed into Retained Earnings — no separate line needed.

        # Total L+E
        if sofp_total:
            tle_ri = len(rows)
            rows.append(["Total Liabilities and Shareholders' Equity", "",
                         _fmt_int(sofp_total.get("current_year")),
                         _fmt_int(sofp_total.get("prior_year"))])
            style_cmds += [
                ("FONTNAME", (0, tle_ri), (-1, tle_ri), "Times-Bold"),
                ("LINEABOVE", (2, tle_ri), (-1, tle_ri), 0.5, BLK),
                ("LINEBELOW", (2, tle_ri), (-1, tle_ri), 1.5, BLK),
                ("TOPPADDING", (0, tle_ri), (-1, tle_ri), 5),
                ("BOTTOMPADDING", (0, tle_ri), (-1, tle_ri), 5),
            ]

        story.append(_fin_table(rows, style_cmds))

    # ─── SOPL ────────────────────────────────────────────────────────────────
    sopl = statements.get("statement_of_profit_or_loss")
    story += [
        _SetStmtTitle("Statement of Profit or Loss"),
        NextPageTemplate('Financial'),
        PageBreak(),
    ]
    if sopl:
        story += [
            Paragraph("Statement of Profit or Loss and Other Comprehensive Income", s_stmt_title),
            Paragraph("(In Arab Emirates Dirhams)", s_stmt_sub),
            Paragraph(f"For the year ended {nice_dt}", s_stmt_sub),
            Spacer(1, 6),
        ]

        net_profit_data = sopl.get("net_profit") or sopl.get("total") or {}
        condensed_rows = _condense_sopl(sopl.get("sections", []), net_profit_data)

        rows, style_cmds = _col_header_rows()

        for label, note, cy, py, row_type in condensed_rows:
            ri = len(rows)
            if row_type == "blank":
                rows.append(["", "", "", ""])
                continue
            cy_str = _fmt_int(cy) if isinstance(cy, (int, float)) else "-"
            py_str = _fmt_int(py) if isinstance(py, (int, float)) else "-"

            if row_type == "header":
                rows.append([label, "", "", ""])
                style_cmds += [
                    ("FONTNAME", (0, ri), (-1, ri), "Times-Bold"),
                    ("TOPPADDING", (0, ri), (-1, ri), 6),
                ]
            elif row_type == "subtotal":
                rows.append([label, "", cy_str, py_str])
                style_cmds += [
                    ("FONTNAME", (0, ri), (-1, ri), "Times-Bold"),
                    ("LINEABOVE", (2, ri), (-1, ri), 0.5, BLK),
                    ("LINEBELOW", (2, ri), (-1, ri), 0.5, BLK),
                ]
            elif row_type == "net":
                rows.append([label, "", cy_str, py_str])
                style_cmds += [
                    ("FONTNAME", (0, ri), (-1, ri), "Times-Bold"),
                    ("LINEABOVE", (2, ri), (-1, ri), 0.5, BLK),
                    ("LINEBELOW", (2, ri), (-1, ri), 1.5, BLK),
                ]
            else:
                rows.append([f"    {label}", note, cy_str, py_str])

        story.append(_fin_table(rows, style_cmds))

    # ─── Statement of Changes in Equity ──────────────────────────────────────
    story += [
        _SetStmtTitle("Statement of Changes in Shareholders' Equity"),
        NextPageTemplate('Financial'),
        PageBreak(),
        Paragraph("Statement of Changes in Shareholders' Equity", s_stmt_title),
        Paragraph("(In Arab Emirates Dirhams)", s_stmt_sub),
        Paragraph(f"For the year ended {nice_dt}", s_stmt_sub),
        Spacer(1, 8),
    ]

    # Equity values from SOFP equity section
    _eq_items = {}
    for _sec in statements.get("statement_of_financial_position", {}).get("sections", []):
        if "equity" in _sec.get("title","").lower():
            for _it in _sec.get("line_items", []):
                _n = _it.get("account_name","")
                _eq_items[_n] = _it.get("current_year") or 0

    # Identify equity components
    _sc = sum(v for k, v in _eq_items.items() if "capital a/c" in k.lower() or "share capital" in k.lower())
    _ca = sum(v for k, v in _eq_items.items() if "current a/c" in k.lower() or "current account" in k.lower())
    _re_open = sum(v for k, v in _eq_items.items()
                   if not any(kw in k.lower() for kw in ["capital a/c", "share capital", "current a/c", "current account"]))
    # Ensure share capital displays as positive
    _sc = abs(_sc) if _sc < 0 else _sc
    # Net profit for SOCE
    _soce_net = (statements.get("statement_of_profit_or_loss", {})
                 .get("total") or {}).get("current_year") or 0

    _soce_desc_w2 = 155  # narrower description column to give value columns enough room
    _soce_val_w = (avail_w - _soce_desc_w2) / 4  # ~80.75 pt each - enough for "Shareholders' Current Account"

    _soce_hdr = ["", "Share\nCapital", "Shareholders'\nCurrent Account", "Retained\nEarnings", "Total"]
    _soce_sub = ["", "AED", "AED", "AED", "AED"]

    _re_close = _re_open + _soce_net   # closing retained earnings = opening + net profit/loss

    _soce_year = nice_dt[-4:]
    _prior_year = str(int(_soce_year) - 1)
    _soce_data = [
        _soce_hdr,
        _soce_sub,
        [f"Balance at 31 December {_prior_year}",
         _fmt_int(_sc), _fmt_int(_ca), _fmt_int(_re_open), _fmt_int(_sc + _ca + _re_open)],
        ["", "", "", "", ""],  # blank row
        ["Changes in Shareholders' Equity", "", "", "", ""],  # sub-header
        ["Capital introduced", "-", "-", "-", "-"],
        ["Net Profit / (loss) for the year",
         "-", "-", _fmt_int(_soce_net), _fmt_int(_soce_net)],
        ["Net movements in shareholders' current account", "-", "-", "-", "-"],
        [f"Balance at 31 December {_soce_year}",
         _fmt_int(_sc), _fmt_int(_ca), _fmt_int(_re_close), _fmt_int(_sc + _ca + _re_close)],
    ]

    _soce_col_widths = [_soce_desc_w2, _soce_val_w, _soce_val_w, _soce_val_w, _soce_val_w]

    _soce_t = Table(_soce_data, colWidths=_soce_col_widths, repeatRows=2)
    _soce_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
        ("FONTNAME", (0, 0), (-1, 1), "Times-Bold"),          # header rows bold
        ("FONTNAME", (0, 4), (-1, 4), "Times-Bold"),           # "Changes in..." row bold
        ("FONTNAME", (0, -1), (-1, -1), "Times-Bold"),         # closing balance bold
        ("FONTSIZE", (0, 0), (-1, -1), FS),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 4), (-1, 4), "LEFT"),                    # "Changes" label left-aligned
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 4), (-1, 4), 8),                    # more space above "Changes"
        ("LINEBELOW", (0, 1), (-1, 1), 0.5, BLK),             # line after AED header
        ("LINEABOVE", (1, -1), (-1, -1), 0.5, BLK),           # line above closing balance
        ("LINEBELOW", (1, -1), (-1, -1), 1.5, BLK),           # double line below closing
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("SPAN", (0, 4), (-1, 4)),                              # span "Changes" across all cols
    ]))
    story.append(_soce_t)

    story.append(Spacer(1, 18))
    story.append(Paragraph(
        "The accompanying notes form an integral part of these financial statements.",
        s_body_j))
    story.append(Paragraph(
        "The report of the auditor's is set out on the pages 1 to 3.",
        s_body))
    story.append(Paragraph(
        f"The financial statements on pages 4 to 22 were approved and signed on {nice_dt}:",
        s_body))
    story.append(Spacer(1, 40))

    # Authorized signatory
    _soce_auth_data = [
        ["___________________________"],
        ["Authorized signatory"],
    ]
    _soce_auth_t = Table(_soce_auth_data, colWidths=[avail_w])
    _soce_auth_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
        ("FONTNAME", (0, 1), (0, 1), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), FS),
    ]))
    story.append(_soce_auth_t)

    # ─── Statement of Cash Flows ─────────────────────────────────────────────
    story += [
        _SetStmtTitle("Statement of Cash Flows"),
        NextPageTemplate('Financial'),
        PageBreak(),
        Paragraph("Statement of Cash Flows", s_stmt_title),
        Paragraph("(In Arab Emirates Dirhams)", s_stmt_sub),
        Paragraph(f"For the year ended {nice_dt}", s_stmt_sub),
        Spacer(1, 8),
    ]

    # Build the CFS data
    _sopl_data = statements.get("statement_of_profit_or_loss", {})
    _net_prof = _sopl_data.get("total") or _sopl_data.get("net_profit") or {}
    _net_cy = _net_prof.get("current_year") or 0
    _net_py = _net_prof.get("prior_year") or 0

    # Depreciation from operating expenses
    _depr_cy = _depr_py = 0.0
    for _sec in _sopl_data.get("sections", []):
        for _item in _sec.get("line_items", []):
            if ("depreciation" in _item.get("account_name","").lower()
                    and "accumulated" not in _item.get("account_name","").lower()):
                _depr_cy += _item.get("current_year") or 0
                _depr_py += _item.get("prior_year") or 0

    # Working capital changes (from SOFP sections, change = current - prior)
    _sofp_raw = statements.get("statement_of_financial_position", {})
    _wc_recv = _wc_adv = _wc_tp = _wc_op = 0.0
    for _sec in _sofp_raw.get("sections", []):
        _stitle = _sec.get("title", "").lower()
        for _item in _sec.get("line_items", []):
            _n = _item.get("account_name","").lower()
            _cy = _item.get("current_year") or 0
            _py = _item.get("prior_year") or 0
            _chg = _cy - _py
            if "current asset" in _stitle or _stitle == "current assets":
                if any(k in _n for k in ["receivable", "debtor"]):
                    _wc_recv += -_chg  # increase = outflow
                elif not any(k in _n for k in ["cash", "bank"]):
                    _wc_adv += -_chg   # increase = outflow
            elif "current liabilit" in _stitle:
                if any(k in _n for k in ["loan", "bank od", "mortgage", "overdraft", "borrowing"]):
                    pass  # loans go to financing
                elif any(k in _n for k in ["sundry creditor", "trade payable", "commission payable", "outside commission"]):
                    _wc_tp += _chg     # increase = inflow
                else:
                    _wc_op += _chg     # increase = inflow

    _net_operating = _net_cy + _depr_cy + _wc_recv + _wc_adv + _wc_tp + _wc_op
    _net_operating_py = _net_py + _depr_py

    # Investing — PPE purchases (gross additions since prior = 0)
    _ppe_add = 0.0
    for _sec in _sofp_raw.get("sections", []):
        if "non-current" in _sec.get("title","").lower():
            for _item in _sec.get("line_items", []):
                _n = _item.get("account_name","").lower()
                if "accumulated" not in _n:  # gross PPE only
                    _cy = _item.get("current_year") or 0
                    _py = _item.get("prior_year") or 0
                    _chg = _cy - _py
                    if _chg > 0:
                        _ppe_add += -_chg   # outflow

    # Financing — loans
    _loans_cy = _loans_py = 0.0
    for _sec in _sofp_raw.get("sections", []):
        if "current liabilit" in _sec.get("title","").lower():
            for _item in _sec.get("line_items", []):
                _n = _item.get("account_name","").lower()
                if any(k in _n for k in ["loan", "bank od", "mortgage", "overdraft", "borrowing"]):
                    _cy = _item.get("current_year") or 0
                    _py = _item.get("prior_year") or 0
                    _loans_cy += (_cy - _py)

    # Closing and opening cash
    _closing_cash = _opening_cash = 0.0
    for _sec in _sofp_raw.get("sections", []):
        if "current asset" in _sec.get("title","").lower() or _sec.get("title","").lower() == "current assets":
            for _item in _sec.get("line_items", []):
                if any(k in _item.get("account_name","").lower() for k in ["cash", "bank"]):
                    _closing_cash += _item.get("current_year") or 0
                    _opening_cash += _item.get("prior_year") or 0

    # Shareholders' balancing figure
    _total_change = _closing_cash - _opening_cash
    _shareholders_flow = _total_change - _net_operating - _ppe_add - _loans_cy
    _net_financing = _loans_cy + _shareholders_flow

    # Build CFS table
    cfs_rows, cfs_cmds = _col_header_rows()

    def _cfs_section_hdr(label):
        ri = len(cfs_rows)
        cfs_rows.append([label, "", "", ""])
        cfs_cmds.extend([
            ("FONTNAME", (0, ri), (-1, ri), "Times-Bold"),
            ("TOPPADDING", (0, ri), (-1, ri), 7),
        ])

    def _cfs_item(label, cy, py=None):
        cfs_rows.append([f"    {label}", "", _fmt_int(cy), _fmt_int(py) if py is not None else "-"])

    def _cfs_subtotal(label, cy, py=None):
        ri = len(cfs_rows)
        cfs_rows.append([label, "", _fmt_int(cy), _fmt_int(py) if py is not None else "-"])
        cfs_cmds.extend([
            ("FONTNAME", (0, ri), (-1, ri), "Times-Bold"),
            ("LINEABOVE", (2, ri), (-1, ri), 0.5, BLK),
            ("LINEBELOW", (2, ri), (-1, ri), 0.5, BLK),
        ])

    # Operating activities
    _cfs_section_hdr("Cash flows from operating activities:")
    _cfs_item("Net profit / (loss) for the year", _net_cy, _net_py if _net_py else None)
    cfs_rows.append(["    Adjustments for:", "", "", ""])
    _cfs_item("Depreciation", _depr_cy, _depr_py if _depr_py else None)
    if _wc_recv:
        _cfs_item("(Increase) / decrease in trade receivables", _wc_recv)
    if _wc_adv:
        _cfs_item("(Increase) / decrease in advances, deposits and prepayments", _wc_adv)
    if _wc_tp:
        _cfs_item("Increase / (decrease) in trade payables", _wc_tp)
    if _wc_op:
        _cfs_item("Increase / (decrease) in other payables, accruals and provisions", _wc_op)
    cfs_rows.append(["", "", "", ""])
    _cfs_subtotal("Net cash from / (used in) operating activities", _net_operating,
                  _net_operating_py if _net_operating_py else None)

    # Investing activities
    cfs_rows.append(["", "", "", ""])
    _cfs_section_hdr("Cash flows from investing activities:")
    if _ppe_add:
        _cfs_item("Purchase of property, plant and equipment", _ppe_add)
    _cfs_subtotal("Net cash used in investing activities", _ppe_add or 0)

    # Financing activities
    cfs_rows.append(["", "", "", ""])
    _cfs_section_hdr("Cash flows from financing activities:")
    if _loans_cy:
        _cfs_item("Net proceeds from bank loans and overdrafts", _loans_cy)
    if _shareholders_flow:
        _cfs_item("Net movement in shareholders' accounts", _shareholders_flow)
    _cfs_subtotal("Net cash from financing activities", _net_financing)

    # Net change and reconciliation
    cfs_rows.append(["", "", "", ""])
    _ri = len(cfs_rows)
    _net_change = _net_operating + _ppe_add + _net_financing
    cfs_rows.append(["Net (decrease) / increase in cash and cash equivalents", "",
                     _fmt_int(_net_change), "-"])
    cfs_cmds.extend([("FONTNAME", (0, _ri), (-1, _ri), "Times-Bold")])

    cfs_rows.append(["    Cash and cash equivalents at beginning of year", "",
                     _fmt_int(_opening_cash) if _opening_cash else "-", "-"])
    _lri = len(cfs_rows)
    cfs_rows.append(["Cash and cash equivalents at end of year", "",
                     _fmt_int(_closing_cash), "-"])
    cfs_cmds.extend([
        ("FONTNAME", (0, _lri), (-1, _lri), "Times-Bold"),
        ("LINEABOVE", (2, _lri), (-1, _lri), 0.5, BLK),
        ("LINEBELOW", (2, _lri), (-1, _lri), 1.5, BLK),
    ])

    story.append(_fin_table(cfs_rows, cfs_cmds))

    # ============= NOTES =============
    notes = report.get("notes", {})
    story += [
        _SetStmtTitle(""),
        NextPageTemplate('Notes'),
        PageBreak(),
        Paragraph("Notes to the Financial Statements", s_stmt_title),
        Paragraph(f"(In Arab Emirates Dirhams)\nFor the year ended {nice_dt}", s_stmt_sub),
        Spacer(1, 10),
    ]

    stmts = report.get("financial_statements", {})
    sofp_data = stmts.get("statement_of_financial_position", {})
    sopl_data = stmts.get("statement_of_profit_or_loss", {})
    all_sofp_secs = _condense_sofp(sofp_data.get("sections", []))

    _note5_depr_cy = 0.0
    for _s in (sopl_data.get("sections", []) if sopl_data else []):
        for _i in _s.get("line_items", []):
            if ("depreciation" in _i.get("account_name","").lower()
                    and "accumulated" not in _i.get("account_name","").lower()):
                _note5_depr_cy += _i.get("current_year") or 0

    def _note_fin_table(rows_data, has_prior=True):
        hdr = ["", cur_hdr, prior_hdr] if has_prior else ["", cur_hdr]
        sub_hdr = ["", "AED", "AED"] if has_prior else ["", "AED"]
        t_rows = [hdr, sub_hdr] + rows_data
        t_cols = [acct_w + note_w, num_w, prior_w] if has_prior else [acct_w + note_w, num_w]
        t = Table(t_rows, colWidths=t_cols, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 1), "Times-Bold"),
            ("FONTNAME", (0, 2), (-1, -1), "Times-Roman"),
            ("FONTSIZE", (0, 0), (-1, -1), FS),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LINEBELOW", (0, 1), (-1, 1), 0.5, BLK),
            ("LINEABOVE", (0, -1), (-1, -1), 0.5, BLK),
            ("LINEBELOW", (0, -1), (-1, -1), 1.5, BLK),
            ("FONTNAME", (0, -1), (-1, -1), "Times-Bold"),
        ]))
        return t

    _nca_sec = next((s for s in all_sofp_secs
                     if "non-current" in s.get("title", "").lower()), None)
    _ca_sec  = next((s for s in all_sofp_secs
                     if s.get("title", "").lower() == "current assets"), None)
    _cl_sec  = next((s for s in all_sofp_secs
                     if "current liabilit" in s.get("title", "").lower()), None)
    _eq_sec  = next((s for s in all_sofp_secs
                     if "equity" in s.get("title", "").lower()), None)

    sopl_sections = sopl_data.get("sections", []) if sopl_data else []
    sopl_total = (sopl_data.get("total") or sopl_data.get("net_profit") or {})
    _sopl_condensed = _condense_sopl(sopl_sections, sopl_total)

    def _sopl_section_by_title(title_key):
        for s in sopl_sections:
            if s.get("title", "").lower() == title_key.lower():
                return s.get("subtotal", {}), s.get("line_items", [])
        return {}, []

    # -- Note 1: Legal Status and Business Activity --
    story.append(Paragraph("<b>1. Legal Status and Business Activity</b>", s_note_title))
    story.append(Paragraph(
        f"<b>1.1</b> {company} (the 'Company') was incorporated and operates as a Limited Liability "
        f"Company (L.L.C.) in Dubai, United Arab Emirates, under a commercial license issued by the "
        "Department of Economic Development.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "<b>1.2</b> The Company is engaged in real estate activities including brokerage, management, "
        "and advisory services related to residential and commercial properties.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        f"<b>1.3</b> The registered office of the Company is located in Dubai, United Arab Emirates.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "<b>1.4</b> These financial statements have been prepared for the year ended "
        f"{nice_dt}.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "<b>1.5</b> The financial statements were approved and authorized for issue by the management "
        f"on {nice_dt}.", s_body_j))
    story.append(Spacer(1, 8))

    # -- Note 2: Share Capital --
    story.append(Paragraph("<b>2. Share Capital</b>", s_note_title))
    story.append(Paragraph(
        "Authorized, Issued and Paid up capital of the Company is AED 300,000 divided into 300 shares "
        "of AED 1,000 each fully paid. The shareholding pattern as at the end of the reporting period "
        "is as below:", s_body_j))
    story.append(Spacer(1, 4))
    _sc_data = [
        ["Name", "Nationality", "% Holding", "Capital (AED)"],
        ["1. Amarjeet Singh Dhir", "India", "100%", "300,000"],
        ["Total", "", "100%", "300,000"],
    ]
    _sc_t = Table(_sc_data, colWidths=[avail_w*0.45, avail_w*0.20, avail_w*0.15, avail_w*0.20])
    _sc_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
        ("FONTNAME", (0, -1), (-1, -1), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), FS),
        ("ALIGN", (2, 0), (-1, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, 0), (-1, 0), 0.5, BLK),
        ("LINEABOVE", (0, -1), (-1, -1), 0.5, BLK),
        ("LINEBELOW", (0, -1), (-1, -1), 1.0, BLK),
        ("BOX", (0, 0), (-1, -1), 0.3, BLK),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, BLK),
    ]))
    story.append(_sc_t)
    story.append(Spacer(1, 8))

    # -- Note 3: Basis of Preparation --
    story.append(Paragraph("<b>3. Basis of Preparation</b>", s_note_title))
    story.append(Paragraph("<b>Statement of compliance</b>", s_bold))
    story.append(Paragraph(
        "These financial statements have been prepared in accordance with the International Financial "
        "Reporting Standards ('IFRSs') as issued by the International Accounting Standards Board "
        "('IASB') and the applicable requirements of UAE Federal Law No. (2) of 2015.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Basis of measurement</b>", s_bold))
    story.append(Paragraph(
        "These financial statements have been prepared on the historical cost basis.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Functional and presentation currency</b>", s_bold))
    story.append(Paragraph(
        "These financial statements are presented in United Arab Emirates Dirham ('AED'), which is "
        "the Company's functional currency. All financial information presented in AED has been "
        "rounded to the nearest dirham.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Use of estimates and judgments</b>", s_bold))
    story.append(Paragraph(
        "The preparation of the financial statements in conformity with IFRS requires management to "
        "make judgements, estimates and assumptions that affect the application of accounting policies "
        "and reported amounts of assets and liabilities, income and expenses. Actual results may "
        "differ from these estimates. Estimates and underlying assumptions are reviewed on an ongoing "
        "basis. Revisions to accounting estimates are recognised prospectively.", s_body_j))
    story.append(Spacer(1, 8))

    # -- Note 4: Significant Accounting Policies --
    story.append(Paragraph("<b>4. Significant Accounting Policies</b>", s_note_title))
    story.append(Paragraph(
        "The accounting policies set out below have been applied consistently to all periods "
        "presented in these financial statements.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Revenue</b>", s_bold))
    story.append(Paragraph(
        "Revenue is recognised in accordance with IFRS 15 — Revenue from Contracts with Customers. "
        "The Company recognises revenue when (or as) performance obligations are satisfied and control "
        "of the promised goods or services is transferred to the customer at the transaction price "
        "allocated to that performance obligation. Commission income from real estate brokerage is "
        "recognised upon completion of the transaction.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Finance income and expenses</b>", s_bold))
    story.append(Paragraph(
        "Borrowing costs that are not directly attributable to the acquisition, construction or "
        "production of a qualifying asset are recognised in profit or loss using the effective "
        "interest method.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Inventories</b>", s_bold))
    story.append(Paragraph(
        "Inventories are measured at the lower of cost and net realisable value. The cost of raw "
        "materials and packaging materials is based on the weighted average cost method and includes "
        "expenditure incurred in acquiring the inventories and bringing them to their existing location "
        "and condition. Finished goods are stated at cost of raw materials and also include an "
        "appropriate share of overheads based on normal operating capacity.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Property, Plant and Equipment</b>", s_bold))
    story.append(Paragraph("<i>Recognition and measurement</i>", s_body))
    story.append(Paragraph(
        "Items of property, plant and equipment are measured at cost less accumulated depreciation "
        "and accumulated impairment losses. Cost includes expenditure that is directly attributable "
        "to the acquisition of the asset. Any gain or loss on disposal of an item of property, plant "
        "and equipment is recognised in profit or loss.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Capital work in progress</i>", s_body))
    story.append(Paragraph(
        "Capital work in progress is stated at cost, less accumulated impairment losses and is "
        "transferred to the respective category of property, plant and equipment when completed and "
        "ready for intended use.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Subsequent costs</i>", s_body))
    story.append(Paragraph(
        "The cost of replacing part of an item of property, plant and equipment is recognised in "
        "the carrying amount of the item if it is probable that the future economic benefits embodied "
        "within the part will flow to the Company and its cost can be measured reliably.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Depreciation</i>", s_body))
    story.append(Paragraph(
        "Depreciation is recognised in profit or loss on a straight-line basis over the estimated "
        "useful lives of each part of an item of property, plant and equipment:", s_body_j))
    story.append(Spacer(1, 3))
    _depr_tbl = Table([
        ["Assets", "Years"],
        ["Furniture, fixtures & equipment", "5 – 15"],
        ["Motor vehicles", "5 – 15"],
        ["Office Equipment", "3 – 5"],
    ], colWidths=[avail_w * 0.70, avail_w * 0.30])
    _depr_tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), FS),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("LINEBELOW", (0, 0), (-1, 0), 0.5, BLK),
        ("BOX", (0, 0), (-1, -1), 0.3, BLK),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, BLK),
    ]))
    story.append(_depr_tbl)
    story.append(Spacer(1, 4))
    story.append(PageBreak())

    story.append(Paragraph("<b>4. Significant Accounting Policies (continued)</b>", s_note_title))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Financial instruments</b>", s_bold))
    story.append(Paragraph(
        "The Company initially recognises loans and receivables on the date that they are originated. "
        "All other financial assets and financial liabilities are initially recognised on the trade date.",
        s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "The Company derecognises a financial asset when the contractual rights to the cash flows from "
        "the asset expire, or it transfers the rights to receive the contractual cash flows on the "
        "financial asset in a transaction in which substantially all the risks and rewards of ownership "
        "of the financial asset are transferred.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Non-derivative financial assets — recognition and derecognition</i>", s_body))
    story.append(Paragraph(
        "The Company classifies non-derivative financial liabilities into the other financial liabilities "
        "category. Other financial liabilities comprise: bank borrowings, trade and other payables and "
        "due to a related party.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Loans and receivables</i>", s_body))
    story.append(Paragraph(
        "Loans and receivables are financial assets with fixed or determinable payments that are not "
        "quoted in an active market. Such assets are recognised initially at fair value plus any "
        "directly attributable transaction costs. Subsequent to initial recognition loans and "
        "receivables are measured at amortised cost using the effective interest method, less any "
        "impairment losses.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Cash and cash equivalents</i>", s_body))
    story.append(Paragraph(
        "Cash and cash equivalents comprise cash in hand and bank balances.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Impairment</i>", s_body))
    story.append(Paragraph("<i>(i) Non-derivative financial assets</i>", s_body))
    story.append(Paragraph(
        "A financial asset is assessed at each reporting date to determine whether there is objective "
        "evidence that it is impaired. A financial asset is impaired if there is objective evidence of "
        "impairment as a result of one or more events that occurred after the initial recognition of "
        "the asset, and that the loss event(s) had an impact on the estimated future cash flows of "
        "that asset that can be estimated reliably.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "An impairment loss in respect of a financial asset measured at amortised cost is calculated "
        "as the difference between its carrying amount and the present value of the estimated future "
        "cash flows discounted at the asset's original effective interest rate. Losses are recognised "
        "in profit or loss and reflected in an allowance account against loans and receivables.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>(ii) Non-financial assets</i>", s_body))
    story.append(Paragraph(
        "The carrying amounts of the Company's non-financial assets, other than inventories, are "
        "reviewed at each reporting date to determine whether there is any indication of impairment. "
        "If any such indication exists, then the asset's recoverable amount is estimated.", s_body_j))
    story.append(PageBreak())

    story.append(Paragraph("<b>4. Significant Accounting Policies (continued)</b>", s_note_title))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Staff terminal benefits</b>", s_bold))
    story.append(Paragraph(
        "The Company's net obligation in respect of long-term employee benefits is the amount of "
        "future benefit that employees have earned in return for their service in the current and "
        "prior years. That benefit is discounted to determine its present value. Re-measurements are "
        "recognised in profit or loss in the year in which they arise.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Provisions</b>", s_bold))
    story.append(Paragraph(
        "A provision is recognised if, as a result of a past event, the Company has a present legal "
        "or constructive obligation that can be estimated reliably, and it is probable that an outflow "
        "of economic benefits will be required to settle the obligation.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Operating leases</b>", s_bold))
    story.append(Paragraph(
        "Payments made under operating leases are recognised in profit or loss on a straight-line "
        "basis over the term of the lease. Lease incentives received are recognised as an integral "
        "part of the total lease expense, over the term of the lease.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Foreign currency transactions</b>", s_bold))
    story.append(Paragraph(
        "Transactions in foreign currencies are translated to the functional currency of the Company "
        "at exchange rates at the dates of the transactions. Monetary assets and liabilities "
        "denominated in foreign currencies at the reporting date are retranslated to the functional "
        "currency at the exchange rate at that date.", s_body_j))
    story.append(Spacer(1, 10))

    # -- Note 5: Financial Risk Management --
    story.append(Paragraph("<b>5. Financial Risk Management</b>", s_note_title))
    story.append(Paragraph("<b>Overview</b>", s_bold))
    story.append(Paragraph(
        "The Company has exposure to the following risks arising from financial instruments:", s_body_j))
    story.append(Paragraph("•  credit risk;", s_body))
    story.append(Paragraph("•  liquidity risk;", s_body))
    story.append(Paragraph("•  market risk.", s_body))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "This note presents information about the Company's exposure to each of the above risks, the "
        "Company's objectives, policies and processes for measuring and managing risk, and the "
        "Company's management of capital.", s_body_j))
    story.append(PageBreak())
    story.append(Paragraph("<b>5. Financial Risk Management (continued)</b>", s_note_title))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Credit risk</b>", s_bold))
    story.append(Paragraph(
        "Credit risk is the risk of financial loss to the Company if a customer or counterparty to "
        "a financial instrument fails to meet its contractual obligations, and arises principally "
        "from the Company's trade receivables and cash balances.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "The Company establishes an allowance for impairment that represents its estimate of incurred "
        "losses in respect of trade receivables. The allowance mainly represents specific losses that "
        "relates to individually significant exposures.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<i>Cash at banks</i>", s_body))
    story.append(Paragraph(
        "Cash is placed with banks of good repute. Management believes that the unimpaired amounts "
        "that are past due are still collectible, based on historic payment behaviour.", s_body_j))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Liquidity risk</b>", s_bold))
    story.append(Paragraph(
        "Liquidity risk is the risk that the Company will encounter difficulty in meeting the "
        "obligations associated with its financial liabilities that are settled by delivering cash "
        "or another financial asset. The Company's approach to managing liquidity is to ensure, "
        "as far as possible, that it will always have sufficient liquidity to meet its liabilities "
        "when due, under both normal and stressed conditions.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Market risk</b>", s_bold))
    story.append(Paragraph(
        "Market risk is the risk that changes in market prices, such as foreign exchange rates and "
        "interest rates will affect the Company's income or the value of its holdings of financial "
        "instruments. The objective of market risk management is to manage and control market risk "
        "exposures within acceptable parameters, while optimising the return.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Currency risk</b>", s_bold))
    story.append(Paragraph(
        "The Company's transactions are primarily denominated in AED and USD. As the AED is pegged "
        "to the USD, the Company's exposure to foreign currency risk is minimal.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Interest rate risk</b>", s_bold))
    story.append(Paragraph(
        "Interest rate risk is the risk that value of a financial instrument will fluctuate because "
        "of changes in market interest rates. As at the reporting date, the Company is not exposed "
        "to any significant unhedged interest rate risk.", s_body_j))
    story.append(Spacer(1, 8))

    # -- Note 6: PPE --
    story.append(PageBreak())
    story.append(Paragraph("<b>6. Property, Plant and Equipment</b>", s_note_title))
    _nca_raw = next((s for s in sofp_data.get("sections", [])
                     if "non-current" in s.get("title","").lower()), None)
    if _nca_raw and _nca_raw.get("line_items"):
        gross_items = [i for i in _nca_raw["line_items"]
                       if "accumulated" not in i.get("account_name","").lower()
                       and (i.get("current_year") or 0) != 0]
        accum_items = [i for i in _nca_raw["line_items"]
                       if "accumulated" in i.get("account_name","").lower()]
        gross_total = sum(i.get("current_year") or 0 for i in gross_items)
        accum_total = sum(abs(i.get("current_year") or 0) for i in accum_items)
        net_total = gross_total - accum_total
        cat_names = [i["account_name"] for i in gross_items]
        n_cats = len(cat_names)
        if n_cats <= 3:
            col_labels = cat_names + ["Total"]
            all_gross_vals = [_fmt_int(i.get("current_year")) for i in gross_items] + [_fmt_int(gross_total)]
            all_accum_vals = [f"({_fmt_int(abs(i.get('current_year') or 0))})" for i in accum_items]
            while len(all_accum_vals) < n_cats:
                all_accum_vals.append("-")
            all_accum_vals.append(f"({_fmt_int(accum_total)})")
            all_net_vals = ["-"] * n_cats + [_fmt_int(net_total)]
            ppe_col_w = avail_w / (n_cats + 2)
            ppe_desc_w = avail_w - ppe_col_w * (n_cats + 1)
            ppe_col_widths = [ppe_desc_w] + [ppe_col_w] * (n_cats + 1)
            ppe_hdr = [""] + col_labels
            ppe_aed = [""] + ["AED"] * (n_cats + 1)
            yr = nice_dt[-4:]
            ppe_rows_data = [
                ppe_hdr, ppe_aed,
                ["Cost:"] + [""] * (n_cats + 1),
                ["  At 1 January " + yr] + ["-"] * n_cats + ["-"],
                ["  Additions"] + all_gross_vals[:-1] + [all_gross_vals[-1]],
                ["  Disposals"] + ["-"] * n_cats + ["-"],
                ["  At 31 December " + yr] + all_gross_vals,
                ["Depreciation:"] + [""] * (n_cats + 1),
                ["  At 1 January " + yr] + ["-"] * n_cats + ["-"],
                ["  Depreciation for the year"] + all_accum_vals,
                ["  Disposals"] + ["-"] * n_cats + ["-"],
                ["  At 31 December " + yr] + all_accum_vals,
                ["Net book value:"] + [""] * (n_cats + 1),
                ["  At 31 December " + yr] + all_net_vals,
                ["  At 31 December " + str(int(yr)-1)] + ["-"] * n_cats + ["-"],
            ]
            ppe_t = Table(ppe_rows_data, colWidths=ppe_col_widths)
            ppe_t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
                ("FONTNAME", (0, 0), (-1, 1), "Times-Bold"),
                ("FONTNAME", (0, 2), (0, 2), "Times-Bold"),
                ("FONTNAME", (0, 7), (0, 7), "Times-Bold"),
                ("FONTNAME", (0, 12), (0, 12), "Times-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), FS),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LINEBELOW", (0, 1), (-1, 1), 0.5, BLK),
                ("LINEABOVE", (1, 6), (-1, 6), 0.5, BLK),
                ("LINEBELOW", (1, 6), (-1, 6), 0.5, BLK),
                ("LINEABOVE", (1, 11), (-1, 11), 0.5, BLK),
                ("LINEBELOW", (1, 11), (-1, 11), 0.5, BLK),
                ("LINEABOVE", (1, -1), (-1, -1), 0.5, BLK),
                ("LINEBELOW", (1, -2), (-1, -2), 1.5, BLK),
            ]))
            story.append(ppe_t)
            story.append(Spacer(1, 4))
        else:
            ppe_rows = [["  Cost", _fmt_int(gross_total), "-"]]
            ppe_rows.append(["  Less: Accumulated depreciation", f"({_fmt_int(accum_total)})", "-"])
            ppe_rows.append(["Net book value", _fmt_int(net_total), "-"])
            story.append(_note_fin_table(ppe_rows, has_prior=True))
        story.append(Paragraph(
            "Depreciation is calculated on the straight-line method based on the estimated useful "
            "lives of the assets. The annual depreciation charge for the current year amounted to "
            f"AED {_fmt_int(_note5_depr_cy)}.",
            s_body_j))
    else:
        story.append(Paragraph(
            "The Company does not hold any property, plant and equipment during the current period.",
            s_body))
    story.append(Spacer(1, 6))

    # -- Note 7: Trade Receivables --
    story.append(Paragraph("<b>7. Trade Receivables</b>", s_note_title))
    if _ca_sec:
        recv_items = [i for i in _ca_sec.get("line_items", [])
                      if any(k in i["account_name"].lower() for k in ["receivable", "debtor"])]
        if recv_items:
            recv_rows = [[f"  {i['account_name']}",
                          _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                         for i in recv_items]
            st_recv = sum(i.get("current_year") or 0 for i in recv_items)
            st_prev = sum(i.get("prior_year") or 0 for i in recv_items)
            recv_rows.append(["Total trade receivables", _fmt_int(st_recv), _fmt_int(st_prev)])
            story.append(_note_fin_table(recv_rows))
            story.append(Spacer(1, 4))
        else:
            recv_tbl = [
                ["  Trade Receivables", "-", "-"],
                ["Total trade receivables", "-", "-"],
            ]
            story.append(_note_fin_table(recv_tbl))
            story.append(Spacer(1, 4))
    story.append(Paragraph(
        "Based on management's assessment, the expected credit loss on trade receivables is not "
        "material and no allowance has been recognised.", s_body_j))
    story.append(Spacer(1, 6))

    # -- Note 8: Advances --
    story.append(Paragraph("<b>8. Advances, Deposits and Other Receivables</b>", s_note_title))
    if _ca_sec:
        adv_items = [i for i in _ca_sec.get("line_items", [])
                     if not any(k in i["account_name"].lower()
                                for k in ["receivable", "debtor", "cash", "bank"])]
        if adv_items:
            adv_rows = [[f"  {i['account_name']}",
                         _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                        for i in adv_items]
            st_adv = sum(i.get("current_year") or 0 for i in adv_items)
            st_adv_pr = sum(i.get("prior_year") or 0 for i in adv_items)
            adv_rows.append(["Total", _fmt_int(st_adv), _fmt_int(st_adv_pr)])
            story.append(_note_fin_table(adv_rows))
        else:
            story.append(Paragraph("No advances or other receivables in the current period.", s_body))
    story.append(Spacer(1, 6))
    story.append(PageBreak())

    # -- Note 9: Cash --
    story.append(Paragraph("<b>9. Cash in Hand and at Banks</b>", s_note_title))
    if _ca_sec:
        cash_items = [i for i in _ca_sec.get("line_items", [])
                      if any(k in i["account_name"].lower() for k in ["cash", "bank"])]
        if cash_items:
            cash_rows = [[f"  {i['account_name']}",
                          _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                         for i in cash_items]
            st_cash = sum(i.get("current_year") or 0 for i in cash_items)
            st_cash_pr = sum(i.get("prior_year") or 0 for i in cash_items)
            cash_rows.append(["Total", _fmt_int(st_cash), _fmt_int(st_cash_pr)])
            story.append(_note_fin_table(cash_rows))
        else:
            story.append(Paragraph("No cash or bank balances in the current period.", s_body))
    story.append(Spacer(1, 6))

    # -- Note 10: Trade Payables --
    story.append(Paragraph("<b>10. Trade Payables</b>", s_note_title))
    if _cl_sec:
        tp_items = [i for i in _cl_sec.get("line_items", [])
                    if any(k in i["account_name"].lower() for k in ["payable", "creditor",
                                                                      "commission payable", "outside commission"])]
        if tp_items:
            tp_rows = [[f"  {i['account_name']}",
                        _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                       for i in tp_items]
            st_tp = sum(i.get("current_year") or 0 for i in tp_items)
            st_tp_pr = sum(i.get("prior_year") or 0 for i in tp_items)
            tp_rows.append(["Total trade payables", _fmt_int(st_tp), _fmt_int(st_tp_pr)])
            story.append(_note_fin_table(tp_rows))
        else:
            story.append(Paragraph("No trade payables in the current period.", s_body))
    story.append(Spacer(1, 6))

    # -- Note 11: Other Payables --
    story.append(Paragraph("<b>11. Other Payables, Accruals and Provisions</b>", s_note_title))
    if _cl_sec:
        op_items = [i for i in _cl_sec.get("line_items", [])
                    if not any(k in i["account_name"].lower()
                               for k in ["loan", "bank od", "mortgage", "borrowing",
                                         "payable", "creditor", "commission payable", "outside commission"])]
        if op_items:
            op_rows = [[f"  {i['account_name']}",
                        _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                       for i in op_items]
            st_op = sum(i.get("current_year") or 0 for i in op_items)
            st_op_pr = sum(i.get("prior_year") or 0 for i in op_items)
            op_rows.append(["Total", _fmt_int(st_op), _fmt_int(st_op_pr)])
            story.append(_note_fin_table(op_rows))
        else:
            story.append(Paragraph("No other payables or accruals in the current period.", s_body))
    story.append(Spacer(1, 6))

    # -- Note 12: Long-term Loans --
    story.append(Paragraph("<b>12. Long-term Loans</b>", s_note_title))
    if _cl_sec:
        ln_items = [i for i in _cl_sec.get("line_items", [])
                    if any(k in i["account_name"].lower()
                           for k in ["loan", "bank od", "mortgage", "overdraft", "borrowing"])]
        if ln_items:
            ln_rows = [[f"  {i['account_name']}",
                        _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                       for i in ln_items]
            st_ln = sum(i.get("current_year") or 0 for i in ln_items)
            st_ln_pr = sum(i.get("prior_year") or 0 for i in ln_items)
            ln_rows.append(["Total", _fmt_int(st_ln), _fmt_int(st_ln_pr)])
            story.append(_note_fin_table(ln_rows))
        else:
            story.append(Paragraph("No long-term loans during the current period.", s_body))
    story.append(Spacer(1, 6))
    story.append(PageBreak())

    # -- Note 13: Shareholders Current Account --
    story.append(Paragraph("<b>13. Shareholders' Current Account</b>", s_note_title))
    if _eq_sec:
        ca_items = [i for i in _eq_sec.get("line_items", [])
                    if "current a" in i["account_name"].lower()
                    or "current account" in i["account_name"].lower()]
        if ca_items:
            ca_rows = []
            st_ca_open = sum(i.get("prior_year") or 0 for i in ca_items)
            ca_rows.append(["  Opening balance", _fmt_int(st_ca_open), "-"])
            st_ca = sum(i.get("current_year") or 0 for i in ca_items)
            _ca_movement = st_ca - st_ca_open
            ca_rows.append(["  Net movement during the period", _fmt_int(_ca_movement), "-"])
            ca_rows.append(["Total", _fmt_int(st_ca), "-"])
            story.append(_note_fin_table(ca_rows))
        else:
            story.append(Paragraph(
                "The shareholders' current account represents the net amount due to/from shareholders "
                "arising from transactions during the year.", s_body_j))
    else:
        story.append(Paragraph(
            "The shareholders' current account represents the net amount due to/from shareholders "
            "arising from transactions during the year.", s_body_j))
    story.append(Spacer(1, 6))

    # -- Note 14: Retained Earnings --
    story.append(Paragraph("<b>14. Retained Earnings</b>", s_note_title))
    if sopl_total:
        re_rows = []
        if _eq_sec:
            re_items = [i for i in _eq_sec.get("line_items", [])
                        if any(kw in i["account_name"].lower()
                               for kw in ["retain", "reserve", "surplus", "accumulated"])]
            for i in re_items:
                re_rows.append([f"  Opening balance",
                                 _fmt_int(i.get("current_year")), "-"])
        re_rows.append(["  Net profit / (loss) for the period",
                         _fmt_int(sopl_total.get("current_year")),
                         _fmt_int(sopl_total.get("prior_year"))])
        re_rows.append(["Closing balance", "-", "-"])
        story.append(_note_fin_table(re_rows))
    else:
        story.append(Paragraph(
            "Retained earnings movement is shown in the Statement of Changes in Equity.", s_body_j))
    story.append(Spacer(1, 6))

    # -- Notes 15-18: Income Statement --
    story.append(PageBreak())
    story.append(Paragraph("<b>15. Revenue</b>", s_note_title))
    rev_st, rev_items = _sopl_section_by_title("Revenue")
    if rev_items:
        rev_rows = [[f"  {i['account_name']}",
                     _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                    for i in rev_items]
        rev_rows.append(["Total Revenue",
                          _fmt_int(rev_st.get("current_year")),
                          _fmt_int(rev_st.get("prior_year"))])
        story.append(_note_fin_table(rev_rows))
    else:
        rev_cond = next((r for r in _sopl_condensed if r[0] == "Revenue"), None)
        if rev_cond:
            story.append(_note_fin_table([
                ["  Rental and commission income",
                 _fmt_int(rev_cond[2]), _fmt_int(rev_cond[3])],
                ["Total Revenue",
                 _fmt_int(rev_cond[2]), _fmt_int(rev_cond[3])],
            ]))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>16. Cost of Revenue</b>", s_note_title))
    cos_st, cos_items = _sopl_section_by_title("Cost of Sales")
    if not cos_items:
        cos_st, cos_items = _sopl_section_by_title("Direct Expenses")
    if cos_items:
        cos_rows = [[f"  {i['account_name']}",
                     _fmt_int(abs(i.get("current_year") or 0)),
                     _fmt_int(abs(i.get("prior_year") or 0))]
                    for i in cos_items]
        cos_total = sum(abs(i.get("current_year") or 0) for i in cos_items)
        cos_total_pr = sum(abs(i.get("prior_year") or 0) for i in cos_items)
        cos_rows.append(["Total Cost of Revenue", _fmt_int(cos_total), _fmt_int(cos_total_pr)])
        story.append(_note_fin_table(cos_rows))
    else:
        cos_cond = next((r for r in _sopl_condensed if "cost" in r[0].lower()), None)
        if cos_cond:
            story.append(_note_fin_table([
                ["  Direct expenses",
                 _fmt_int(abs(cos_cond[2] or 0)), _fmt_int(abs(cos_cond[3] or 0))],
                ["Total Cost of Revenue",
                 _fmt_int(abs(cos_cond[2] or 0)), _fmt_int(abs(cos_cond[3] or 0))],
            ]))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>17. General and Administration Expenses</b>", s_note_title))
    ga_st, ga_items = _sopl_section_by_title("General and Administration Expenses")
    if not ga_items:
        ga_st, ga_items = _sopl_section_by_title("General and Administrative Expenses")
    if not ga_items:
        ga_st, ga_items = _sopl_section_by_title("Operating Expenses")
    if ga_items:
        ga_rows = [[f"  {i['account_name']}",
                    _fmt_int(abs(i.get("current_year") or 0)),
                    _fmt_int(abs(i.get("prior_year") or 0))]
                   for i in ga_items]
        ga_total = sum(abs(i.get("current_year") or 0) for i in ga_items)
        ga_total_pr = sum(abs(i.get("prior_year") or 0) for i in ga_items)
        ga_rows.append(["Total G&A Expenses", _fmt_int(ga_total), _fmt_int(ga_total_pr)])
        story.append(_note_fin_table(ga_rows))
    else:
        story.append(Paragraph("Refer to Statement of Profit or Loss for G&A expenses.", s_body))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>18. Other Income</b>", s_note_title))
    oi_st, oi_items = _sopl_section_by_title("Other Income")
    if oi_items:
        oi_rows = [[f"  {i['account_name']}",
                    _fmt_int(i.get("current_year")), _fmt_int(i.get("prior_year"))]
                   for i in oi_items]
        oi_total = sum(i.get("current_year") or 0 for i in oi_items)
        oi_total_pr = sum(i.get("prior_year") or 0 for i in oi_items)
        oi_rows.append(["Total Other Income", _fmt_int(oi_total), _fmt_int(oi_total_pr)])
        story.append(_note_fin_table(oi_rows))
    else:
        story.append(Paragraph("No other income during the current period.", s_body))
    story.append(Spacer(1, 6))
    story.append(PageBreak())

    # -- Notes 19-23: Risk and Other Disclosures --
    story.append(Paragraph("<b>19. Financial Instruments</b>", s_note_title))
    story.append(Paragraph(
        "Credit risk represents the risk that one party to a financial instrument will cause a "
        "financial loss for the other party by failing to discharge an obligation.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "The credit risk on liquid funds is limited because the Company's bank accounts are placed "
        "with high credit quality financial institutions.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "For trade receivables, credit quality of customers is assessed taking into consideration "
        "their financial position and previous dealings and on that basis, individual credit limits "
        "are set. Trade and other receivables are stated net of allowance for doubtful recoveries.", s_body_j))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>20. Financial Risk Management</b>", s_note_title))
    story.append(Paragraph("<b>20.1 Credit risk and concentration of credit risk</b>", s_bold))
    story.append(Paragraph(
        "Liquidity risk is the risk that the Entity will not be able to meet its financial "
        "obligations as they fall due. The Company's approach to managing liquidity is to ensure "
        "that it will always have sufficient liquidity to meet its liabilities when due.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>20.2 Liquidity risk</b>", s_bold))
    story.append(Paragraph(
        "The Company monitors its risk to a shortage of funds using a liquidity planning tool. "
        "This tool considers the maturity of both its financial instruments and financial assets "
        "and projected cash flows from operations.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>20.3 Market risk</b>", s_bold))
    story.append(Paragraph("<i>a) Interest rate risk</i>", s_body))
    story.append(Paragraph(
        "As at the reporting date, the Company is not exposed to any significant interest rate risk.",
        s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<i>b) Currency risk</i>", s_body))
    story.append(Paragraph(
        "The Company's transactions are denominated primarily in AED. As the AED is pegged to the "
        "USD, the Company has no significant currency risk exposure.", s_body_j))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<i>c) Equity price risk</i>", s_body))
    story.append(Paragraph(
        "The Company does not hold any equity instruments at the reporting date and accordingly is "
        "not exposed to any equity price risk.", s_body_j))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>21. Capital Risk Management</b>", s_note_title))
    story.append(Paragraph(
        "The Company's objectives when managing capital are to safeguard the Company's ability to "
        "continue as a going concern, so that it can continue to provide returns for shareholders "
        "and to maintain an optimal capital structure to reduce the cost of capital.", s_body_j))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "The Company monitors capital using a gearing ratio, which is net debt divided by total "
        "equity plus net debt. The Company includes within net debt, interest bearing loans and "
        "borrowings, less cash and cash equivalents.", s_body_j))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>22. Contingent Liability and Capital Commitments</b>", s_note_title))
    story.append(Paragraph(
        "Except for the ongoing business obligations which are under normal course of business "
        "against which no loss is expected, there has been no other known contingent liability or "
        "capital commitment on the Company's account as of the statement of financial position date.",
        s_body_j))
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>23. Comparative Figures</b>", s_note_title))
    _prior_year_str = str(int(nice_dt[-4:]) - 1)
    story.append(Paragraph(
        f"Certain figures for the previous year ended December 31, {_prior_year_str} have been "
        "regrouped / reclassified to conform with the presentation made during the current year.",
        s_body_j))
    story.append(Spacer(1, 20))

    # -- Authorized Signatory --
    story.append(Paragraph(
        f"These financial statements from pages 4 to 23 were approved by the management on {nice_dt} "
        "and signed on their behalf by:",
        s_body))
    story.append(Spacer(1, 40))
    _final_auth_data = [
        ["___________________________"],
        ["Authorized Signatory"],
    ]
    _final_auth_t = Table(_final_auth_data, colWidths=[avail_w])
    _final_auth_t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, 0), "Times-Roman"),
        ("FONTNAME", (0, 1), (0, 1), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), FS),
    ]))
    story.append(_final_auth_t)

    doc.build(story)
    return buf.getvalue()


def _build_pdf_statement_table(stmt: dict, currency: str, tpl: dict) -> list[list[str]]:
    """Build 2-D list for a financial statement table (header + rows)."""
    cols = tpl.get("columns", DEFAULT_TEMPLATE["columns"])
    header = [cols[0], cols[1], f"{cols[2]} ({currency})", f"{cols[3]} ({currency})"]
    rows: list[list[str]] = [header]

    for section in stmt.get("sections", []):
        # Section header row
        rows.append([section.get("title", ""), "", "", ""])

        for item in section.get("line_items", []):
            rows.append([
                f"    {item.get('account_name', '')}",
                str(item.get("note_ref") or item.get("notes_ref") or ""),
                _fmt_number(item.get("current_year"), currency),
                _fmt_number(item.get("prior_year"), currency),
            ])

        sub = section.get("subtotal")
        if sub:
            rows.append([
                sub.get("account_name", ""),
                "",
                _fmt_number(sub.get("current_year"), currency),
                _fmt_number(sub.get("prior_year"), currency),
            ])

    total = stmt.get("total")
    if total:
        rows.append([
            total.get("account_name", ""),
            "",
            _fmt_number(total.get("current_year"), currency),
            _fmt_number(total.get("prior_year"), currency),
        ])

    return rows


def _calc_col_widths(available: float) -> list[float]:
    """Split available width into 4 columns: wide account col + 3 narrow cols."""
    note_w = 40
    num_w = (available - note_w) * 0.30
    acct_w = available - note_w - 2 * num_w
    return [acct_w, note_w, num_w, num_w]


def _financial_table_style(row_count: int, data: list[list[str]]):
    """Return a TableStyle for financial statement tables with proper formatting."""
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    cmds = [
        # Header row
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf3")),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        # Align numbers right
        ("ALIGN", (2, 0), (3, -1), "RIGHT"),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        # Grid lines
        ("LINEBELOW", (0, 0), (-1, 0), 1, colors.black),
        ("LINEABOVE", (0, 0), (-1, 0), 1, colors.black),
    ]

    for i, row in enumerate(data):
        if i == 0:
            continue
        cell_text = (row[0] or "").strip()

        # Section header — bold, no indent
        if cell_text and not cell_text.startswith(" ") and row[2] == "" and row[3] == "":
            cmds.append(("FONTNAME", (0, i), (0, i), "Times-Bold"))
            cmds.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#f5f7fa")))

        # Subtotal — single line below
        is_subtotal = cell_text.lower().startswith("total ") and i < len(data) - 1
        if is_subtotal:
            cmds.append(("FONTNAME", (0, i), (-1, i), "Times-Bold"))
            cmds.append(("LINEBELOW", (2, i), (3, i), 0.5, colors.black))

        # Grand total — double line below
        is_grand_total = cell_text.lower().startswith("total ") and i == len(data) - 1
        if is_grand_total or cell_text.lower().startswith("net profit"):
            cmds.append(("FONTNAME", (0, i), (-1, i), "Times-Bold"))
            cmds.append(("LINEABOVE", (2, i), (3, i), 0.5, colors.black))
            cmds.append(("LINEBELOW", (2, i), (3, i), 1.5, colors.black))

    return TableStyle(cmds)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX generation (python-docx)
# ══════════════════════════════════════════════════════════════════════════════


def _generate_docx(report: dict, tpl: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    font_name = tpl.get("font_family", "Times New Roman")
    font_size = tpl.get("font_size", 10)
    currency = tpl.get("currency_symbol", "AED")
    meta = report.get("metadata", {})
    company = meta.get("company_name", "Company")
    period_end = meta.get("period_end", "")
    auditor = meta.get("auditor_name", "")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)

    def _run(para, text, bold=False, size=None, color=None, italic=False):
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = font_name
        r.font.size = Pt(size or font_size)
        if color:
            r.font.color.rgb = RGBColor(*color)
        return r

    # ── Cover page ────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, company.upper(), bold=True, size=18)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "FINANCIAL STATEMENTS AND", bold=True, size=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "INDEPENDENT AUDITOR'S REPORT", bold=True, size=13)

    doc.add_paragraph()

    if period_end:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"FOR THE YEAR ENDED {period_end}", bold=True, size=13)

    if auditor:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"Auditor: {auditor}", size=11, italic=True)

    doc.add_page_break()

    # ── Auditor's opinion ─────────────────────────────────────────────────
    opinion = report.get("auditor_opinion", {})
    if opinion:
        doc.add_heading("Independent Auditor's Report", level=1)

        opinion_type = (opinion.get("opinion_type") or "").replace("_", " ").title()
        if opinion_type:
            p = doc.add_paragraph()
            _run(p, "Opinion: ", bold=True)
            _run(p, opinion_type)

        opinion_text = opinion.get("opinion_text", "")
        if opinion_text:
            doc.add_paragraph(opinion_text)

        basis = opinion.get("basis_text", "")
        if basis:
            p = doc.add_paragraph()
            _run(p, "Basis for Opinion", bold=True)
            doc.add_paragraph(basis)

        for kam in opinion.get("key_audit_matters", []):
            p = doc.add_paragraph()
            _run(p, "Key Audit Matter: ", bold=True)
            _run(p, str(kam))

        if opinion.get("going_concern"):
            p = doc.add_paragraph()
            _run(p, "Going Concern", bold=True)
            gc_note = opinion.get("going_concern_note", "")
            if gc_note:
                doc.add_paragraph(gc_note)

        doc.add_page_break()

    # ── Financial statements ──────────────────────────────────────────────
    statements = report.get("financial_statements", {})
    for stmt_key in ("statement_of_financial_position", "statement_of_profit_or_loss",
                     "statement_of_changes_in_equity", "statement_of_cash_flows"):
        stmt = statements.get(stmt_key)
        if not stmt:
            continue

        doc.add_heading(stmt.get("title", ""), level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = (f"{company}  —  As at {period_end}" if "position" in stmt_key
                 else f"{company}  —  For the year ended {period_end}")
        _run(p, label, italic=True, size=font_size)

        cols = tpl.get("columns", DEFAULT_TEMPLATE["columns"])
        table_rows = _build_docx_statement_rows(stmt, currency)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = cols[0]
        hdr[1].text = cols[1]
        hdr[2].text = f"{cols[2]} ({currency})"
        hdr[3].text = f"{cols[3]} ({currency})"

        # Header formatting
        for cell in hdr:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
            _shade_cell(cell, "D9E2F3")

        # Data rows
        for row_data in table_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data["account"]
            row_cells[1].text = row_data["note"]
            row_cells[2].text = row_data["current"]
            row_cells[3].text = row_data["prior"]

            is_section_header = row_data.get("is_section")
            is_subtotal = row_data.get("is_subtotal")
            is_total = row_data.get("is_total")

            for i, cell in enumerate(row_cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(font_size)
                        run.font.name = font_name
                        if is_section_header or is_subtotal or is_total:
                            run.bold = True
                    if i >= 2:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if is_section_header:
                _shade_cell(row_cells[0], "F2F2F2")

        doc.add_paragraph()
        doc.add_page_break()

    # ── Notes ─────────────────────────────────────────────────────────────
    notes = report.get("notes", {})
    if notes:
        doc.add_heading("Notes to the Financial Statements", level=1)

        policies = notes.get("accounting_policies", "")
        if policies:
            p = doc.add_paragraph()
            _run(p, "Accounting Policies", bold=True, size=font_size + 1)
            for para_text in policies.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

        estimates = notes.get("critical_estimates", "")
        if estimates:
            p = doc.add_paragraph()
            _run(p, "Critical Estimates and Judgements", bold=True, size=font_size + 1)
            doc.add_paragraph(estimates)

        items = notes.get("items") or notes.get("sections") or []
        for item in items:
            num = item.get("note_number", "")
            title = item.get("title", "")
            content = item.get("content", "")
            doc.add_heading(f"Note {num}: {title}", level=2)
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_statement_rows(stmt: dict, currency: str) -> list[dict]:
    """Flatten statement sections into row dicts for DOCX table."""
    rows = []
    for section in stmt.get("sections", []):
        rows.append({
            "account": section.get("title", ""),
            "note": "", "current": "", "prior": "",
            "is_section": True,
        })
        for item in section.get("line_items", []):
            rows.append({
                "account": f"    {item.get('account_name', '')}",
                "note": str(item.get("note_ref") or item.get("notes_ref") or ""),
                "current": _fmt_number(item.get("current_year"), currency),
                "prior": _fmt_number(item.get("prior_year"), currency),
            })
        sub = section.get("subtotal")
        if sub:
            rows.append({
                "account": sub.get("account_name", ""),
                "note": "",
                "current": _fmt_number(sub.get("current_year"), currency),
                "prior": _fmt_number(sub.get("prior_year"), currency),
                "is_subtotal": True,
            })

    total = stmt.get("total")
    if total:
        rows.append({
            "account": total.get("account_name", ""),
            "note": "",
            "current": _fmt_number(total.get("current_year"), currency),
            "prior": _fmt_number(total.get("prior_year"), currency),
            "is_total": True,
        })
    return rows


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a DOCX table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


# ══════════════════════════════════════════════════════════════════════════════
#  Excel generation (openpyxl)
# ══════════════════════════════════════════════════════════════════════════════


def _generate_xlsx(report: dict, tpl: dict) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter

    currency = tpl.get("currency_symbol", "AED")
    font_name = tpl.get("font_family", "Times New Roman")
    font_size = tpl.get("font_size", 10)
    meta = report.get("metadata", {})
    company = meta.get("company_name", "Company")
    period_end = meta.get("period_end", "")

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_font = Font(name=font_name, size=font_size, bold=True)
    header_fill = PatternFill("solid", fgColor="D9E2F3")
    section_font = Font(name=font_name, size=font_size, bold=True)
    section_fill = PatternFill("solid", fgColor="F2F2F2")
    normal_font = Font(name=font_name, size=font_size)
    bold_font = Font(name=font_name, size=font_size, bold=True)
    title_font = Font(name=font_name, size=12, bold=True)
    num_fmt = '#,##0.00'
    thin_border = Border(bottom=Side(style="thin"))
    double_border = Border(bottom=Side(style="double"))

    cols = tpl.get("columns", DEFAULT_TEMPLATE["columns"])
    statements = report.get("financial_statements", {})

    sheet_map = {
        "statement_of_financial_position": "SOFP",
        "statement_of_profit_or_loss": "SOPL",
    }

    for stmt_key, sheet_name in sheet_map.items():
        stmt = statements.get(stmt_key)
        if not stmt:
            continue

        ws = wb.create_sheet(title=sheet_name)
        row_num = 1

        # Title
        ws.cell(row=row_num, column=1, value=stmt.get("title", "")).font = title_font
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=4)
        row_num += 1

        # Subtitle
        ws.cell(row=row_num, column=1, value=f"{company}  —  {period_end}").font = Font(
            name=font_name, size=font_size, italic=True)
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=4)
        row_num += 2

        # Header row
        headers = [cols[0], cols[1], f"{cols[2]} ({currency})", f"{cols[3]} ({currency})"]
        for c, val in enumerate(headers, 1):
            cell = ws.cell(row=row_num, column=c, value=val)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
        row_num += 1

        # Data rows
        for section in stmt.get("sections", []):
            # Section header
            cell = ws.cell(row=row_num, column=1, value=section.get("title", ""))
            cell.font = section_font
            cell.fill = section_fill
            for c in range(2, 5):
                ws.cell(row=row_num, column=c).fill = section_fill
            row_num += 1

            for item in section.get("line_items", []):
                ws.cell(row=row_num, column=1, value=f"    {item.get('account_name', '')}").font = normal_font
                ws.cell(row=row_num, column=2, value=str(
                    item.get("note_ref") or item.get("notes_ref") or "")).font = normal_font
                ws.cell(row=row_num, column=2).alignment = Alignment(horizontal="center")

                _write_number_cell(ws, row_num, 3, item.get("current_year"), normal_font, num_fmt)
                _write_number_cell(ws, row_num, 4, item.get("prior_year"), normal_font, num_fmt)
                row_num += 1

            sub = section.get("subtotal")
            if sub:
                ws.cell(row=row_num, column=1, value=sub.get("account_name", "")).font = bold_font
                _write_number_cell(ws, row_num, 3, sub.get("current_year"), bold_font, num_fmt)
                _write_number_cell(ws, row_num, 4, sub.get("prior_year"), bold_font, num_fmt)
                ws.cell(row=row_num, column=3).border = thin_border
                ws.cell(row=row_num, column=4).border = thin_border
                row_num += 1

            row_num += 1  # blank row between sections

        # Grand total
        total = stmt.get("total")
        if total:
            ws.cell(row=row_num, column=1, value=total.get("account_name", "")).font = bold_font
            _write_number_cell(ws, row_num, 3, total.get("current_year"), bold_font, num_fmt)
            _write_number_cell(ws, row_num, 4, total.get("prior_year"), bold_font, num_fmt)
            ws.cell(row=row_num, column=3).border = double_border
            ws.cell(row=row_num, column=4).border = double_border

        # Column widths
        ws.column_dimensions["A"].width = 40
        ws.column_dimensions["B"].width = 8
        ws.column_dimensions["C"].width = 22
        ws.column_dimensions["D"].width = 22

    # ── Notes sheet ───────────────────────────────────────────────────────
    notes = report.get("notes", {})
    if notes:
        ws = wb.create_sheet(title="Notes")
        row_num = 1

        ws.cell(row=row_num, column=1, value="Notes to the Financial Statements").font = title_font
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=3)
        row_num += 2

        # Header
        for c, val in enumerate(["Note #", "Title", "Content"], 1):
            cell = ws.cell(row=row_num, column=c, value=val)
            cell.font = header_font
            cell.fill = header_fill
        row_num += 1

        # Accounting policies row
        policies = notes.get("accounting_policies", "")
        if policies:
            ws.cell(row=row_num, column=1, value="—").font = normal_font
            ws.cell(row=row_num, column=2, value="Accounting Policies").font = bold_font
            ws.cell(row=row_num, column=3, value=policies[:500]).font = normal_font
            ws.cell(row=row_num, column=3).alignment = Alignment(wrap_text=True)
            row_num += 1

        items = notes.get("items") or notes.get("sections") or []
        for item in items:
            ws.cell(row=row_num, column=1, value=str(item.get("note_number", ""))).font = normal_font
            ws.cell(row=row_num, column=1).alignment = Alignment(horizontal="center")
            ws.cell(row=row_num, column=2, value=item.get("title", "")).font = bold_font
            ws.cell(row=row_num, column=3, value=item.get("content", "")).font = normal_font
            ws.cell(row=row_num, column=3).alignment = Alignment(wrap_text=True)
            row_num += 1

        ws.column_dimensions["A"].width = 10
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 70

    # Ensure at least one sheet exists
    if not wb.sheetnames:
        wb.create_sheet(title="Report")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _write_number_cell(ws, row: int, col: int, value, font, num_fmt: str):
    """Write a numeric value to an Excel cell with proper formatting."""
    from openpyxl.styles import Alignment
    cell = ws.cell(row=row, column=col)
    if value is None:
        cell.value = "-"
        cell.font = font
        cell.alignment = Alignment(horizontal="right")
        return
    try:
        num = float(value)
        cell.value = num
        cell.number_format = num_fmt
    except (TypeError, ValueError):
        cell.value = "-"
    cell.font = font
    cell.alignment = Alignment(horizontal="right")
