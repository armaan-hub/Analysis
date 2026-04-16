"""
Template Report Generator — generates a DOCX audit report by cloning the
prior year template's structure and filling it with current year data.

The output matches the prior year layout page-to-page, section-to-section.
Uses the template dict produced by DocumentFormatAnalyzer (Task 2) and
classified rows from AccountPlacementEngine (Task 3).
"""
from __future__ import annotations

import io
import logging
import re
from collections import OrderedDict
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ── Default formatting constants ──────────────────────────────────────────────

_DEFAULT_CURRENCY_FORMAT = "#,##0"
_DEFAULT_NEGATIVE_FORMAT = "(X,XXX)"

_FALLBACK_SECTIONS = [
    {"title": "Independent Auditors' Report", "level": 1, "content_type": "narrative"},
    {"title": "Statement of Financial Position", "level": 1, "content_type": "table"},
    {"title": "Statement of Profit or Loss and Other Comprehensive Income", "level": 1, "content_type": "table"},
    {"title": "Statement of Changes in Shareholders' Equity", "level": 1, "content_type": "table"},
    {"title": "Statement of Cash Flows", "level": 1, "content_type": "table"},
    {"title": "Notes to the Financial Statements", "level": 1, "content_type": "narrative"},
]


# ═══════════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════════

def generate_from_template(
    current_data: dict,
    template: dict | None,
    prior_year_rows: list[dict] | None = None,
) -> bytes:
    """
    Generate a DOCX audit report by cloning prior year template structure.

    Args:
        current_data: dict with company_name, location, period_end,
            opinion_type, draft_content, rows (list of classified accounts).
        template: dict from DocumentFormatAnalyzer with keys
            document_structure, account_grouping, terminology, formatting_rules.
            May be None/empty for graceful fallback.
        prior_year_rows: optional list of {account_name, prior_year_value}.

    Returns:
        bytes — DOCX file content.
    """
    template = template or {}

    doc_structure = template.get("document_structure") or {}
    account_grouping = template.get("account_grouping") or {}
    terminology = template.get("terminology") or {}
    formatting_rules = template.get("formatting_rules") or {}

    fmt = _resolve_formatting(formatting_rules)

    # Determine amount format from template
    fmt_rules = (template or {}).get("formatting_rules", {})
    table_fmt = fmt_rules.get("table_formatting", {})
    currency_format = table_fmt.get("currency_format", "#,##0")
    negative_format = table_fmt.get("negative_number_format", "(X,XXX)")

    def fmt_amount(value: float) -> str:
        """Format a numeric amount according to the extracted template rules."""
        if currency_format == "#,##0.00":
            formatted = f"{abs(value):,.2f}"
        else:
            formatted = f"{abs(value):,.0f}"
        if value < 0:
            if "(X,XXX)" in negative_format:
                return f"({formatted})"
            else:
                return f"-{formatted}"
        return formatted

    sections = doc_structure.get("sections") or _FALLBACK_SECTIONS
    page_break_sections = set(
        formatting_rules.get("page_break_after_sections", [])
    )

    company_name = current_data.get("company_name", "Company")
    location = current_data.get("location", "")
    period_end = current_data.get("period_end", "")
    draft_content = current_data.get("draft_content", "")
    rows = current_data.get("rows", [])

    # Build prior-year lookup (normalised account name → value)
    prior_lookup = _build_prior_year_lookup(prior_year_rows)

    # ── Create document ───────────────────────────────────────────
    doc = Document()
    _set_page_margins(doc)

    # 1. Cover page
    _build_cover_page(doc, company_name, location, period_end, terminology)

    # 2. Table of contents
    _build_toc(doc, sections)

    # 3–5. Iterate sections in template order
    for section_info in sections:
        title = section_info.get("title", "")
        content_type = section_info.get("content_type", "heading")

        if content_type == "narrative":
            _build_narrative_section(doc, title, draft_content)
        elif content_type == "table":
            section_rows = _collect_section_rows(title, rows, account_grouping)
            _build_financial_table(
                doc, title, section_rows, fmt, prior_lookup, fmt_amount,
            )
        else:
            # Generic heading placeholder
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Page break if template says so
        if title in page_break_sections:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Formatting helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _resolve_formatting(formatting_rules: dict) -> dict:
    """Merge template formatting rules with sensible defaults."""
    table_fmt = formatting_rules.get("table_formatting") or {}
    font_hier = formatting_rules.get("font_hierarchy") or {}

    currency_format = table_fmt.get("currency_format", _DEFAULT_CURRENCY_FORMAT)
    negative_format = table_fmt.get("negative_number_format", _DEFAULT_NEGATIVE_FORMAT)

    return {
        "currency_format": currency_format,
        "negative_format": negative_format,
        "use_decimals": ".00" in currency_format,
        "use_parenthetical": "(" in negative_format,
        "heading_1_bold": font_hier.get("heading_1_bold", True),
        "table_header_bold": font_hier.get("table_header_bold", True),
    }


def _format_number(value: float | None, fmt: dict) -> str:
    """Format a numeric value according to the resolved template rules."""
    if value is None:
        return "-"

    is_negative = value < 0
    abs_val = abs(value)

    if fmt["use_decimals"]:
        formatted = f"{abs_val:,.2f}"
    else:
        formatted = f"{abs_val:,.0f}"

    if is_negative:
        if fmt["use_parenthetical"]:
            return f"({formatted})"
        else:
            return f"-{formatted}"
    return formatted


# ═══════════════════════════════════════════════════════════════════════════════
# Run / paragraph helpers (mirrors audit_formatter.py patterns)
# ═══════════════════════════════════════════════════════════════════════════════

def _add_run(
    para, text: str, bold: bool = False, size_pt: int = 11,
    color: Optional[tuple] = None, italic: bool = False,
):
    """Append a run to a paragraph with formatting."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _set_page_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


# ═══════════════════════════════════════════════════════════════════════════════
# Cover page
# ═══════════════════════════════════════════════════════════════════════════════

def _build_cover_page(
    doc: Document,
    company_name: str,
    location: str,
    period_end: str,
    terminology: dict,
):
    """Insert cover page, using template's terminology for title text."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_company, company_name.upper(), bold=True, size_pt=16)

    if location:
        p_loc = doc.add_paragraph()
        p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_loc, location.upper(), bold=True, size_pt=13)

    doc.add_paragraph()
    doc.add_paragraph()

    # Use terminology from template if available
    headings_seen = terminology.get("headings_seen", [])
    title_line_1 = "FINANCIAL STATEMENTS AND"
    title_line_2 = "INDEPENDENT AUDITOR'S REPORT"
    for h in headings_seen:
        h_upper = h.upper()
        if "FINANCIAL" in h_upper and "STATEMENT" in h_upper:
            title_line_1 = h_upper
            title_line_2 = ""
            break

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, title_line_1, bold=True, size_pt=13)

    if title_line_2:
        p_title2 = doc.add_paragraph()
        p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_title2, title_line_2, bold=True, size_pt=13)

    doc.add_paragraph()

    if period_end:
        p_period = doc.add_paragraph()
        p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(
            p_period, f"FOR THE YEAR ENDED {period_end.upper()}",
            bold=True, size_pt=13,
        )

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# Table of contents
# ═══════════════════════════════════════════════════════════════════════════════

def _build_toc(doc: Document, sections: list[dict]):
    """Build table of contents from template's section list."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(heading, "Table of Contents", bold=False, size_pt=12)
    doc.add_paragraph()

    page_counter = 1
    for sec in sections:
        title = sec.get("title", "")
        if not title:
            continue
        start_page = sec.get("start_page", page_counter)
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        level = sec.get("level", 1)
        indent = "    " * (level - 1)
        run_title = p.add_run(f"{indent}{title}")
        run_title.font.size = Pt(11)
        run_page = p.add_run(f"\t{start_page}")
        run_page.font.size = Pt(11)
        run_page.bold = True
        page_counter += 1

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# Narrative sections
# ═══════════════════════════════════════════════════════════════════════════════

def _build_narrative_section(doc: Document, title: str, content: str):
    """Add a narrative section (e.g. auditor's report) from text/markdown."""
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not content:
        doc.add_paragraph()
        return

    cleaned = re.sub(r"^#{1,3}\s+", "", content, flags=re.MULTILINE)
    for para_text in cleaned.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", para_text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                _add_run(p, part[2:-2], bold=True, size_pt=11)
            else:
                _add_run(p, part, size_pt=11)


# ═══════════════════════════════════════════════════════════════════════════════
# Financial tables
# ═══════════════════════════════════════════════════════════════════════════════

def _build_financial_table(
    doc: Document,
    title: str,
    rows: list[dict],
    fmt: dict,
    prior_lookup: dict,
    fmt_amount=None,
):
    """
    Build a financial statement table with columns:
    Account Name | Notes | Current Year | Prior Year
    """
    doc.add_heading(title, level=2)

    has_prior = any(
        r.get("prior_year") is not None or _lookup_prior(r, prior_lookup) is not None
        for r in rows
    )
    col_count = 4 if has_prior else 3

    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # ── Dynamic column widths ─────────────────────────────────────
    _set_column_widths(table, col_count)

    # ── Header row ────────────────────────────────────────────────
    hdr = table.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = "Notes"
    hdr[2].text = "Current Year\nAED"
    if has_prior:
        hdr[3].text = "Prior Year\nAED"

    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = fmt["table_header_bold"]
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Data rows ─────────────────────────────────────────────────
    for row in rows:
        account_name = row.get("account_name", "")
        indent_level = row.get("indent_level", 1)
        is_total = row.get("is_total", False)
        is_subtotal = row.get("is_subtotal", False)
        notes_ref = row.get("notes_ref") or ""
        current_val = row.get("net")
        prior_val = row.get("prior_year")
        if prior_val is None:
            prior_val = _lookup_prior(row, prior_lookup)

        data_row = table.add_row()
        cells = data_row.cells

        # Account name with indentation
        prefix = "  " * max(0, indent_level - 1) if not (is_total or is_subtotal) else ""
        cells[0].text = f"{prefix}{account_name}"
        cells[1].text = str(notes_ref)
        if fmt_amount is not None and current_val is not None:
            cells[2].text = fmt_amount(current_val)
        else:
            cells[2].text = _format_number(current_val, fmt)
        if has_prior:
            if fmt_amount is not None and prior_val is not None:
                cells[3].text = fmt_amount(prior_val)
            else:
                cells[3].text = _format_number(prior_val, fmt)

        # Cell formatting
        for i, cell in enumerate(cells):
            for p in cell.paragraphs:
                for r in p.runs:
                    if is_total or is_subtotal:
                        r.bold = True
                    r.font.size = Pt(10)
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
                )

    doc.add_paragraph()


def _set_column_widths(table, col_count: int):
    """Set column widths for financial table (60/20/20 % split)."""
    if col_count == 4:
        widths = [Inches(3.0), Inches(0.6), Inches(1.2), Inches(1.2)]
    else:
        widths = [Inches(3.6), Inches(1.2), Inches(1.2)]

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = width
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width


# ═══════════════════════════════════════════════════════════════════════════════
# Row collection / lookup helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _normalise(name: str) -> str:
    """Lowercase, strip, collapse whitespace for matching."""
    return re.sub(r"\s+", " ", name.lower().strip())


def _build_prior_year_lookup(prior_year_rows: list[dict] | None) -> dict:
    """Build normalised account_name → prior_year_value lookup."""
    if not prior_year_rows:
        return {}
    lookup = {}
    for entry in prior_year_rows:
        name = entry.get("account_name", "")
        value = entry.get("prior_year_value")
        if name:
            lookup[_normalise(name)] = value
    return lookup


def _lookup_prior(row: dict, prior_lookup: dict) -> float | None:
    """Look up a row's prior year value from the external lookup."""
    if not prior_lookup:
        return None
    name = row.get("account_name", "")
    return prior_lookup.get(_normalise(name))


def _collect_section_rows(
    section_title: str,
    all_rows: list[dict],
    account_grouping: dict,
) -> list[dict]:
    """
    Collect rows for a given section.

    First tries matching rows by their ``section`` field against the
    section_title.  Falls back to looking up the template's
    account_grouping to find which accounts belong to this section.
    """
    norm_title = _normalise(section_title)

    # Direct match on row.section
    matched = [
        r for r in all_rows
        if _normalise(r.get("section", "")) == norm_title
    ]
    if matched:
        return matched

    # Fallback: check account_grouping keys
    template_accounts: set[str] = set()
    for key, accounts in account_grouping.items():
        if _normalise(key) == norm_title:
            for acct in accounts:
                template_accounts.add(_normalise(acct.get("account_name", "")))

    if template_accounts:
        return [
            r for r in all_rows
            if _normalise(r.get("account_name", "")) in template_accounts
        ]

    return []
