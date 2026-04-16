"""
Export Converter — converts markdown text to Word, PDF, or Excel.
Used by the chat export endpoint to let users download LLM responses.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


def _parse_markdown_tables(text: str) -> list[list[list[str]]]:
    """Extract all markdown tables from text."""
    tables = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.count("|") >= 2:
            header_cells = [c.strip() for c in line.split("|") if c.strip()]
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    row_line = lines[j].strip()
                    if not row_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in row_line.split("|") if c.strip()]
                    if row_cells:
                        rows.append(row_cells)
                    j += 1
                if len(rows) > 1:
                    tables.append(rows)
                i = j
                continue
        i += 1
    return tables


def to_word(markdown_text: str) -> bytes:
    """Convert markdown text to a Word DOCX file. Returns bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        stripped = line.strip()

        if stripped.startswith("|") and stripped.count("|") >= 2:
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                header_cells = [c.strip() for c in stripped.split("|") if c.strip()]
                table_rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    r_line = lines[j].strip()
                    if not r_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in r_line.split("|") if c.strip()]
                    if row_cells:
                        table_rows.append(row_cells)
                    j += 1

                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            if r_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                doc.add_paragraph()
                i = j
                continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(markdown_text: str) -> bytes:
    """Convert markdown text to PDF. Returns bytes."""
    try:
        return _to_pdf_weasyprint(markdown_text)
    except Exception as exc:
        logger.warning(f"weasyprint PDF failed ({exc}), trying reportlab")
        try:
            return _to_pdf_reportlab(markdown_text)
        except Exception as exc2:
            logger.error(f"Both PDF methods failed: {exc2}")
            raise RuntimeError("PDF export unavailable — install weasyprint or reportlab") from exc2


def _to_pdf_weasyprint(markdown_text: str) -> bytes:
    import markdown as md_lib
    from weasyprint import HTML, CSS

    html_body = md_lib.markdown(markdown_text, extensions=["tables", "fenced_code"])
    css = CSS(string="""
        @page { size: A4; margin: 2cm; }
        body { font-family: "Times New Roman", Times, serif; font-size: 11pt; line-height: 1.6; }
        h1 { font-size: 16pt; margin-top: 20pt; }
        h2 { font-size: 13pt; margin-top: 14pt; }
        table { border-collapse: collapse; width: 100%; margin: 10pt 0; }
        th, td { border: 1px solid #ccc; padding: 6pt 8pt; text-align: left; }
        th { background: #f0f0f0; font-weight: bold; }
    """)
    full_html = f"<html><body>{html_body}</body></html>"
    return HTML(string=full_html).write_pdf(stylesheets=[css])


def _to_pdf_reportlab(markdown_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Heading1"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            story.append(Paragraph(f"&bull; {stripped[2:]}", styles["Normal"]))
        else:
            clean = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", stripped)
            story.append(Paragraph(clean, styles["Normal"]))

    doc_rl.build(story)
    return buf.getvalue()


def to_excel(markdown_text: str) -> bytes:
    """Convert markdown tables to Excel. Returns empty bytes if no tables found."""
    tables = _parse_markdown_tables(markdown_text)
    if not tables:
        return b""

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    lines = markdown_text.split("\n")
    table_line_indices = []
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("|") and stripped.count("|") >= 2:
            if idx + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[idx + 1].strip()):
                table_line_indices.append(idx)

    for t_idx, table_rows in enumerate(tables):
        sheet_name = f"Table {t_idx + 1}"
        if t_idx < len(table_line_indices):
            tli = table_line_indices[t_idx]
            for li in range(tli - 1, max(tli - 10, -1), -1):
                if lines[li].strip().startswith("#"):
                    sheet_name = re.sub(r"^#+\s*", "", lines[li].strip())[:31]
                    break

        ws = wb.create_sheet(title=sheet_name)
        header_fill = PatternFill("solid", fgColor="E0E7FF")
        header_font = Font(bold=True)

        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_value in enumerate(row):
                cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=cell_value)
                if r_idx == 0:
                    cell.font = header_font
                    cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True)

                if r_idx > 0:
                    try:
                        numeric = float(cell_value.replace(",", "").replace("%", ""))
                        cell.value = numeric
                        if "%" in cell_value:
                            cell.number_format = "0.00%"
                        else:
                            cell.number_format = "#,##0.00"
                    except (ValueError, AttributeError):
                        pass

        for col in ws.columns:
            max_len = max(
                (len(str(c.value)) for c in col if c.value is not None),
                default=10,
            )
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
