"""
Format Extractor — parses uploaded template files and extracts
structured section information for LLM-guided report generation.
"""
from __future__ import annotations

import os
import re
import shutil
import tempfile
from typing import Any


def extract_format_structure(tmp_path: str, ext: str) -> dict[str, Any]:
    """
    Extract document structure from a template file.

    Returns:
        {
            "sections": [{"heading": str, "level": int, "sample_text": str}],
            "detected_format": str,   # "docx" | "pdf" | "xlsx" | "txt"
            "section_count": int,
            "raw_preview": str,       # first 500 chars of raw text
        }
    """
    sections: list[dict] = []
    raw_text = ""

    if ext == "txt":
        with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()
        sections = _extract_sections_from_text(raw_text)

    elif ext in ("docx", "doc"):
        try:
            import docx as _docx
            doc = _docx.Document(tmp_path)
            raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            sections = []
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                style_name = (para.style.name or "").lower()
                if "heading 1" in style_name or style_name.startswith("heading 1"):
                    sections.append({"heading": para.text.strip(), "level": 1, "sample_text": ""})
                elif "heading 2" in style_name:
                    sections.append({"heading": para.text.strip(), "level": 2, "sample_text": ""})
                elif "heading 3" in style_name:
                    sections.append({"heading": para.text.strip(), "level": 3, "sample_text": ""})
                elif sections:
                    # Append to last section's sample_text (up to 120 chars)
                    last = sections[-1]
                    if len(last["sample_text"]) < 120:
                        last["sample_text"] += " " + para.text.strip()
        except Exception:
            raw_text = raw_text or ""
            sections = _extract_sections_from_text(raw_text)

    elif ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(tmp_path)
            raw_text = "\n".join(page.get_text() for page in doc)
            sections = _extract_sections_from_text(raw_text)
        except ImportError:
            raw_text = ""
            sections = []

    elif ext in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"[Sheet: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    line = "\t".join(str(c) for c in row if c is not None)
                    if line.strip():
                        lines.append(line)
            raw_text = "\n".join(lines)
            # For Excel, treat sheet names as top-level sections
            sections = [{"heading": ws.title, "level": 1, "sample_text": ""} for ws in wb.worksheets]
        except Exception:
            raw_text = raw_text or ""
            sections = []

    return {
        "sections": sections,
        "detected_format": ext,
        "section_count": len(sections),
        "raw_preview": raw_text[:500],
    }


_HEADING_PATTERNS = [
    (re.compile(r"^#{1,3}\s+(.+)$", re.MULTILINE), "markdown"),
    (re.compile(r"^([A-Z][A-Z\s\-/]{3,}):?\s*$", re.MULTILINE), "allcaps"),
    (re.compile(r"^\d+[\.\)]\s+([A-Z].{2,60})$", re.MULTILINE), "numbered"),
    (re.compile(r"^[A-Z].{5,60}$", re.MULTILINE), "titlecase"),
]


def _extract_sections_from_text(text: str) -> list[dict]:
    """Heuristic section extraction from plain text."""
    if not text.strip():
        return []

    for pattern, style in _HEADING_PATTERNS[:3]:  # try structured patterns first
        matches = pattern.findall(text)
        if len(matches) >= 2:
            return [{"heading": m.strip(), "level": 1, "sample_text": ""} for m in matches[:20]]

    # Fallback: lines that look like headings (short, capitalised, no period)
    sections = []
    for line in text.splitlines():
        line = line.strip()
        if 4 < len(line) < 70 and not line.endswith(".") and line[0].isupper() and line == line.title():
            sections.append({"heading": line, "level": 1, "sample_text": ""})
        if len(sections) >= 20:
            break
    return sections
