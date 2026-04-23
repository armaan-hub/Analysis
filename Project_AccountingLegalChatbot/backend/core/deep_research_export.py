"""
Deep Research Export — branded PDF, DOCX, and XLSX exports for deep-research results.
Produces professional documents with cover page, headings, body, and sources appendix.
"""
from __future__ import annotations

import html
import io
import logging
import re
from datetime import datetime
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_markdown_tables(text: str) -> list[list[list[str]]]:
    """Extract all markdown tables from text (reuses pattern from export_converter)."""
    tables: list[list[list[str]]] = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.count("|") >= 2:
            header_cells = [c.strip() for c in line.split("|") if c.strip()]
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    row_line = lines[j].strip()
                    if not row_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in row_line.split("|") if c.strip()]
                    if row_cells:
                        rows.append(row_cells)
                    j += 1
                if len(rows) > 1:
                    tables.append(rows)
                i = j
                continue
        i += 1
    return tables


def _strip_markdown_bold(text: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


# ---------------------------------------------------------------------------
# PDF (ReportLab)
# ---------------------------------------------------------------------------

def to_branded_pdf(content: str, sources: list[dict[str, Any]], query: str) -> bytes:
    """Return branded PDF bytes with cover page, body, and sources appendix."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    # Custom styles
    title_style = ParagraphStyle(
        "CoverTitle", parent=styles["Title"],
        fontSize=24, leading=30,
        textColor=HexColor("#1a365d"),
        spaceAfter=20,
    )
    subtitle_style = ParagraphStyle(
        "CoverSubtitle", parent=styles["Normal"],
        fontSize=12, leading=16,
        textColor=HexColor("#4a5568"),
        spaceAfter=6,
    )
    heading1 = ParagraphStyle(
        "BrandH1", parent=styles["Heading1"],
        fontSize=16, leading=20,
        textColor=HexColor("#1a365d"),
        spaceBefore=16, spaceAfter=8,
    )
    heading2 = ParagraphStyle(
        "BrandH2", parent=styles["Heading2"],
        fontSize=13, leading=17,
        textColor=HexColor("#2d3748"),
        spaceBefore=12, spaceAfter=6,
    )
    heading3 = ParagraphStyle(
        "BrandH3", parent=styles["Heading3"],
        fontSize=11, leading=14,
        textColor=HexColor("#4a5568"),
        spaceBefore=8, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BrandBody", parent=styles["Normal"],
        fontSize=10, leading=14,
    )
    source_style = ParagraphStyle(
        "SourceItem", parent=styles["Normal"],
        fontSize=9, leading=12,
        textColor=HexColor("#4a5568"),
        leftIndent=12,
    )

    story: list[Any] = []

    # ── Cover page ──
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph("Deep Research Report", title_style))
    story.append(Spacer(1, 12))
    safe_query = query.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    story.append(Paragraph(f"<b>Query:</b> {safe_query}", subtitle_style))
    story.append(Paragraph(
        f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        subtitle_style,
    ))
    story.append(Paragraph(
        f"<b>Sources analysed:</b> {len(sources)}",
        subtitle_style,
    ))
    story.append(PageBreak())

    # ── Body ──
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue

        # Headings
        if stripped.startswith("### "):
            story.append(Paragraph(_strip_markdown_bold(stripped[4:]), heading3))
        elif stripped.startswith("## "):
            story.append(Paragraph(_strip_markdown_bold(stripped[3:]), heading2))
        elif stripped.startswith("# "):
            story.append(Paragraph(_strip_markdown_bold(stripped[2:]), heading1))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            clean = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", stripped[2:])
            story.append(Paragraph(f"&bull; {clean}", body_style))
        elif re.match(r"^\d+\.\s", stripped):
            clean = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>",
                           re.sub(r"^\d+\.\s", "", stripped))
            story.append(Paragraph(clean, body_style))
        else:
            clean = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", stripped)
            # Skip table separator lines
            if not re.match(r"^\|[\s\-\|:]+\|$", stripped):
                story.append(Paragraph(clean, body_style))

    # ── Sources appendix ──
    if sources:
        story.append(PageBreak())
        story.append(Paragraph("Sources &amp; References", heading1))
        story.append(Spacer(1, 8))
        for idx, src in enumerate(sources, 1):
            name = (src.get("source") or "Unknown").replace("&", "&amp;")
            excerpt = (src.get("excerpt") or "")[:200].replace("&", "&amp;").replace("<", "&lt;")
            page = src.get("page", "")
            page_str = f" (p. {page})" if page else ""
            score = src.get("score")
            score_str = f" — relevance {score:.0%}" if score is not None else ""
            story.append(Paragraph(
                f"<b>[{idx}]</b> {name}{page_str}{score_str}",
                source_style,
            ))
            url = src.get("url")
            if url:
                story.append(Paragraph(
                    f'<link href="{html.escape(url, quote=True)}" color="#1a365d">'
                    f'<u>{html.escape(url)}</u></link>',
                    source_style,
                ))
            if excerpt:
                story.append(Paragraph(f"<i>{excerpt}…</i>", source_style))
            story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a clickable hyperlink to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1A365D")
    rPr.append(color_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)
    new_run.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    new_run.append(t_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def to_branded_docx(content: str, sources: list[dict[str, Any]], query: str) -> bytes:
    """Return branded DOCX bytes with cover page, body, and sources appendix."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover page ──
    for _ in range(6):
        doc.add_paragraph()  # spacing

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Deep Research Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"Query: {query}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    src_para = doc.add_paragraph()
    src_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src_para.add_run(f"Sources analysed: {len(sources)}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.add_page_break()

    # ── Body ──
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(_strip_markdown_bold(stripped[4:]), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(_strip_markdown_bold(stripped[3:]), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(_strip_markdown_bold(stripped[2:]), level=1)
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and stripped.count("|") >= 2:
            if i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1].strip()):
                header_cells = [c.strip() for c in stripped.split("|") if c.strip()]
                table_rows = [header_cells]
                j = i + 2
                while j < len(lines):
                    r_line = lines[j].strip()
                    if not r_line.startswith("|"):
                        break
                    row_cells = [c.strip() for c in r_line.split("|") if c.strip()]
                    if row_cells:
                        table_rows.append(row_cells)
                    j += 1

                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            if r_idx == 0:
                                for para in cell.paragraphs:
                                    for r in para.runs:
                                        r.bold = True
                doc.add_paragraph()
                i = j
                continue

        # Lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            i += 1
            continue

        # Normal paragraph with bold handling
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1

    # ── Sources appendix ──
    if sources:
        doc.add_page_break()
        doc.add_heading("Sources & References", level=1)
        for idx, src in enumerate(sources, 1):
            name = src.get("source") or "Unknown"
            excerpt = (src.get("excerpt") or "")[:200]
            page = src.get("page", "")
            page_str = f" (p. {page})" if page else ""
            score = src.get("score")
            score_str = f" — relevance {score:.0%}" if score is not None else ""

            p = doc.add_paragraph()
            run = p.add_run(f"[{idx}] {name}{page_str}{score_str}")
            run.bold = True
            run.font.size = Pt(10)
            url = src.get("url")
            if url:
                p_url = doc.add_paragraph()
                _add_hyperlink(p_url, url, url)
            if excerpt:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f"{excerpt}…")
                run2.italic = True
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# XLSX (openpyxl)
# ---------------------------------------------------------------------------

def to_branded_xlsx(content: str, sources: list[dict[str, Any]], query: str) -> bytes:
    """Return branded XLSX bytes. Tables → separate sheets; always includes sources sheet."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="1A365D")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    body_font = Font(size=10)
    wrap_align = Alignment(wrap_text=True, vertical="top")

    tables = _parse_markdown_tables(content)
    lines = content.split("\n")

    if tables:
        # Find heading above each table for sheet naming
        table_line_indices: list[int] = []
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("|") and stripped.count("|") >= 2:
                if idx + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[idx + 1].strip()):
                    table_line_indices.append(idx)

        for t_idx, table_rows in enumerate(tables):
            sheet_name = f"Table {t_idx + 1}"
            if t_idx < len(table_line_indices):
                tli = table_line_indices[t_idx]
                for li in range(tli - 1, max(tli - 10, -1), -1):
                    if lines[li].strip().startswith("#"):
                        sheet_name = re.sub(r"^#+\s*", "", lines[li].strip())[:31]
                        break

            ws = wb.create_sheet(title=sheet_name)
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_value in enumerate(row):
                    cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=cell_value)
                    cell.alignment = wrap_align
                    if r_idx == 0:
                        cell.font = header_font
                        cell.fill = header_fill
                    else:
                        cell.font = body_font
                        try:
                            numeric = float(cell_value.replace(",", "").replace("%", ""))
                            if "%" in cell_value:
                                cell.value = numeric / 100
                                cell.number_format = "0.00%"
                            else:
                                cell.number_format = "#,##0.00"
                        except (ValueError, AttributeError):
                            pass

            for col in ws.columns:
                max_len = max(
                    (len(str(c.value)) for c in col if c.value is not None),
                    default=10,
                )
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)
    else:
        # No tables — put full text in column A
        ws = wb.create_sheet(title="Research Report")
        ws.cell(row=1, column=1, value="Deep Research Report")
        ws.cell(row=1, column=1).font = Font(bold=True, size=14)
        ws.cell(row=2, column=1, value=f"Query: {query}")
        ws.cell(row=3, column=1, value=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        ws.cell(row=4, column=1, value="")

        row_num = 5
        for line in lines:
            ws.cell(row=row_num, column=1, value=line)
            ws.cell(row=row_num, column=1).alignment = wrap_align
            row_num += 1

        ws.column_dimensions["A"].width = 120

    # ── Sources sheet ──
    if sources:
        ws_src = wb.create_sheet(title="Sources")
        headers = ["#", "Source", "Page", "Relevance", "Excerpt"]
        for c_idx, h in enumerate(headers, 1):
            cell = ws_src.cell(row=1, column=c_idx, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = wrap_align

        for s_idx, src in enumerate(sources, 1):
            ws_src.cell(row=s_idx + 1, column=1, value=s_idx).font = body_font
            ws_src.cell(row=s_idx + 1, column=2, value=src.get("source", "")).font = body_font
            ws_src.cell(row=s_idx + 1, column=3, value=src.get("page", "")).font = body_font
            score = src.get("score")
            if score is not None:
                cell = ws_src.cell(row=s_idx + 1, column=4, value=score)
                cell.number_format = "0.00%"
                cell.font = body_font
            ws_src.cell(row=s_idx + 1, column=5, value=src.get("excerpt", "")).font = body_font
            ws_src.cell(row=s_idx + 1, column=5).alignment = wrap_align

        col_widths = [5, 30, 8, 12, 60]
        for c_idx, w in enumerate(col_widths, 1):
            ws_src.column_dimensions[get_column_letter(c_idx)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
